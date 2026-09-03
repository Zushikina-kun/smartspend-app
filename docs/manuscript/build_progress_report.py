"""
build_progress_report.py
Generates Lorma BSIT Capstone Progress Report DOCX files
matching the exact template format from Capstone_Progress_Report_BSIT.docx

Usage:
    python build_progress_report.py week4
    python build_progress_report.py week5
    python build_progress_report.py all

Output:
    docs/manuscript/SmartSpend_Progress_Report_Week4.docx
    docs/manuscript/SmartSpend_Progress_Report_Week5.docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Lorma template colours ────────────────────────────────────────────────────
MAROON_DARK  = RGBColor(0x5C, 0x0E, 0x24)  # title banner + section headings
MAROON_LOGO  = RGBColor(0x7A, 0x12, 0x30)  # "LORMA COLLEGES"
DARK_TEXT    = RGBColor(0x22, 0x22, 0x22)  # body / adviser name
GREY_TEXT    = RGBColor(0x6B, 0x6B, 0x6B)  # subtitles, prompts, labels
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CREAM_BG     = "F2EFEC"   # instructions box background (hex string)
MAROON_HEX   = "5C0E24"   # banner background

FONT = "Tahoma"

# ── Helper utilities ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_col_width(cell, width_twips: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def add_para(cell_or_doc, text: str, bold=False, size_pt=9,
             color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=0) -> None:
    """Add a paragraph to a cell or document with given formatting."""
    if hasattr(cell_or_doc, 'paragraphs') and hasattr(cell_or_doc, '_body'):
        # It's a cell
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def set_table_borders(table, border_color="BFBFBF", border_size=4):
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def clear_table_borders(table):
    """Remove borders from a table (for header/logo table)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_table_width(table, width_twips: int):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove any existing tblW elements first to avoid duplicates
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ── Main builder ──────────────────────────────────────────────────────────────

def build_report(data: dict, output_path: Path):
    doc = Document()

    # ── Page setup: Letter, 0.7" margins ─────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    margin = Inches(0.7)
    section.top_margin    = margin
    section.bottom_margin = margin
    section.left_margin   = margin
    section.right_margin  = margin

    # Remove default paragraph spacing
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(9)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(2)

    # ── HEADER TABLE (3-col, borderless) ─────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=3)
    set_table_width(hdr, 10224)
    clear_table_borders(hdr)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left logo placeholder
    lc = hdr.cell(0, 0)
    set_col_width(lc, 2100)
    p = lc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Logo]')
    run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = GREY_TEXT

    # Center — school name block
    cc = hdr.cell(0, 1)
    set_col_width(cc, 6024)
    cc.paragraphs[0].clear()
    for txt, bold, sz, col in [
        ('LORMA COLLEGES',                               True,  15, MAROON_LOGO),
        ('College of Computer Studies and Engineering',  False, 10, DARK_TEXT),
        ('Bachelor of Science in Information Technology',False,  9, GREY_TEXT),
    ]:
        p = cc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(txt)
        r.font.name = FONT; r.font.size = Pt(sz); r.bold = bold; r.font.color.rgb = col

    # Right logo placeholder
    rc = hdr.cell(0, 2)
    set_col_width(rc, 2100)
    p = rc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Logo]')
    run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = GREY_TEXT

    # ── Spacer ────────────────────────────────────────────────────────────────
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(0)

    # ── TITLE BANNER ─────────────────────────────────────────────────────────
    banner = doc.add_table(rows=1, cols=1)
    set_table_width(banner, 10224)
    clear_table_borders(banner)
    bc = banner.cell(0, 0)
    set_cell_bg(bc, MAROON_HEX)
    p = bc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('CAPSTONE PROJECT PROGRESS REPORT')
    r.font.name = FONT; r.font.size = Pt(13); r.bold = True; r.font.color.rgb = WHITE

    # ── Spacer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()

    # ── GENERAL INSTRUCTIONS BOX ─────────────────────────────────────────────
    ibox = doc.add_table(rows=1, cols=1)
    set_table_width(ibox, 10224)
    set_table_borders(ibox, "CCCCCC", 4)
    ic = ibox.cell(0, 0)
    set_cell_bg(ic, CREAM_BG)
    p = ic.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('GENERAL INSTRUCTIONS')
    r.font.name = FONT; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = DARK_TEXT
    instructions = [
        'Use this format for all progress report submissions.',
        'Encode the report clearly and concisely.',
        'Submit on or before the deadline set by the adviser.',
        'Attach photo documentation as required.',
    ]
    for inst in instructions:
        p2 = ic.add_paragraph(style='List Bullet')
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(1)
        p2.paragraph_format.left_indent  = Inches(0.25)
        r2 = p2.add_run(inst)
        r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # ── SECTION: Report Details ───────────────────────────────────────────────
    _section_label(doc, 'Report Details')

    details = doc.add_table(rows=3, cols=2)
    set_table_width(details, 10224)
    set_table_borders(details, "AAAAAA", 4)
    details.alignment = WD_TABLE_ALIGNMENT.LEFT
    label_w, val_w = 2200, 8024
    rows_data = [
        ('Project / Capstone Title', data['project_title']),
        ('Report No.',               data['report_no']),
        ('Date Submitted',           data['date_submitted']),
    ]
    for i, (lbl, val) in enumerate(rows_data):
        lc2 = details.cell(i, 0)
        vc  = details.cell(i, 1)
        set_col_width(lc2, label_w)
        set_col_width(vc,  val_w)
        lc2.paragraphs[0].clear()
        p = lc2.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(lbl)
        r.font.name = FONT; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = DARK_TEXT
        vc.paragraphs[0].clear()
        p2 = vc.paragraphs[0]
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run(val)
        r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # ── SECTION: Members ─────────────────────────────────────────────────────
    _section_label(doc, 'Members')

    members_tbl = doc.add_table(rows=len(data['members']) + 1, cols=2)
    set_table_width(members_tbl, 10224)
    set_table_borders(members_tbl, "AAAAAA", 4)
    num_w, name_w = 800, 9424

    # Header row
    for col_i, htext in enumerate(['#', 'Member Name']):
        c = members_tbl.cell(0, col_i)
        set_col_width(c, num_w if col_i == 0 else name_w)
        set_cell_bg(c, "E8E0DA")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(htext)
        r.font.name = FONT; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = DARK_TEXT

    for i, member in enumerate(data['members'], start=1):
        nc = members_tbl.cell(i, 0)
        mc = members_tbl.cell(i, 1)
        set_col_width(nc, num_w)
        set_col_width(mc, name_w)
        for c2, txt in [(nc, str(i)), (mc, member)]:
            p = c2.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(txt)
            r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = DARK_TEXT

    doc.add_paragraph()

    # ── SECTION: Narrative ────────────────────────────────────────────────────
    _section_label(doc, 'Narrative')

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(0)
    p_intro.paragraph_format.space_after  = Pt(4)
    r = p_intro.add_run('Describe the project\'s current status. Write each part in paragraph form.')
    r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = GREY_TEXT

    narrative_sections = [
        ('1. Tasks Already Completed',
         'What has the team finished since the last report?',
         data['completed']),
        ('2. Ongoing Activities',
         'What is currently in progress?',
         data['ongoing']),
        ('3. Problems or Challenges Encountered',
         'What issues or roadblocks came up?',
         data['problems']),
        ('4. Solutions Applied or Actions Taken',
         'How were the above issues addressed?',
         data['solutions']),
        ('5. Next Steps / Plans Before the Next Report',
         'What will the team do next?',
         data['next_steps']),
    ]

    for heading, prompt, content in narrative_sections:
        # Subsection heading
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(6)
        ph.paragraph_format.space_after  = Pt(1)
        rh = ph.add_run(heading)
        rh.font.name = FONT; rh.font.size = Pt(10.5); rh.bold = True
        rh.font.color.rgb = MAROON_DARK

        # Prompt question
        pp2 = doc.add_paragraph()
        pp2.paragraph_format.space_before = Pt(0)
        pp2.paragraph_format.space_after  = Pt(3)
        rp = pp2.add_run(prompt)
        rp.font.name = FONT; rp.font.size = Pt(8.5); rp.font.color.rgb = GREY_TEXT

        # Content box (single-cell bordered table)
        ctbl = doc.add_table(rows=1, cols=1)
        set_table_width(ctbl, 10224)
        set_table_borders(ctbl, "AAAAAA", 4)
        cc2 = ctbl.cell(0, 0)
        cc2.paragraphs[0].clear()

        # Write content paragraphs
        lines = content.strip().split('\n')
        first = True
        for line in lines:
            if first:
                p2 = cc2.paragraphs[0]
                first = False
            else:
                p2 = cc2.add_paragraph()
            p2.paragraph_format.space_before = Pt(3)
            p2.paragraph_format.space_after  = Pt(3)
            r2 = p2.add_run(line)
            r2.font.name = FONT; r2.font.size = Pt(9.5); r2.font.color.rgb = DARK_TEXT

        doc.add_paragraph()

    # ── SECTION: Photo Documentation ──────────────────────────────────────────
    _section_label(doc, 'Photo Documentation')

    p_photo_inst = doc.add_paragraph()
    p_photo_inst.paragraph_format.space_before = Pt(0)
    p_photo_inst.paragraph_format.space_after  = Pt(4)
    r = p_photo_inst.add_run(
        'Attach clear photos showing actual progress of the project. '
        'Each photo should have a short caption describing what is shown. '
        'Add more rows if needed.')
    r.font.name = FONT; r.font.size = Pt(9); r.font.color.rgb = GREY_TEXT

    photos = doc.add_table(rows=6, cols=2)
    set_table_width(photos, 10224)
    set_table_borders(photos, "AAAAAA", 4)
    col_w = 5112
    photo_labels = [
        ('[ Insert Photo 1 Here ]', '[ Insert Photo 2 Here ]'),
        ('Caption: ',               'Caption: '),
        ('[ Insert Photo 3 Here ]', '[ Insert Photo 4 Here ]'),
        ('Caption: ',               'Caption: '),
        ('[ Insert Photo 5 Here ]', '[ Insert Photo 6 Here ]'),
        ('Caption: ',               'Caption: '),
    ]
    for row_i, (left_txt, right_txt) in enumerate(photo_labels):
        for col_i, txt in [(0, left_txt), (1, right_txt)]:
            c = photos.cell(row_i, col_i)
            set_col_width(c, col_w)
            is_placeholder = 'Insert' in txt
            min_h = 36 if is_placeholder else 12
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_placeholder else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(min_h)
            p.paragraph_format.space_after  = Pt(min_h)
            r = p.add_run(txt)
            r.font.name = FONT
            r.font.size = Pt(8.5 if is_placeholder else 9)
            r.font.color.rgb = GREY_TEXT if is_placeholder else DARK_TEXT

    doc.add_paragraph()

    # ── SIGNATURE BLOCK ───────────────────────────────────────────────────────
    sig = doc.add_table(rows=3, cols=3)
    set_table_width(sig, 10224)
    clear_table_borders(sig)
    left_w, mid_w, right_w = 4912, 400, 4912

    for row_i in range(3):
        for col_i in range(3):
            c = sig.cell(row_i, col_i)
            set_col_width(c, [left_w, mid_w, right_w][col_i])
            c.paragraphs[0].clear()

    # Top labels
    for col_i, txt, bold, sz, col_rgb in [
        (0, 'Group Leader / Representative', False, 8.5, GREY_TEXT),
        (2, data['adviser_name'],             True,  10,  DARK_TEXT),
    ]:
        p = sig.cell(0, col_i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(txt)
        r.font.name = FONT; r.font.size = Pt(sz); r.bold = bold; r.font.color.rgb = col_rgb

    # Signature space (middle row)
    for col_i in [0, 2]:
        p = sig.cell(1, col_i).paragraphs[0]
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(4)

    # Underline / role row
    for col_i, txt in [(0, 'Signature over Printed Name'), (2, 'Teacher in Charge')]:
        p = sig.cell(2, col_i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt)
        r.font.name = FONT; r.font.size = Pt(8.5); r.font.color.rgb = GREY_TEXT

    # ── Footer ───────────────────────────────────────────────────────────────
    footer = section.footer
    footer.paragraphs[0].clear()
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after  = Pt(0)

    # Left: institution text
    fr = fp.add_run(
        'Lorma Colleges \u2013 College of Computer Studies and Engineering '
        '| BS Information Technology Capstone Project')
    fr.font.name = FONT; fr.font.size = Pt(7); fr.font.color.rgb = GREY_TEXT

    # Tab to right side
    fp.add_run('\t')

    # "Page " text run
    fr_pg = fp.add_run('Page ')
    fr_pg.font.name = FONT; fr_pg.font.size = Pt(7); fr_pg.font.color.rgb = GREY_TEXT

    # PAGE field — correct OOXML: fldChar(begin) + instrText + fldChar(end)
    # each must be a child of its own <w:r>
    def _add_field_run(para, instr: str, rpr_xml: str = None):
        """Add a complete field (begin + instrText + end) as three <w:r> elements."""
        p_xml = para._p
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for fld_type, extra in [('begin', None), (None, instr), ('end', None)]:
            r_el = OxmlElement('w:r')
            if rpr_xml:
                rpr = OxmlElement('w:rPr')
                r_el.append(rpr)
            if fld_type is not None:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), fld_type)
                r_el.append(fc)
            else:
                it = OxmlElement('w:instrText')
                it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                it.text = extra
                r_el.append(it)
            p_xml.append(r_el)

    _add_field_run(fp, ' PAGE ')

    fr_of = fp.add_run(' of ')
    fr_of.font.name = FONT; fr_of.font.size = Pt(7); fr_of.font.color.rgb = GREY_TEXT

    _add_field_run(fp, ' NUMPAGES ')

    doc.save(output_path)
    print(f"✓ Saved: {output_path}")


def _section_label(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = MAROON_DARK


# ── Report data ───────────────────────────────────────────────────────────────

WEEK4 = {
    'project_title': (
        'SmartSpend: An AI-Assisted Multi-Modal Personal Financial Management '
        'Application for Filipino Users Using Agentic Large Language Model Architecture'
    ),
    'report_no':        '04',
    'date_submitted':   'September 3, 2026',
    'adviser_name':     'Janelli M. Mendez, DIT',
    'members': [
        'Brix A. Directo — Lead Developer',
        'Cyrille John M. Rubis — UI/UX Designer & Documentation Lead',
        'Djaunathan Albert S. Madayag — Project Manager & QA Lead',
    ],
    'completed': (
        'Completed a comprehensive documentation audit resolving 30+ inconsistencies '
        'across all project documents (CRITICAL: screen count 36→37, service count 23→26, '
        'stale model names GPT-4o→GPT-5.6 / Claude 3.5→Claude Fable 5, input modality '
        'count 7→6, FMS formula corrected from 3×33.3 pts to 4×25 pts matching code).\n\n'
        'Implemented 6 AI coverage gaps identified through a systematic audit: '
        '(1) Added set_spending_limit agentic action — AI can now set daily/weekly/'
        'monthly/yearly spending caps directly from chat; '
        '(2) Added add_insurance_policy action — AI can create SSS, PhilHealth, '
        'Pag-IBIG, and insurance entries; '
        '(3) Added Rule 12 (Taglish triggers) to system prompt for all 31 action types; '
        '(4) Budget queries now route to smart model tier (LLaMA 70B/Gemini) instead '
        'of fast 8B; '
        '(5) set_budget fallback parser added; '
        '(6) Auto-categorization rules injected into AI context.\n\n'
        'Added Financial Management Score (FMS) mini-cards to both Home screen '
        '(compact strip) and Analytics screen (full breakdown after FHS section). '
        'FHS "Unmeasured" label added when income is not set (grey chip, dashed bar).\n\n'
        'Fixed ScanReviewScreen code duplication — removed 295-line embedded class '
        'from smart_camera_screen.dart, replaced with import of standalone file.\n\n'
        'Fixed recurring transaction "Add Recurring" button — now saves entry directly '
        'without navigating to an empty screen.\n\n'
        'Rebuilt SMARTSPEND_FINAL_V5.docx (17 sections, 63 refs, 7 images). '
        'Action count updated 29→31 across all 12 documentation files. '
        'Released v2.9.8 to GitHub with 3 APK variants.'
    ),
    'ongoing': (
        'Manuscript revision: Chapters 1–4 complete in SMARTSPEND_REVISED_MANUSCRIPT.md. '
        'Google Docs version pending Fixes 11–16 from MANUSCRIPT_GUIDE.md (account type '
        'flexibility, target population reframing, respondent criteria, validator roles).\n\n'
        'Pre-Final Defense preparation: demo flow rehearsal (8–9 minutes per '
        'DEFENSE_GUIDE.md Part 2), presentation slides in progress.\n\n'
        'SUS survey instrument finalization for Week 7 administration.'
    ),
    'problems': (
        'Multi-item AI logging was silently failing for messages with typos '
        '(e.g. "spen 30 for transport") — max_tokens estimate was too low and '
        'the multi-item detection regex did not catch common Filipino spelling variants.\n\n'
        'The 50/30/20 analytics card and income-based overview cards were rendering '
        'in Lightweight Mode using stale income data (₱650 set months ago), '
        'producing meaningless percentages for a student user with income tracking OFF.'
    ),
    'solutions': (
        'Multi-item fix: broadened verb detection regex to catch typos (spen, spe, '
        'nagastos, ginastos); added connector-aware detection (comma/and); '
        'scaled token budget 800–1200 by item count; multi-item messages now route '
        'to smart tier model instead of fast 8B. System prompt Rules 1+2 updated '
        'with explicit examples.\n\n'
        'Lightweight Mode fix: added _incomeWalletMode state field to AnalyticsScreen, '
        'loaded from DB on _loadData(). All 3 income-dependent cards (50/30/20, '
        'Tax+Savings, Allowance Overview) gated on _incomeWalletMode flag.'
    ),
    'next_steps': (
        'Week 5: Configure GitHub Actions secrets (KEYSTORE_BASE64, KEY_PROPERTIES, '
        'GOOGLE_SERVICES_JSON, APP_CONFIG_DART) for automated release pipeline.\n\n'
        'Apply Google Docs manuscript fixes (Fixes 11–16): account type flexibility, '
        'target population reframing, respondent criteria, validator roles.\n\n'
        'Create Figure 1.1 bar chart (BSP financial literacy data) and insert in Google Docs.\n\n'
        'Obtain validator signatures on Appendix A validation certificates.\n\n'
        'Complete Pre-Final Defense presentation slides and conduct full rehearsal.'
    ),
}

WEEK5 = {
    'project_title': (
        'SmartSpend: An AI-Assisted Multi-Modal Personal Financial Management '
        'Application for Filipino Users Using Agentic Large Language Model Architecture'
    ),
    'report_no':        '05',
    'date_submitted':   'September 4, 2026',
    'adviser_name':     'Janelli M. Mendez, DIT',
    'members': [
        'Brix A. Directo — Lead Developer',
        'Cyrille John M. Rubis — UI/UX Designer & Documentation Lead',
        'Djaunathan Albert S. Madayag — Project Manager & QA Lead',
    ],
    'completed': (
        'Analyzed real device debug data (Debug Log, JSON backup, CSV export from '
        'Poco X6 Pro test device) and identified 5 specific bugs:\n\n'
        '(1) Lightweight Mode analytics: 50/30/20 card, Tax+Savings card, and '
        'Allowance Overview were rendering using stale ₱650 income even when income '
        'tracking was OFF. Fixed by adding _incomeWalletMode state to AnalyticsScreen '
        'and gating all 3 cards on the flag.\n\n'
        '(2) Logging Consistency scoring formula: activeDays was computed from the '
        'user\'s first logged entry this month — causing 1 entry on Sep 2 to give '
        '25/25 "Logging every active day". Fixed: if first entry is on day 1–7 of '
        'month, use full daysPassed (honest). Grace period preserved for mid-month '
        'starters (after day 7). Applied to 3 places in score_service.dart.\n\n'
        '(3) FMS "See breakdown" tap: navigated to Profile top with no scroll. Fixed '
        'using static flag ProfileScreen.scrollToFMS + GlobalKey + ScrollController. '
        'Now auto-scrolls to FMS section after data loads.\n\n'
        '(4) AI language detection: AI was replying in Taglish when user wrote in '
        'English. Fixed by adding Rule 12 to system prompt and updating persona line.\n\n'
        '(5) Score history chart: showed 2 data points with no explanation. Added '
        'placeholder card explaining scores are recorded on days app is opened.\n\n'
        'Created GitHub Actions CI/CD workflow (.github/workflows/release.yml): '
        'auto-builds all 3 APK variants and creates GitHub Release on version tag push.\n\n'
        'Built and released v2.9.9 to GitHub. Rebuilt progress reports '
        'using exact Lorma BSIT progress report template format.\n\n'
        'Added UX/Behavioral backlog (8 items) and Document Tooling backlog '
        '(8 tools to research) to PROJECT_STATUS.md for future planning.'
    ),
    'ongoing': (
        'Configuring GitHub Actions secrets (KEYSTORE_BASE64, KEY_PROPERTIES, '
        'GOOGLE_SERVICES_JSON, APP_CONFIG_DART) to enable automated release pipeline.\n\n'
        'Pre-Final Defense preparation: presentation slides and demo flow rehearsal '
        '(8–9 minutes per DEFENSE_GUIDE.md).\n\n'
        'Google Docs manuscript: applying Fixes 11–16 (account type, target population, '
        'respondent criteria, validator roles).\n\n'
        'SUS survey instrument finalization and respondent recruitment coordination.'
    ),
    'problems': (
        'GitHub Actions automated release workflow requires 4 secrets to be configured '
        'manually in repository settings before it can build — keystore, '
        'key.properties, google-services.json, and app_config.dart. These contain '
        'signing credentials and API keys that cannot be committed to the repository.\n\n'
        'Logging Consistency score was giving inflated results (25/25) to users who '
        'log on the first day of the month, because activeDays was computed from the '
        'first logged entry rather than from the start of the month.'
    ),
    'solutions': (
        'GitHub Actions: workflow file created and committed. Detailed setup '
        'instructions documented in the workflow comments. Once secrets are configured '
        'in GitHub Settings → Secrets and variables → Actions, future releases require '
        'only: git tag v2.9.X && git push origin v2.9.X\n\n'
        'Logging Consistency: fixed activeDays baseline to use daysPassed (days elapsed '
        'since month start). Grace period retained: if user started logging after the '
        '7th of the month, span from first entry is still used. Applied consistently '
        'across FHS full mode, FHS lightweight mode, and FMS (3 locations).'
    ),
    'next_steps': (
        'Configure GitHub Actions secrets to enable automated release pipeline.\n\n'
        'Complete Pre-Final Defense preparation: finalize slides, conduct full '
        '8–9 minute demo rehearsal using DEFENSE_GUIDE.md flow.\n\n'
        'Apply remaining Google Docs manuscript fixes (Fixes 11–16).\n\n'
        'Create Figure 1.1 bar chart and obtain validator signatures for Appendix A.\n\n'
        'Prepare SUS survey instruments and begin respondent recruitment '
        '(target: 30 respondents — 20 parents 35–55, 10 young professionals 21–35).'
    ),
}


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    out_dir = Path(__file__).parent
    target = sys.argv[1].lower() if len(sys.argv) > 1 else 'all'

    if target in ('week4', 'all'):
        build_report(WEEK4, out_dir / 'SmartSpend_Progress_Report_Week4.docx')
    if target in ('week5', 'all'):
        build_report(WEEK5, out_dir / 'SmartSpend_Progress_Report_Week5.docx')


if __name__ == '__main__':
    main()
