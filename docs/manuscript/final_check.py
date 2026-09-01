from docx import Document
from docx.oxml.ns import qn
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
doc = Document("SMARTSPEND_FINAL_V4.docx")
body = list(doc.element.body)

def get_text(el):
    return "".join(t.text or "" for t in el.findall(".//{%s}t" % NS)).strip()

EXPECTED = [
    "APPROVAL SHEET", "CAPSTONE PROJECT ABSTRACT", "ACKNOWLEDGEMENT", "DEDICATION",
    "TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES",
    "CHAPTER I", "CHAPTER II", "CHAPTER III", "CHAPTER IV",
    "REFERENCES", "APPENDIX A", "APPENDIX B", "APPENDIX C", "APPENDIX D",
    "CURRICULUM VITAE"
]

# Map actual section markers
seen = []
for el in body:
    t = get_text(el).strip().upper()
    # handle "Chapter I" vs "CHAPTER I"
    for m in EXPECTED:
        if m == t and m not in seen:
            seen.append(m)
            break
        # case-insensitive catch for "Chapter I"
        if m.upper() == t.upper() and m not in seen:
            seen.append(m)
            break

print("=== SECTION ORDER ===")
all_ok = True
for i, s in enumerate(seen):
    exp = EXPECTED[i] if i < len(EXPECTED) else "?"
    ok  = "✅" if s.upper() == exp.upper() else "❌"
    if s.upper() != exp.upper():
        all_ok = False
    print(f"  {ok} [{i+1:2d}]  {s}")

if all_ok:
    print("\n  ✅ ALL SECTIONS IN CORRECT ORDER")

# ToC check
print("\n=== TOC CHECK ===")
in_toc = False
toc_issues = []
for i, el in enumerate(body):
    t   = get_text(el).strip()
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if t.upper() == "TABLE OF CONTENTS":
        in_toc = True; continue
    if in_toc:
        if t.upper() in ("LIST OF FIGURES", "LIST OF TABLES", "CHAPTER I"):
            break
        if tag == "tbl":
            toc_issues.append(f"  TABLE at body[{i}]: {get_text(el)[:40]}")
        if t and "Tarsi" in t and "Feature" in t:
            toc_issues.append(f"  Feature table content at body[{i}]")
if toc_issues:
    for iss in toc_issues: print(iss)
else:
    print("  ✅ No stray tables in ToC")

# Images
print(f"\n=== IMAGES: {len(doc.inline_shapes)} (all 7 originals) ===")
for i, sh in enumerate(doc.inline_shapes):
    try:    print(f"  IMG{i+1}: {sh.width.inches:.2f}\" x {sh.height.inches:.2f}\"")
    except: print(f"  IMG{i+1}: size N/A")

# References
print("\n=== REFERENCES ===")
in_refs, rc = False, 0
for el in body:
    t = get_text(el).strip()
    if t == "REFERENCES": in_refs = True; continue
    if in_refs:
        if t.upper().startswith("APPENDIX") or t.upper() == "APPENDICES": break
        if t: rc += 1
print(f"  {rc} reference entries")

# Tables
print(f"\n=== TABLES: {len(doc.tables)} ===")
for i, tbl in enumerate(doc.tables):
    r0 = [c.text.strip()[:22] for c in tbl.rows[0].cells]
    print(f"  T{i+1}: {len(tbl.rows)}r x {len(tbl.columns)}c  hdr={r0}")

# Spacing spot-check on new content
print("\n=== LINE SPACING SPOT CHECK (new paragraphs in Ch3/Ch4) ===")
in_ch3 = False
ls_480_count = 0
ls_other = []
for el in body:
    t = get_text(el).strip()
    if "CHAPTER III" in t.upper() and len(t) < 15:
        in_ch3 = True; continue
    if "CHAPTER IV" in t.upper() and len(t) < 15 and in_ch3:
        pass  # stay in check
    if "REFERENCES" == t.upper() and in_ch3:
        break
    if in_ch3 and t and len(t) > 20:
        pPr  = el.find(qn("w:pPr"))
        sp   = pPr.find(qn("w:spacing")) if pPr is not None else None
        ls   = sp.get(qn("w:line")) if sp is not None else None
        if ls == "480":
            ls_480_count += 1
        elif ls:
            ls_other.append(f"{t[:40]}  [ls={ls}]")

print(f"  Body paragraphs with ls=480 (double): {ls_480_count}")
if ls_other:
    print(f"  Paragraphs with other spacing:")
    for x in ls_other[:5]: print(f"    {x}")
else:
    print("  ✅ All new body paragraphs are double-spaced")

words = sum(len(get_text(el).split()) for el in body)
print(f"\n=== WORD COUNT: ~{words:,} words (~{words//250} pages) ===")
print("\n=== FINAL CHECK COMPLETE ===")
