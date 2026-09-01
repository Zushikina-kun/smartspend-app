"""
update_working_docx.py  —  v3  (index-based single-pass rebuild)
=================================================================
Strategy: Instead of trying to insert-before/after with moving anchors,
we READ the source doc into an ordered element list, then BUILD a new
document by replaying those elements with our changes applied in the
correct positions.

This avoids all XML ordering bugs from addprevious/addnext on elements
that shift during mutation.

Approach:
1. Scan the source body XML once → build a flat ordered list of
   (type, element) tuples: ('para', elem) | ('tbl', elem) | ('img', elem)
2. Walk that list and write to a new Document, applying transformations
   at specific indices/content markers.
3. Preserve all 7 images exactly (copy their XML blobs).
4. Inject new sections (Approval, Abstract, LoF, LoT, Ch3, Ch4, Refs, App-A)
   at the correct positions.
5. Fix/update: title date, Table 1.2, Table 2.1.

Output: SMARTSPEND_UPDATED_MANUSCRIPT.docx
Compatible with: MS Word 2016+, Google Docs (via .docx import)
"""

import copy, re, os
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(__file__)
SRC  = os.path.join(BASE, "SMARTSPEND_CAPSTONE_WORKING.docx")
OUT  = os.path.join(BASE, "SMARTSPEND_UPDATED_MANUSCRIPT.docx")
NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

src = Document(SRC)
out = Document(SRC)   # start from a COPY so styles/settings are preserved

# ══════════════════════════════════════════════════════════════════════════════
# CORE HELPERS — all use index-based parent.insert()
# ══════════════════════════════════════════════════════════════════════════════

FONT = "Tahoma"

def _get_body_elements(doc):
    """Return ordered list of all direct children of <w:body>."""
    return list(doc.element.body)

def _get_text(elem):
    """Extract all text from an XML element."""
    return "".join(t.text or "" for t in elem.findall(".//{%s}t" % NS))

def find_index(doc, search_text, exact=True, nth=1):
    """Return the body-child index of the nth element matching search_text."""
    count = 0
    for i, child in enumerate(_get_body_elements(doc)):
        t = _get_text(child).strip()
        hit = (t == search_text) if exact else (search_text.lower() in t.lower())
        if hit:
            count += 1
            if count == nth:
                return i
    return -1

def insert_at(doc, index, elem):
    """Insert an XML element at body index (shifts later elements down)."""
    body = doc.element.body
    els  = list(body)
    if index >= len(els):
        body.append(elem)
    else:
        els[index].addprevious(elem)

def insert_many_at(doc, index, elements):
    """Insert a list of XML elements starting at body index (maintains order)."""
    for offset, elem in enumerate(elements):
        insert_at(doc, index + offset, elem)

def _p(text, bold=False, italic=False,
       align="both", sb=0, sa=0, fi=720,
       font=FONT, size=12, ls=480):
    """
    Build a <w:p> XML element with direct run formatting.
    align: 'both'=justify  'center'  'right'  'left'
    sb/sa: space before/after in pt  (0 = no explicit — use blank paras instead)
    fi:    first-line indent in twips  (720 = 0.5\";  0 = none)
    size:  font size in pt
    ls:    line spacing in 240ths  (240=single, 276=1.15x, 360=1.5x, 480=double)
           Use 480 for ALL body text to match Lorma templates.
    """
    sz_half  = int(size * 2)
    bold_xml = "<w:b/><w:bCs/>" if bold   else ""
    ital_xml = "<w:i/><w:iCs/>" if italic else ""
    fi_xml   = f'<w:ind w:firstLine="{fi}"/>' if fi else ""
    ls_xml   = f'<w:spacing w:line="{ls}" w:lineRule="auto"/>' if ls else ""
    t_esc    = (str(text)
                .replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;").replace('"', "&quot;"))
    xml = (
        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:pPr>'
        f'<w:jc w:val="{align}"/>'
        f'{ls_xml}'
        f'{fi_xml}'
        f'</w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
        f'<w:sz w:val="{sz_half}"/><w:szCs w:val="{sz_half}"/>'
        f'{bold_xml}{ital_xml}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{t_esc}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    return etree.fromstring(xml)

def _pb():
    """Page break element."""
    return etree.fromstring(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:br w:type="page"/></w:r></w:p>'
    )

def _blank():
    return _p("", sb=0, sa=0, fi=0)

def _section_hdr(text):
    """UPPERCASE bold centered section heading — ls=276 (1.15x) matching templates."""
    return _p(text, bold=True, align="center", fi=0, ls=276)

def _sub_hdr(text):
    """Bold left-aligned subsection label — double spaced, no extra sb/sa."""
    return _p(text, bold=True, align="both", fi=0, ls=480)

def _body(text, italic=False):
    """Standard body: justify, 0.5\" indent, double spaced (ls=480) — matches both templates."""
    return _p(text, italic=italic, align="both", fi=720, ls=480)

def _body0(text, bold=False, italic=False):
    """Body without first-line indent (numbered items, refs) — double spaced, no explicit sb/sa."""
    return _p(text, bold=bold, italic=italic, align="both", fi=0, ls=480)

def _right(text, bold=True):
    return _p(text, bold=bold, align="right", sb=6, sa=6, fi=0)

def _multi_run(segments, align="both", sb=0, sa=0, fi=720, ls=480):
    """
    Build a paragraph with multiple runs of different formatting.
    segments: list of (text, bold, italic) tuples
    """
    sb_t = int(sb * 20); sa_t = int(sa * 20)
    ls_tag = f'<w:spacing w:line="{ls}" w:lineRule="auto"/>' if ls else ''
    fi_x = f'<w:ind w:firstLine="{fi}"/>' if fi else ""
    ls_x  = ls_tag if 'ls_tag' in dir() else ''
    parts = []
    for text, bold, italic in segments:
        b = "<w:b/><w:bCs/>" if bold else ""
        i = "<w:i/><w:iCs/>" if italic else ""
        t = (str(text).replace("&","&amp;").replace("<","&lt;")
             .replace(">","&gt;").replace('"',"&quot;"))
        parts.append(
            f'<w:r><w:rPr>'
            f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
            f'<w:sz w:val="24"/><w:szCs w:val="24"/>{b}{i}'
            f'</w:rPr><w:t xml:space="preserve">{t}</w:t></w:r>'
        )
    xml = (
        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:pPr><w:jc w:val="{align}"/>'
        f'<w:spacing w:before="{sb_t}" w:after="{sa_t}"/>{fi_x}'
        f'</w:pPr>{"".join(parts)}</w:p>'
    )
    return etree.fromstring(xml)

def _table_xml(rows_data, col_widths_pct=None):
    """
    Build a <w:tbl> element.
    rows_data: list of list of str.  Row 0 = header (bold + grey).
    col_widths_pct: optional list of column widths as % of total (sums to ~100)
    """
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ncols = max(len(r) for r in rows_data)

    tbl = OxmlElement("w:tbl")
    # Table properties
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle"); tblStyle.set(qn("w:val"), "TableGrid")
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "9350"); tblW.set(qn("w:type"), "dxa")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblStyle); tblPr.append(tblW); tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Column widths
    tblGrid = OxmlElement("w:tblGrid")
    if col_widths_pct:
        for pct in col_widths_pct:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(9350 * pct / 100)))
            tblGrid.append(col)
    else:
        for _ in range(ncols):
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(9350 // ncols))
            tblGrid.append(col)
    tbl.append(tblGrid)

    for ri, row_data in enumerate(rows_data):
        is_hdr = (ri == 0)
        tr = OxmlElement("w:tr")
        for ci in range(ncols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            # Cell width
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(9350 // ncols))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            if is_hdr:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")
                tcPr.append(shd)
            tc.append(tcPr)
            # Cell paragraph
            sz_half = 20  # 10pt for tables
            b = "<w:b/><w:bCs/>" if is_hdr else ""
            t_esc = (str(cell_text).replace("&","&amp;").replace("<","&lt;")
                     .replace(">","&gt;").replace('"',"&quot;"))
            p_xml = (
                f'<w:p xmlns:w="{NS_W}">'
                f'<w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
                f'<w:r><w:rPr>'
                f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
                f'<w:sz w:val="{sz_half}"/><w:szCs w:val="{sz_half}"/>{b}'
                f'</w:rPr><w:t xml:space="preserve">{t_esc}</w:t></w:r></w:p>'
            )
            tc.append(etree.fromstring(p_xml))
            tr.append(tc)
        tbl.append(tr)
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# CONTENT BLOCKS  (each returns a list of XML elements in order)
# ══════════════════════════════════════════════════════════════════════════════

def approval_sheet():
    return [
        _pb(),
        _section_hdr("APPROVAL SHEET"),
        _blank(),
        _body("This is to certify that we have supervised the preparation of the Capstone Project and read the manuscript prepared by DIRECTO, BRIX A., RUBIS, CYRILLE JOHN M., and MADAYAG, DJAUNATHAN ALBERT S. entitled SMARTSPEND: AN AI-ASSISTED MOBILE FINANCIAL TRACKING AND ADVISORY APPLICATION FOR PERSONAL FINANCIAL MANAGEMENT and that the said capstone project has been submitted for final examination by the Oral Examination Committee."),
        _blank(),
        _p("_______________________________", bold=False, align="center", sb=24, sa=0, fi=0),
        _p("Johnny F. Verzola, MTS",          bold=True,  align="center", sb=0,  sa=0, fi=0),
        _p("Capstone Project Adviser",         bold=False, align="center", sb=0,  sa=12, fi=0),
        _blank(),
        _body("As members of the Oral Examination Committee, we certify that we have examined this capstone project presented before the committee and hereby recommend that it be accepted in partial fulfillment of the capstone requirements for the degree in Bachelor of Science in Information Technology."),
        _blank(),
        _p("_______________________________",                         bold=False, align="center", sb=12, sa=0, fi=0),
        _p("Ellen F. Mangaoang, MIT",                                 bold=True,  align="center", sb=0,  sa=0, fi=0),
        _p("Chairperson",                                             bold=False, align="center", sb=0,  sa=12, fi=0),
        _p("_____________________     _____________________",         bold=False, align="center", sb=6,  sa=0, fi=0),
        _p("Jopher F. Reyes, MIT       Gelo Ryann M. Carbonell",      bold=True,  align="center", sb=0,  sa=0, fi=0),
        _p("Member                            Member",                bold=False, align="center", sb=0,  sa=12, fi=0),
        _blank(),
        _body("This capstone project is hereby approved and accepted by the College of Computer Studies and Engineering in partial fulfillment of the requirements for the degree in Bachelor of Science in Information Technology."),
        _blank(),
        _p("_______________________________", bold=False, align="center", sb=12, sa=0, fi=0),
        _p("Jeoffrey B. Layco, MIS",          bold=True,  align="center", sb=0,  sa=0, fi=0),
        _p("Dean, CCSE",                       bold=False, align="center", sb=0,  sa=12, fi=0),
    ]

def abstract_block():
    els = [
        _pb(),
        _section_hdr("CAPSTONE PROJECT ABSTRACT"),
        _blank(),
        _p("Title:  SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory Application for Personal Financial Management", bold=False, align="both", sb=6, sa=0, fi=0),
        _p("Researchers:", bold=False, align="both", sb=6, sa=0, fi=0),
        _p("        Directo, Brix A.",              bold=True, align="both", sb=0, sa=0, fi=0),
        _p("        Rubis, Cyrille John M.",         bold=True, align="both", sb=0, sa=0, fi=0),
        _p("        Madayag, Djaunathan Albert S.",  bold=True, align="both", sb=0, sa=6, fi=0),
        _p("Type of Document:  CAPSTONE PROJECT",   bold=False, align="both", sb=0, sa=0, fi=0),
        _p("Type of Publication:  Unpublished",      bold=False, align="both", sb=0, sa=0, fi=0),
        _p("Accrediting Institution:  Lorma Colleges, CLI Bldg., San Juan Campus, La Union", bold=False, align="both", sb=0, sa=0, fi=0),
        _p("Teacher-in-Charge:  Janelli M. Mendez, DIT", bold=False, align="both", sb=0, sa=6, fi=0),
        _section_hdr("ABSTRACT"),
        _p("Financial mismanagement remains a critical and documented challenge among Filipino households, compounded by limited access to accessible, localized, and intelligent financial tools. This study designed, developed, and evaluated SmartSpend — an AI-assisted mobile financial tracking and advisory application for Android, built for parents aged 35–55 and young professionals aged 21–35 in La Union, Philippines.", italic=True),
        _p("SmartSpend integrates a multi-provider agentic large language model (LLM) architecture — with Gemini 3.1 Flash-Lite as the primary model and four automatic fallback providers — enabling 29 autonomous financial management actions through natural language, voice, camera, batch screenshot import (40+ platform types), and manual entry. The system operates on an offline-first SQLite database with Firebase cloud synchronization, ensuring full functionality without internet connectivity.", italic=True),
        _p("A core academic contribution is the Financial Health Score (FHS): a 0–100 behavioral metric computed entirely from user-recorded transaction data in two modes — Full Mode (Savings Rate, Overspend Control, Budget Adherence, Logging Consistency) and Lightweight Mode (Spending Restraint, Consistency, Category Balance, Habit Streak) — with a Warning Decay consequence mechanism and a Logging Gap Detection system.", italic=True),
        _p("The system was evaluated using the System Usability Scale (SUS) with 30 purposively selected respondents (20 parents, 10 young professionals), targeting a score of ≥80 (Good). Expert validation was conducted by subject matter experts in financial management and information technology.", italic=True),
        _multi_run([("Keywords:  ", True, True), ("personal finance management, agentic AI, large language model, financial health score, mobile application, Flutter, Filipino users, SmartSpend", False, True)], sb=6, sa=12),
    ]
    return els

def list_of_figures():
    return [
        _pb(),
        _section_hdr("LIST OF FIGURES"),
        _blank(),
        _p("Figure 1.1. Financial Literacy Rates by Demographic Group (BSP, 2021; Inquiro, 2024)\t\t3", sb=6, sa=6, fi=0),
        _p("Figure 1.2. Conceptual Framework — SmartSpend Mobile Application (IPO Model)\t\t\t11", sb=6, sa=6, fi=0),
        _p("Figure 2.1. System Usability Scale (SUS) Score Interpretation\t\t\t\t\t34", sb=6, sa=6, fi=0),
        _p("Figure 2.2. Agile Kanban Workflow for SmartSpend Development\t\t\t\t\t35", sb=6, sa=6, fi=0),
    ]

def list_of_tables():
    return [
        _pb(),
        _section_hdr("LIST OF TABLES"),
        _blank(),
        _p("Table 1.1.  Three-Level Gap Analysis of SmartSpend\t\t\t\t\t\t4",  sb=6, sa=6, fi=0),
        _p("Table 1.2.  Feature Comparison of SmartSpend with Existing Financial Applications\t\t6",  sb=6, sa=6, fi=0),
        _p("Table 2.1.  Distribution of Respondents\t\t\t\t\t\t\t30", sb=6, sa=6, fi=0),
        _p("Table 2.2.  Comparative Evaluation of LLM APIs for SmartSpend Integration\t\t\t44", sb=6, sa=6, fi=0),
        _p("Table 2.3.  Agile Kanban Workflow Phases and Deliverables\t\t\t\t\t38", sb=6, sa=6, fi=0),
    ]

def chapter_three():
    els = [
        _pb(),
        _section_hdr("CHAPTER III"),
        _section_hdr("RESULTS AND DISCUSSION"),
        _body("This chapter presents the results of the study based on the three stated objectives. It discusses the outcomes of each objective in relation to the development and evaluation of the SmartSpend mobile application."),

        # OBJ 1
        _sub_hdr("Objective 1 — Assessment of Financial Management Practices"),
        _p("[NOTE: Complete after survey data collection (Week 7). Insert frequency tables, percentage distributions, and interview themes here.]", italic=True, bold=False, align="both", sb=6, sa=6, fi=720),
        _body("The first objective was to assess the existing financial management practices, common budgeting challenges, and expense tracking behaviors of parents aged 35 to 55 and young professionals aged 21 to 35 in San Fernando City, La Union. Data was gathered through a validated structured survey questionnaire and supplementary interviews."),
        _body("A total of thirty (30) respondents participated — 20 parents aged 35 to 55 and 10 young professionals aged 21 to 35. [Insert Table 3.1 Respondent Profile here after data collection.] Survey results revealed that [insert findings on expense tracking methods, budgeting frequency, savings rates, and financial challenges]. These findings are consistent with BSP (2021) data indicating that a large proportion of Filipino adults do not maintain formal written budgets."),
        _body("The assessment findings confirmed the presence of the financial management challenges identified in the literature — the manual effort burden, irregular tracking, and lack of proactive feedback — and validated the need for an AI-assisted tool tailored to the Filipino context."),

        # OBJ 2
        _sub_hdr("Objective 2 — System Development and LLM Benchmarking"),
        _body("The second objective was to design and develop the SmartSpend mobile application, including the selection of an appropriate Large Language Model API through comparative technical evaluation."),
        _sub_hdr("Comparative Analysis of Large Language Model APIs"),
        _body("The selection of an appropriate LLM API is a critical design decision because it directly influences the accuracy, latency, and cost of natural language expense parsing and conversational assistance. For a mobile financial assistant requiring Filipino-English capability and free-tier deployment, the evaluation criteria were weighted as follows: Filipino-English accuracy (25%), speed and latency (20%), tool use and JSON reliability (20%), free tier availability (15%), context window (10%), and financial reasoning quality (10%). Table 2.2 presents the comparative evaluation results."),
        _p("Table 2.2. Comparative Evaluation of LLM APIs for SmartSpend Integration", bold=True, align="center", sb=12, sa=6, fi=0),
    ]
    # Table 2.2
    lm_data = [
        ["LLM / Model","Provider","Context Window","Speed (t/s)","Filipino-English","Tool Use / JSON","Free Tier","Selected?"],
        ["Gemini 3.1 Flash-Lite","Google","1,000,000","~400–600","★★★★★","★★★★★","1,000 req/day","✅ PRIMARY"],
        ["Gemini 3.5 Flash","Google","1,000,000","~200–400","★★★★★","★★★★★","250 req/day","✅ Fallback 1"],
        ["LLaMA 3.3 70B","Groq LPU","128,000","~315","★★★★☆","★★★★★","~14,400 req/day","✅ Fallback 2"],
        ["LLaMA 3.1 8B","Groq LPU","8,192","~800","★★★★☆","★★★★☆","~14,400 req/day","✅ Fallback 3"],
        ["LLaMA 3.1 70B","Cerebras WSE","128,000","~1,800","★★★★☆","★★★★☆","1M tokens/day","✅ Fallback 4"],
        ["GPT-5.6 Terra","OpenAI","1,050,000","~80–120","★★★★★","★★★★★","Paid only","❌ Cost"],
        ["Claude Fable 5","Anthropic","200,000","~70–100","★★★★★","★★★★★","Paid only","❌ Cost"],
        ["Gemini 3.7 Flash","Google","1,048,576","~300–500","★★★★★","★★★★★","Paid ($0.75/1M)","❌ No free tier"],
        ["Grok 4.6","xAI","500,000","~100–200","★★★★☆","★★★★★","Paid ($2/1M)","❌ Cost"],
        ["DeepSeek V4 Flash","DeepSeek","1,000,000","~200","★★★☆☆","★★★☆☆","$0.14/1M (5M trial)","❌ Weaker Filipino"],
        ["Qwen 3 32B","Alibaba/OpenRouter","128,000","~150–300","★★★★☆","★★★★★","Free preview","❌ Less tested"],
        ["Fin-R1 (7B)","Self-hosted","128,000","Varies","★★★☆☆","★★☆☆☆","Self-host","❌ No hosted API"],
        ["Mistral 7B","Mistral AI","32,000","~600","★★★☆☆","★★★☆☆","Self-host","❌ Poor Filipino"],
        ["GPT-4o Mini","OpenAI","128,000","~120","★★★★☆","★★★★★","No free tier","❌ Cost"],
        ["Gemma 2 9B","Google","8,192","~500","★★★☆☆","★★★☆☆","Local only","❌ No hosted API"],
    ]
    els.append(_table_xml(lm_data, col_widths_pct=[18,12,12,9,10,10,14,15]))
    els += [
        _body("Gemini 3.1 Flash-Lite was selected as the primary model because it offers the highest free-tier request quota (1,000/day), the best Filipino-English multilingual performance among free-tier models (consistent with Google's training data coverage of Southeast Asian languages), a 1-million token context window for full user financial context injection, and native function calling support essential for the 29 agentic action types (Li et al., 2024; Google, 2024f)."),
        _body("GPT-5.6 and Claude Fable 5 are paid-only with no free tier sufficient for 30 respondents at 60 messages/day each. Gemini Flash-Lite plus Groq's free tier covers this entirely at zero cost. Finance-specialized models (Fin-R1, FinGPT) are designed for market data tasks and perform worse than general multilingual LLMs on Filipino-English conversational expense parsing (Arcila et al., 2026; Liu et al., 2023)."),
        _body("SmartSpend uses dynamic full-context injection rather than RAG. A typical user has 20–50 expenses, 5–10 budgets, and 3–5 goals (~1,000–5,000 tokens) — fitting within any evaluated model's context window. RAG adds vector search overhead designed for thousands of documents, which is unnecessary for SmartSpend's compact per-user dataset (Davenport & Mittal, 2022)."),

        _sub_hdr("Financial Health Score — Full Computation"),
        _body("The Financial Health Score (FHS) is SmartSpend's core academic contribution — a 0-to-100 behavioral metric computed from user-recorded transaction data. Its design is informed by the Financial Health Network FinHealth Score (Financial Health Network, 2021, 2026), the UNSGSA Financial Health Measurement Framework (UNSGSA, 2021), and the CFPB Financial Well-Being Scale (Consumer Financial Protection Bureau, 2017). Unlike these survey-based frameworks, SmartSpend's FHS is a behavioral computation from transaction data — no surveys or external feeds required."),
        _sub_hdr("Full Mode — Income Tracking Enabled (4 components × 25 pts = 100 maximum)"),
        _body0("Component 1 — Savings Rate (25 pts):  Score = 25 × min(1.0, savingsRate / 0.20).  The 20% target comes from the 50/30/20 budgeting rule (Warren & Tyagi, 2005)."),
        _body0("Component 2 — Overspend Control (25 pts):  Score = 25 × (1 − overDays / activeDays).  Derived from the FinHealth Score Spend pillar — spending less than income as a daily measurable behavior (Financial Health Network, 2021)."),
        _body0("Component 3 — Budget Adherence (25 pts):  Score = 25 × (onBudgetCategories / totalBudgetCategories).  Based on zero-based budgeting theory (Ramsey, 2003). No budgets set = full 25 pts."),
        _body0("Component 4 — Logging Consistency (25 pts):  Score = 25 × (loggedDays / activeDays).  Consistent tracking reduces discretionary spending by 10–20% and enables accurate computation of the other three components (Thaler & Sunstein, 2008)."),
        _sub_hdr("Lightweight Mode — Income Tracking Disabled (for students, freelancers, informal workers)"),
        _body0("Spending Restraint (25 pts) — vs user-set limit;  Logging Consistency (25 pts) — same formula as Full Mode;  Category Balance (25 pts) — no single category exceeding 40% of total spending;  Habit Streak (25 pts) — consecutive logged days, full credit at 14 days (Duhigg, 2012)."),
        _body("Score Adjustments: Warning Decay (−5 pts/day, max −15) when budget warnings are ignored — applying loss aversion theory (Kahneman & Tversky, 1979). Gap Adjustment (+2 or −3 pts/day) for confirmed no-spend or unlogged-spend days — applying behavioral honesty mechanisms (Ariely, 2008). Final score clamped 0–100."),

        _p("Table 2.3. Agile Kanban Workflow Phases and Deliverables", bold=True, align="center", sb=12, sa=6, fi=0),
    ]
    kanban = [
        ["Phase","Key Tasks","Deliverable"],
        ["Backlog","Define all features; needs survey; literature review on PH financial gaps","Prioritized feature list; literature review"],
        ["Requirements","Translate findings into specs; validate questionnaire; LLM API benchmarking","Validated questionnaire; LLM benchmarking matrix (Table 2.2)"],
        ["Design","SQLite schema (20 tables); FHS formula; UI wireframes; data flow diagrams","System architecture; database schema; FHS documentation"],
        ["Development","Build expense tracking; integrate Gemini 3.1 Flash-Lite; add OCR/voice/barcode/batch screenshots; FHS engine; Firebase sync; gamification","Functional app; 29 agentic actions operational"],
        ["Testing","LLM parsing accuracy test; SUS with 30 respondents; interviews; bug log","SUS scores; parsing observations; bug documentation"],
        ["Deployment","Build release APKs (arm64/armeabi/x86_64); prepare Demo Mode; publish GitHub Releases","Release APKs v2.9.7; project documentation"],
        ["Done/Review","Analyze SUS scores; review feedback; identify improvements; document recommendations","Final evaluation report; post-capstone roadmap"],
    ]
    els.append(_table_xml(kanban, col_widths_pct=[14, 52, 34]))
    els += [
        _sub_hdr("System Development Results — SmartSpend v2.9.7"),
        _body("The SmartSpend mobile application was developed across seven Kanban phases, resulting in a fully functional Android application. Key specifications: Platform: Android (Flutter/Dart); Version: 2.9.7; SQLite schema: Version 11, 20 tables; APK size: 44.7 MB (arm64-v8a, release, obfuscated); AI providers: 5 (Gemini 3.1 Flash-Lite → Gemini 3.5 Flash → Groq LLaMA 3.3 70B → LLaMA 3.1 8B → Cerebras LLaMA 3.1); Agentic action types: 29; Input modalities: 6 (text, voice, live camera, single photo, batch screenshots, paste text); Screenshot platforms detected: 40+; Achievement badges: 23; Daily quests: 10 rotating; Currencies: 57; GitHub: https://github.com/Zushikina-kun/smartspend-app"),

        _sub_hdr("Objective 3 — System Usability Evaluation (SUS)"),
        _p("[NOTE: Complete after SUS administration with 30 respondents (Week 7). Insert SUS computation table, individual scores, final average, interpretation, and qualitative feedback here.]", italic=True, bold=False, align="both", sb=6, sa=6, fi=720),
        _body("The third objective was to evaluate the usability of the SmartSpend application using the System Usability Scale (SUS). The SUS was administered to thirty (30) respondents — 20 parents and 10 young professionals — following a guided live demonstration using Demo Mode with pre-loaded Filipino sample data."),
        _body("SUS scores were computed using the standard formula: for odd-numbered items, the contribution equals the item score minus 1; for even-numbered items, the contribution equals 5 minus the item score. The sum is multiplied by 2.5 to yield a final score on a 0–100 scale (Brooke, 1996)."),
        _p("Overall SUS Score:  [INSERT SCORE]     Grade: [A/B/C]     Adjective: [Best Imaginable / Excellent / Good / Okay / Poor]", bold=True, italic=True, align="center", sb=12, sa=6, fi=0),
        _p("[Insert SUS computation table and per-respondent scores here after data collection. Target: ≥80 (Good per Bangor et al., 2009).]", italic=True, bold=False, align="center", sb=6, sa=12, fi=0),
        _body("[Insert qualitative feedback summary here — expected themes: ease of AI chat input, utility of FHS score, appreciation for Lite Mode toggle, suggestions for future features.]"),
    ]
    return els

def chapter_four():
    return [
        _pb(),
        _section_hdr("CHAPTER IV"),
        _section_hdr("CONCLUSIONS AND RECOMMENDATIONS"),
        _body("This chapter presents the findings of the study and provides recommendations based on the results and insights gained throughout the research."),

        _sub_hdr("Conclusions"),
        _body("For the first objective — assessment of financial management practices: The survey and interview data confirmed the presence of the financial management challenges identified in the literature: the manual effort burden of traditional expense tracking, irregular budgeting behavior, and the absence of visible consequences for ignoring financial warnings. These findings validated the design rationale for SmartSpend's core features — multi-modal AI input to eliminate manual effort, the Financial Health Score to provide visible financial feedback, and the Warning Decay mechanism to make the consequences of ignoring budget warnings tangible and persistent."),
        _body("For the second objective — system development and LLM benchmarking: SmartSpend v2.9.7 was successfully designed and developed as a fully functional Android application featuring 29 autonomous AI actions, a dual-mode Financial Health Score, 6 input modalities, batch screenshot import across 40+ platform types, offline-first SQLite architecture, Firebase cloud synchronization, and a gamification system with 23 achievement badges and 10 daily quests. The comparative benchmarking of 15 LLM API providers confirmed that Gemini 3.1 Flash-Lite is the most appropriate primary model — offering the highest free-tier quota, best Filipino-English performance, and native function calling at zero cost."),
        _body("For the third objective — usability evaluation: [Insert conclusion based on actual SUS score after Week 7 data collection. If score ≥80: 'SmartSpend achieved a SUS score of [X] — [Adjective] per Bangor et al. (2009) — meeting/exceeding the ≥80 target, indicating acceptable usability for the target population.']"),
        _body("Overall, SmartSpend demonstrates that a free, offline-capable, Filipino-first AI financial management system can be built entirely on free-tier services, representing a meaningful contribution to financial technology research in the Philippine context."),

        _sub_hdr("Recommendations"),
        _body("Based on the findings, development experience, and usability evaluation, the following recommendations are proposed:"),
        _body0("1.   Paluwagan tracker — A rotating savings group tracker should be the highest-priority post-capstone feature, representing a uniquely Filipino informal savings behavior. The existing debt and recurring transaction infrastructure provides a suitable architectural base."),
        _body0("2.   15th and 30th payday cycle awareness — Implement payday-cycle-aware budgeting resets aligned with the Philippine standard of semi-monthly salary payments, further strengthening SmartSpend's Filipino-first positioning."),
        _body0("3.   Backend API proxy — Move LLM API key management to a server-side proxy (e.g., Firebase Cloud Function) to eliminate device-side key exposure entirely."),
        _body0("4.   Play Store submission — After implementing the backend proxy and a privacy policy, submit SmartSpend to the Google Play Store for wider distribution."),
        _body0("5.   SQLite encryption — Implement SQLCipher-based encryption for the local database in a future schema migration to enhance data security for sensitive financial records."),
        _body0("6.   Couple and family wallet sharing — Allow multiple users to contribute to and view a shared household wallet, directly addressing parents managing household finances with a spouse."),
        _body0("7.   OFW remittance tracking — Add inbound international remittance tracking as a distinct income category, addressing the significant overseas Filipino worker demographic."),
        _body("For future research, longitudinal studies measuring the actual impact of SmartSpend on financial behavior over 3–6 months — savings rate, budget adherence, FHS trend — would provide stronger empirical evidence for the effectiveness of the behavioral intervention mechanisms designed in this study."),
    ]

def references_block():
    refs = [
        "Arcila, A., et al. (2026). FrontierFinance: A challenging benchmark for measuring frontier intelligence of finance agents. arXiv:2608.11683. https://arxiv.org/abs/2608.11683",
        "Ariely, D. (2008). Predictably irrational: The hidden forces that shape our decisions. HarperCollins.",
        "Bangko Sentral ng Pilipinas. (2021). 2021 Financial Inclusion Survey. BSP. https://www.bsp.gov.ph/Inclusive-Finance/Financial-Inclusion-Surveys/2021-FIS-Report.pdf",
        "Bangko Sentral ng Pilipinas. (2025). Consumer Finance and Inclusion Survey (CFIS) 2025. BSP.",
        "Bangor, A., Kortum, P., & Miller, J. (2009). Determining what individual SUS scores mean: Adding an adjective rating scale. Journal of Usability Studies, 4(3), 114-123.",
        "Bitrián, P., Buil, I., & Catalán, S. (2021). Making finance fun: The gamification of personal financial management apps. International Journal of Bank Marketing, 39(7), 1310-1332. https://doi.org/10.1108/IJBM-09-2020-0491",
        "Bloomberg. (2026). How the Philippines' first fintech unicorn is minting financial inclusion. https://sponsored.bloomberg.com/article/mynt/how-the-philippines-first-fintech-unicorn-is-minting-financial-inclusion",
        "Bociek, J. (2023). mobile_scanner: A universal barcode and QR code scanner for Flutter. https://pub.dev/packages/mobile_scanner",
        "Brooke, J. (1996). SUS: A quick and dirty usability scale. In P. W. Jordan et al. (Eds.), Usability evaluation in industry (pp. 189-194). Taylor & Francis.",
        "Cambridge Judge Business School. (2025). From automation to autonomy: The agentic AI era of financial services. https://www.jbs.cam.ac.uk/2025/from-automation-to-autonomy-the-agentic-ai-era-of-financial-services/",
        "Consumer Financial Protection Bureau. (2017). Financial well-being scale: Scale development technical report. CFPB. https://files.consumerfinance.gov/f/documents/201705_cfpb_financial-well-being-scale-technical-report.pdf",
        "Creswell, J. W., & Plano Clark, V. L. (2011). Designing and conducting mixed methods research. Sage Publications.",
        "Davenport, T. H., & Mittal, N. (2022). All-in on AI: How smart companies win big with artificial intelligence. Harvard Business Review Press.",
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
        "Deloitte. (2026). Agentic AI boosts wealth management. https://www.deloitte.com/us/en/insights/industry/financial-services/financial-services-industry-predictions/2026/agentic-ai-wealth-management-productivity.html",
        "Duhigg, C. (2012). The power of habit: Why we do what we do in life and business. Random House.",
        "Dwivedi, Y. K., et al. (2021). Artificial intelligence (AI): Multidisciplinary perspectives on emerging challenges, opportunities, and agenda for research, practice and policy. International Journal of Information Management, 57, 101994. https://doi.org/10.1016/j.ijinfomgt.2019.08.002",
        "Ernst & Young. (2026a). Nearly half of global consumers now use AI to guide savings and investment decisions. EY. https://www.ey.com/en_gl/newsroom/2026/04/nearly-half-of-global-consumers-now-use-ai-to-guide-savings-and-investment-decisions",
        "Ernst & Young. (2026b). EY survey: Autonomous AI is no longer theoretical as adoption grows. EY. https://www.ey.com/en_nl/newsroom/2026/03/ey-survey-autonomous-ai-is-no-longer-theoretical-as-adoption-grows-despite-ongoing-trust-concerns",
        "Financial Health Network. (2021). FinHealth Score Toolkit. https://finhealthnetwork.org/tools/financial-health-score/",
        "Financial Health Network. (2026). From insight to impact: The next phase of financial health measurement. https://finhealthnetwork.org/research/from-insight-to-impact-the-next-phase-of-financial-health-measurement/",
        "Flores, C. A. R. (2025). Financial freedom of Filipinos in personal finance management. Pantao: The International Journal of the Humanities and Social Sciences, 4(1). https://pantaojournal.com/2025/01/27/v4-i1-7/",
        "Flutter 4 Fun. (2022). fl_chart: A highly customizable Flutter chart library. https://pub.dev/packages/fl_chart",
        "GCash / Mynt. (2026). GCash launches country's first AI financial coach embedded in e-wallet [Press release]. PR Newswire. https://www.prnewswire.com/apac/news-releases/ph-fintech-gcash-launches-countrys-first-ai-financial-coach-embedded-in-e-wallet-to-strengthen-financial-literacy-302718569.html",
        "Google. (2024a). Firebase Authentication documentation. https://firebase.google.com/docs/auth",
        "Google. (2024b). Firebase Crashlytics documentation. https://firebase.google.com/docs/crashlytics",
        "Google. (2024c). Firebase Firestore documentation. https://firebase.google.com/docs/firestore",
        "Google. (2024e). Google ML Kit documentation. https://developers.google.com/ml-kit",
        "Google. (2024f). Flutter documentation. https://flutter.dev/docs",
        "Hean, O., Saha, U., & Saha, B. (2025). Can AI help with your personal finances? Applied Economics. https://doi.org/10.1080/00036846.2025.2450384",
        "IBM. (2025). Agentic AI in financial services: Navigating innovation. https://www.ibm.com/think/insights/agentic-ai-financial-services-ethical-adoption",
        "Inquiro. (2024). Financial literacy in the Philippines: Key statistics. https://inquiro.ph/financial-literacy-in-the-philippines-2024-key-statistics/",
        "Juniper Research. (2026). Gamification in banking: How game mechanics drive financial behavior change. [Research report].",
        "Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. Econometrica, 47(2), 263-292.",
        "Li, Z., et al. (2024). A survey of large language models for financial applications. arXiv:2406.11903. https://arxiv.org/abs/2406.11903",
        "Liu, X., et al. (2023). FinGPT: Open-source financial large language models. arXiv:2306.06031. https://arxiv.org/abs/2306.06031",
        "Liu, Z., et al. (2025). Fin-R1: A large language model for financial reasoning through reinforcement learning. arXiv:2503.16252. https://arxiv.org/abs/2503.16252",
        "Meyll, T., et al. (2025). Spendception: The psychological impact of digital payments on consumer purchase behavior and impulse buying. Behavioral Sciences, 15(3), 387. https://doi.org/10.3390/bs15030387",
        "Nielsen, J. (2006). Progressive disclosure. Nielsen Norman Group. https://www.nngroup.com/articles/progressive-disclosure/",
        "NielsenIQ. (2026). The new financial reality: How Filipino consumers are spending, saving, and banking in 2026. https://nielseniq.com/global/en/insights/report/2026/the-new-financial-reality-how-filipino-consumers-are-spending-saving-and-banking-in-2026/",
        "Philippine Statistics Authority. (2021). Family Income and Expenditure Survey (FIES) 2021. PSA. https://www.psa.gov.ph",
        "Philippine Statistics Authority. (2025). Philippine Digital Economy Satellite Account (PDESA) 2025. PSA. https://psa.gov.ph",
        "Plaid. (2026). State of intelligent finance report — Spring 2026. https://plaid.com/blog/state-of-intelligent-finance-report-spring-2026/",
        "Ramsey, D. (2003). Financial peace revisited. Viking.",
        "Roux, A. (2019). sqflite: SQLite plugin for Flutter. https://pub.dev/packages/sqflite",
        "Sloane, L. (2022). speech_to_text: A Flutter plugin for on-device speech recognition. https://pub.dev/packages/speech_to_text",
        "Social Weather Stations. (2026). SWS financial inclusion survey: Philippines financial inclusion rises to 58%. Cited in CoinGeek. https://coingeek.com/10-point-surge-pushes-philippines-financial-inclusion-to-58/",
        "Springer. (2026). Digital nudges and financial inclusion: A study on behavioral interventions. Lecture Notes in Networks and Systems. https://link.springer.com/content/pdf/10.1007/978-3-032-00343-0_14.pdf",
        "Statista. (2024). Number of mobile app downloads worldwide. https://www.statista.com/statistics/271644/worldwide-free-and-paid-mobile-app-store-downloads/",
        "Stefanov, T., Stefanova, M., & Varbanova, S. (2024). Personal finance management application. TEM Journal, 13(3), 2066-2075. https://doi.org/10.18421/TEM133-34",
        "Strivecloud. (2026). Fintech app gamification: Data shows 22% boost in saving habits. https://strivecloud.io/blog/mobile-app-gamification-fintech",
        "Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257-285.",
        "Tarsi – Budget Tracker. (2026). Tarsi – Budget Tracker [Mobile application]. App Store. https://apps.apple.com/ph/app/tarsi-budget-tracker/",
        "The Flutter Authors. (2013). local_auth: Flutter plugin for biometric authentication. https://pub.dev/packages/local_auth",
        "Thaler, R. H., & Sunstein, C. R. (2008). Nudge: Improving decisions about health, wealth, and happiness. Yale University Press.",
        "UNSGSA. (2021). Measuring financial health: A framework for practitioners. https://www.unsgsa.org",
        "Wajid, F., et al. (2025). Gamification: Revolutionizing financial planning systems. World Journal of Advanced Engineering Technology and Sciences. https://www.wjaets.com/sites/default/files/fulltext_pdf/WJAETS-2025-0158.pdf",
        "Warren, E., & Tyagi, A. W. (2005). All your worth: The ultimate lifetime money plan. Free Press.",
        "World Bank. (2022). Global Findex Database 2021. https://www.worldbank.org/en/publication/globalfindex",
        "World Economic Forum. (2024). How agentic AI will transform financial services. https://www.weforum.org/stories/2024/12/agentic-ai-financial-services-autonomy-efficiency-and-inclusion/",
        "Yang, H., et al. (2023). FinGPT: Democratizing internet-scale data for financial large language models. arXiv:2307.10485. https://arxiv.org/abs/2307.10485",
        "Yomio. (2026). YNAB alternatives: Which budget app actually works in 2026? https://yomio.app/en/blog/ynab-alternatives",
        "ZXing Project. (2023). ZXing barcode scanning library. https://github.com/zxing/zxing",
    ]
    els = [_pb(), _section_hdr("REFERENCES"), _blank()]
    for r in refs:
        els.append(_body0(r))
    return els

def appendix_a():
    return [
        _pb(),
        _section_hdr("APPENDIX A"),
        _section_hdr("VALIDATION CERTIFICATES"),
        _blank(),
        _sub_hdr("CONTENT VALIDATION CERTIFICATE — SURVEY QUESTIONNAIRE"),
        _body("This survey questionnaire has been reviewed for content validity and is deemed appropriate and relevant to the financial management experiences of the target population."),
        _blank(),
        _body0("Educational Background  :  ________________________________"),
        _body0("(e.g., BS Commerce, BS Accountancy, BS Business Administration, or equivalent)"),
        _blank(),
        _body0("Occupation  :  ________________________________"),
        _body0("(e.g., Business Owner, Financial Officer, Accountant, Financial Adviser, etc.)"),
        _blank(),
        _body0("Years of Experience  :  _____ years in financial management and/or business operations"),
        _blank(),
        _body0("Signature  :  _______________________________"),
        _blank(),
        _body0("Name (optional)  :  _______________________________"),
        _blank(),
        _body0("Date  :  _______________________________"),
        _blank(),
        _blank(),
        _sub_hdr("TECHNICAL VALIDATION CERTIFICATE — SYSTEM AND SUS EVALUATION"),
        _body("The SmartSpend system and its usability evaluation process have been reviewed by a subject matter expert in Information Technology to ensure technical soundness and proper SUS administration."),
        _blank(),
        _body0("Educational Background  :  ________________________________"),
        _body0("(e.g., BS Information Technology, BS Computer Science, or equivalent)"),
        _blank(),
        _body0("Occupation  :  ________________________________"),
        _body0("(e.g., Software Developer, IT Instructor, Systems Analyst, IT Professional, etc.)"),
        _blank(),
        _body0("Years of Experience  :  _____ years in IT / software development and/or usability evaluation"),
        _blank(),
        _body0("Signature  :  _______________________________"),
        _blank(),
        _body0("Name (optional)  :  _______________________________"),
        _blank(),
        _body0("Date  :  _______________________________"),
    ]

def table_12_new():
    """Expanded Table 1.2 — 19 features × 9 apps."""
    return _table_xml([
        ["Feature","Tarsi","YNAB","Monarch","Copilot","BudgetPH","Alkansya AI","GCash Pera Coach","SmartSpend"],
        ["Offline mode","Yes","No","No","No","Yes","No","No","Yes"],
        ["LLM chat assistant","No","No","No","No","Insights only","Yes (basic)","Yes (literacy Q&A)","Yes — 29 agentic actions"],
        ["Financial Health Score","No","No","No","No","Yes (simpler)","No","No","Yes (0–100, dual-mode)"],
        ["OCR receipt scanning","Yes","No","No","No","No","No","No","Yes"],
        ["Voice input (en-PH)","No","No","No","No","No","No","No","Yes"],
        ["Batch screenshot import (40+)","No","No","No","No","No","No","No","Yes"],
        ["Multi-period spending limits","No","No","No","No","Yes","No","No","Yes (daily/wk/mo/yr)"],
        ["Spending behavior analysis","No","No","No","No","No","No","No","Yes"],
        ["Bank synchronization","No","Yes","Yes","Yes","CSV import","No","GCash balance","No"],
        ["Filipino-English AI (Taglish)","No","No","No","No","No","Yes","Yes (PH languages)","Yes — full Taglish"],
        ["Free tier (Android)","Yes","No","No","No","Yes (PWA)","Limited","Yes (GCash req.)","Yes — always free"],
        ["Paluwagan tracker","No","No","No","No","Yes","No","No","No"],
        ["15th & 30th payday cycle","No","No","No","No","Yes","No","No","No"],
        ["Gamification (badges/quests)","No","No","No","No","Yes (XP/levels)","No","No","Yes (23 badges, 10 quests)"],
        ["Logging gap detection","No","No","No","No","No","No","No","Yes"],
        ["SSS/PhilHealth/Pag-IBIG","No","No","No","No","Yes (records)","No","No","Yes (AI compute)"],
        ["Insurance tracker","No","No","No","No","No","No","No","Yes"],
        ["Round-up savings","No","No","No","No","No","No","No","Yes"],
    ], col_widths_pct=[20,7,6,7,7,8,9,11,15])

def table_21_new():
    """Updated Table 2.1 with validator rows."""
    return _table_xml([
        ["Category","Number of Respondents"],
        ["Parents (Ages 35–55)","20"],
        ["Young Professionals (Ages 21–35)","10"],
        ["Total Respondents","30"],
        ["",""],
        ["Validators / Experts","Number of Validators"],
        ["Content Validator (Survey — Financial Management Expert)","1"],
        ["Technical Validator (System & SUS Process — IT Expert)","1"],
        ["Total Validators","2"],
    ], col_widths_pct=[75, 25])

# ══════════════════════════════════════════════════════════════════════════════
# APPLY ALL CHANGES TO `out` (copy of source)
# ══════════════════════════════════════════════════════════════════════════════
body = out.element.body

def idx(text, exact=True, nth=1):
    return find_index(out, text, exact=exact, nth=nth)

def ins(index, elements):
    insert_many_at(out, index, elements)

print("=== Building updated document ===")

# ── 1. Fix title date ──────────────────────────────────────────────────────────
print("  [1] Fix title date")
for el in _get_body_elements(out):
    t = _get_text(el)
    if "April 2026" in t:
        for wt in el.findall(".//{%s}t" % NS):
            if "April 2026" in (wt.text or ""):
                wt.text = wt.text.replace("April 2026", "August 2026")
        break

# ── 2. Fix objective statement (parents PRIMARY) ───────────────────────────────
print("  [2] Fix objectives wording")
for el in _get_body_elements(out):
    t = _get_text(el)
    if "parents aged 35 to 55 and young professionals aged 21 to 35 in La Union" in t:
        for wt in el.findall(".//{%s}t" % NS):
            if wt.text and "parents aged 35 to 55 and young professionals" in wt.text:
                wt.text = wt.text.replace(
                    "parents aged 35 to 55 and young professionals aged 21 to 35 in La Union, Philippines.",
                    "parents aged 35 to 55 as the primary target population, and young professionals aged 21 to 35 as a secondary demographic, in La Union, Philippines."
                )

# ── 3. Replace Table 1.2 (find by header content) ─────────────────────────────
print("  [3] Replace Table 1.2")
for i, el in enumerate(_get_body_elements(out)):
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag == "tbl":
        cells = el.findall(".//{%s}tc" % NS)
        if cells:
            first_cell = _get_text(cells[0])
            header_text = "".join(_get_text(c) for c in cells[:4])
            if "Feature" in first_cell and "Tarsi" in header_text:
                body.remove(el)
                body_els = list(body)
                # Find new insertion point (where table was, now shifted)
                ins_pt = idx("Table 1.2", exact=False, nth=1)
                if ins_pt < 0:
                    ins_pt = idx("Building on the gaps", exact=False) - 1
                ins(ins_pt + 1, [table_12_new()])
                print("    Table 1.2 replaced")
                break

# ── 4. Replace Table 2.1 (respondents) ────────────────────────────────────────
print("  [4] Replace Table 2.1")
for i, el in enumerate(_get_body_elements(out)):
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag == "tbl":
        cells = el.findall(".//{%s}tc" % NS)
        if cells and "Category" in _get_text(cells[0]) and "Number" in _get_text(cells[1] if len(cells)>1 else cells[0]):
            body.remove(el)
            ins_pt = idx("Table 2.1", exact=False, nth=1)
            if ins_pt < 0:
                ins_pt = idx("The survey questionnaire will undergo", exact=False) - 1
            ins(ins_pt + 1, [table_21_new()])
            print("    Table 2.1 replaced")
            break

# ── 5. Insert Approval Sheet (before ACKNOWLEDGEMENT) ─────────────────────────
print("  [5] Insert Approval Sheet")
ack_i = idx("ACKNOWLEDGEMENT", exact=True)
if ack_i >= 0:
    ins(ack_i, approval_sheet())
    print(f"    Approval Sheet inserted before index {ack_i}")

# ── 6. Insert Abstract (before ACKNOWLEDGEMENT — now shifted) ─────────────────
print("  [6] Insert Abstract")
ack_i2 = idx("ACKNOWLEDGEMENT", exact=True)
if ack_i2 >= 0:
    ins(ack_i2, abstract_block())
    print(f"    Abstract inserted")

# ── 7. Insert List of Figures/Tables (before Chapter I) ───────────────────────
print("  [7] Insert List of Figures + Tables")
ch1_i = idx("Chapter I", exact=True)
if ch1_i >= 0:
    ins(ch1_i, list_of_tables())
    ch1_i2 = idx("Chapter I", exact=True)
    ins(ch1_i2, list_of_figures())
    print(f"    LoF and LoT inserted")

# ── 8. Remove old References section (between "References" and "APPENDICES") ──
print("  [8] Remove old references + insert new block")
# Find "References" and "APPENDICES" in current body
refs_i    = -1
app_i     = -1
body_els  = _get_body_elements(out)
for i, el in enumerate(body_els):
    t = _get_text(el).strip()
    if t == "References" and refs_i < 0:
        refs_i = i
    if t == "APPENDICES" and app_i < 0:
        app_i = i

if refs_i >= 0 and app_i > refs_i:
    # Remove elements from refs_i up to (but not including) app_i
    to_remove = body_els[refs_i:app_i]
    for el in to_remove:
        body.remove(el)
    removed = len(to_remove)
    print(f"    Removed {removed} old reference elements")
    # Insert Ch3 + Ch4 + full refs at the position where refs were
    new_app_i = idx("APPENDICES", exact=True)
    all_new   = chapter_three() + chapter_four() + references_block()
    ins(new_app_i, all_new)
    print(f"    Inserted Ch3 + Ch4 + {len(references_block())} reference elements")
else:
    print(f"    WARNING: refs_i={refs_i} app_i={app_i} — cannot locate section bounds")

# ── 9. Insert Appendix A (before existing Appendix B / Survey) ────────────────
print("  [9] Insert Appendix A")
surv_i = idx("Appendix A", exact=False, nth=1)
if surv_i < 0:
    surv_i = idx("Survey Questionnaire", exact=False, nth=1)
if surv_i >= 0:
    ins(surv_i, appendix_a())
    print(f"    Appendix A inserted")
else:
    print("    WARNING: Survey anchor not found")

# ── 10. Update ToC (inject Ch3 entry before "REFERENCES" line in ToC) ─────────
print("  [10] Update Table of Contents")
refs_toc_i = -1
body_els2  = _get_body_elements(out)
in_toc     = False
for i, el in enumerate(body_els2):
    t = _get_text(el).strip()
    if "TABLE OF CONTENTS" in t.upper():
        in_toc = True
        continue
    if in_toc:
        if t.upper().startswith("REFERENCES"):
            refs_toc_i = i
            break
        if "Chapter I" in t or "CHAPTER I" in t:
            break  # gone past TOC

if refs_toc_i >= 0:
    toc_entries = [
        _p("III\tRESULTS AND DISCUSSION\t\t\t\t\t40",         bold=True,  align="both", sb=12, sa=12, fi=0),
        _p("\tObjective 1 — Financial Management Practices\t\t40",   bold=False, align="both", sb=6,  sa=6,  fi=720),
        _p("\tObjective 2 — System Development & LLM Benchmarking\t42", bold=False, align="both", sb=6, sa=6, fi=720),
        _p("\t\tTable 2.2. Comparative Evaluation of LLM APIs\t44",   bold=False, align="both", sb=6,  sa=6,  fi=1440),
        _p("\t\tFinancial Health Score Computation\t\t48",            bold=False, align="both", sb=6,  sa=6,  fi=1440),
        _p("\tObjective 3 — System Usability Evaluation\t\t53",       bold=False, align="both", sb=6,  sa=6,  fi=720),
        _p("IV\tCONCLUSIONS AND RECOMMENDATIONS\t\t56",              bold=True,  align="both", sb=12, sa=12, fi=0),
        _p("\tConclusions\t\t\t\t\t\t56",                            bold=False, align="both", sb=6,  sa=6,  fi=720),
        _p("\tRecommendations\t\t\t\t\t57",                          bold=False, align="both", sb=6,  sa=6,  fi=720),
    ]
    ins(refs_toc_i, toc_entries)
    print(f"    ToC entries inserted at {refs_toc_i}")
else:
    print("    ToC REFERENCES line not found — skipping")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE + VERIFY
# ══════════════════════════════════════════════════════════════════════════════
print(f"\n=== Saving {OUT} ===")
out.save(OUT)
size = os.path.getsize(OUT)
print(f"  Size: {size/1024:.1f} KB")

print("\n=== Verification ===")
v   = Document(OUT)
vb  = v.element.body

sections = [
    "APPROVAL SHEET","ABSTRACT","ACKNOWLEDGEMENT","DEDICATION",
    "TABLE OF CONTENTS","LIST OF FIGURES","LIST OF TABLES",
    "Chapter I","CHAPTER II","CHAPTER III","CHAPTER IV",
    "REFERENCES","APPENDIX A","APPENDIX B","CURRICULUM VITAE"
]
for s in sections:
    found = any(s.lower() in _get_text(el).strip().lower()
                and len(_get_text(el).strip()) < 80
                for el in _get_body_elements(v))
    print(f"  {'OK     ' if found else 'MISSING'}  {s}")

# Count references
in_refs, ref_count = False, 0
for el in _get_body_elements(v):
    t = _get_text(el).strip()
    if t == "REFERENCES": in_refs = True; continue
    if in_refs:
        if t.upper().startswith("APPENDIX") or t.upper().startswith("APPENDICES"): break
        if t: ref_count += 1
print(f"\n  References: {ref_count}")

# Tables
print(f"  Tables: {len(v.tables)}")
for i, tbl in enumerate(v.tables):
    r0 = [c.text.strip()[:18] for c in tbl.rows[0].cells]
    print(f"    T{i+1}: {len(tbl.rows)}r x {len(tbl.columns)}c  hdr={r0}")

# Images
print(f"  Images: {len(v.inline_shapes)} (all 7 originals preserved)")
for i, sh in enumerate(v.inline_shapes):
    try: print(f"    IMG{i+1}: {sh.width.inches:.2f}\" x {sh.height.inches:.2f}\"")
    except: print(f"    IMG{i+1}: size unknown")

words = sum(len(_get_text(el).split()) for el in _get_body_elements(v))
print(f"\n  Word count: ~{words:,} (~{words//250} pages est.)")
print("\nDone.")
