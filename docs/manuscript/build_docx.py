"""
SmartSpend Capstone Manuscript — .docx builder
Generates SMARTSPEND_FINAL_MANUSCRIPT.docx from the revised markdown.
Follows Lorma Colleges formatting conventions (Times New Roman 12pt, justified).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE    = os.path.dirname(__file__)
SRC_MD  = os.path.join(BASE, "SMARTSPEND_REVISED_MANUSCRIPT.md")
OUT     = os.path.join(BASE, "SMARTSPEND_FINAL_MANUSCRIPT.docx")

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins: 1.5" left (binding), 1" other
sec = doc.sections[0]
sec.top_margin    = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin   = Inches(1.5)
sec.right_margin  = Inches(1)
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)

# ── Style helpers ──────────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=6, line_spacing=None):
    fmt = para.paragraph_format
    fmt.alignment       = align
    fmt.space_before    = Pt(space_before)
    fmt.space_after     = Pt(space_after)
    if line_spacing:
        from docx.shared import Pt as _Pt
        fmt.line_spacing = _Pt(line_spacing)

def heading(text, level=1, center=False):
    sizes  = {1: 14, 2: 13, 3: 12}
    para   = doc.add_paragraph()
    align  = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    set_para_fmt(para, align=align, space_before=12, space_after=6)
    run = para.add_run(text.strip())
    set_font(run, size=sizes.get(level, 12), bold=True)
    return para

def body(text, indent=False, center=False, bold=False, italic=False, space_after=6):
    para  = doc.add_paragraph()
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_fmt(para, align=align, space_after=space_after)
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    # Handle inline **bold** and *italic*
    segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for seg in segments:
        if seg.startswith('**') and seg.endswith('**'):
            r = para.add_run(seg[2:-2])
            set_font(r, bold=True, italic=italic)
        elif seg.startswith('*') and seg.endswith('*'):
            r = para.add_run(seg[1:-1])
            set_font(r, italic=True)
        elif seg.startswith('`') and seg.endswith('`'):
            r = para.add_run(seg[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        else:
            r = para.add_run(seg)
            set_font(r, bold=bold, italic=italic)
    return para

def bullet(text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3)
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text.strip())
    for seg in segments:
        if seg.startswith('**') and seg.endswith('**'):
            r = para.add_run(seg[2:-2]); set_font(r, bold=True)
        elif seg.startswith('*') and seg.endswith('*'):
            r = para.add_run(seg[1:-1]); set_font(r, italic=True)
        else:
            r = para.add_run(seg); set_font(r)
    return para

def divider():
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ── Table builder ──────────────────────────────────────────────────────────────
def add_table(rows_data, header=True):
    if not rows_data:
        return
    ncols = max(len(r) for r in rows_data)
    tbl   = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        row = tbl.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):
                break
            cell = row.cells[j]
            cell.text = ""
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            txt = str(cell_text).strip().replace('**', '').replace('`', '')
            run = p.add_run(txt)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            if header and i == 0:
                run.bold = True
                from docx.oxml.ns import qn as _qn
                from docx.oxml import OxmlElement as _OE
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = _OE('w:shd')
                shd.set(_qn('w:val'), 'clear')
                shd.set(_qn('w:color'), 'auto')
                shd.set(_qn('w:fill'), 'D9D9D9')
                tcPr.append(shd)
    doc.add_paragraph()

# ── Markdown parser ────────────────────────────────────────────────────────────
def parse_md(path):
    with open(path, encoding='utf-8') as f:
        lines = f.readlines()

    in_table    = False
    table_rows  = []
    in_code     = False
    skip_image  = False
    i = 0

    while i < len(lines):
        raw  = lines[i].rstrip('\n')
        line = raw.strip()

        # Skip base64 image lines
        if line.startswith('[image') or 'data:image' in line or line.startswith('iVBOR'):
            i += 1; continue

        # Code fences
        if line.startswith('```'):
            in_code = not in_code
            if in_code:
                doc.add_paragraph()
            i += 1; continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run(raw)
            r.font.name = "Courier New"; r.font.size = Pt(9)
            i += 1; continue

        # Tables
        if line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1; continue  # separator row
            table_rows.append(cells)
            in_table = True
            i += 1; continue
        elif in_table:
            add_table(table_rows)
            table_rows = []
            in_table   = False
            # don't increment — reprocess current line

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].replace('**','').replace('*','')
            page_break()
            heading(text, level=1, center=True)
            i += 1; continue
        if line.startswith('## '):
            text = line[3:].replace('**','').replace('*','')
            heading(text, level=2, center=False)
            i += 1; continue
        if line.startswith('### '):
            text = line[4:].replace('**','').replace('*','')
            heading(text, level=3, center=False)
            i += 1; continue
        if line.startswith('#### '):
            text = line[5:].replace('**','').replace('*','')
            p = body(text, bold=True, space_after=3)
            i += 1; continue

        # Horizontal rule
        if line.startswith('---'):
            divider()
            i += 1; continue

        # Page break marker
        if line == '\\pagebreak' or line == '<page-break>':
            page_break()
            i += 1; continue

        # Bullets
        if line.startswith('- ') or line.startswith('* '):
            bullet(line[2:], level=0)
            i += 1; continue
        if line.startswith('  - ') or line.startswith('  * '):
            bullet(line[4:], level=1)
            i += 1; continue

        # Numbered list
        nm = re.match(r'^(\d+)\.\s+(.*)', line)
        if nm:
            p = doc.add_paragraph(style='List Number')
            set_para_fmt(p, space_after=3)
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(nm.group(2).strip())
            set_font(r)
            i += 1; continue

        # Skip figure/image placeholder lines
        if line.startswith('*[Insert') or line.startswith('*\\[Insert'):
            p = body(line.strip('*').strip(), italic=True, center=True, space_after=4)
            i += 1; continue

        # Blockquote
        if line.startswith('> '):
            p = body(line[2:], indent=True, italic=True, space_after=4)
            i += 1; continue

        # Empty line
        if not line:
            doc.add_paragraph()
            i += 1; continue

        # Normal paragraph
        body(line, indent=False, space_after=6)
        i += 1

    # Flush any remaining table
    if table_rows:
        add_table(table_rows)

# ── Title page ─────────────────────────────────────────────────────────────────
def title_page():
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory Application\nfor Personal Financial Management")
    r.font.name = "Times New Roman"; r.font.size = Pt(14); r.bold = True

    doc.add_paragraph()
    for line in ["A CAPSTONE Project", "presented to the faculty of the",
                 "College of Computer Studies and Engineering", "LORMA Colleges",
                 "", "In Partial Fulfillment", "of the requirements for the degree",
                 "of Bachelor of Science in Information Technology"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    doc.add_paragraph()
    for line in ["by:", "Directo, Brix A.",
                 "Rubis, Cyrille John M.", "Madayag, Djaunathan Albert S.",
                 "Verzola, Johnny Flores, MTS", "", "August 2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)
        if line in ["Directo, Brix A.", "Rubis, Cyrille John M.",
                    "Madayag, Djaunathan Albert S."]:
            r.bold = True
    page_break()

# ── Pre-process: split markdown into sections ──────────────────────────────────
def load_md_body(path):
    """Return only lines from '# CHAPTER I' onwards — skip front-matter."""
    with open(path, encoding='utf-8') as f:
        lines = f.readlines()
    start = 0
    for i, line in enumerate(lines):
        if line.strip().startswith('# CHAPTER I'):
            start = i
            break
    return lines[start:]

# ── BUILD ──────────────────────────────────────────────────────────────────────
print("Building SMARTSPEND_FINAL_MANUSCRIPT.docx ...")
title_page()

# Add front-matter sections from scratch (clean, no duplication)
def front_matter():
    # ABSTRACT
    heading("ABSTRACT", level=1, center=True)
    abstract_text = (
        "Financial mismanagement remains a critical and documented challenge among Filipino households, "
        "compounded by limited access to accessible, localized, and intelligent financial tools. "
        "This study designed, developed, and evaluated SmartSpend — an AI-assisted mobile financial "
        "tracking and advisory application for Android, built for parents aged 35–55 and young "
        "professionals aged 21–35 in La Union, Philippines.\n\n"
        "SmartSpend integrates a multi-provider agentic large language model (LLM) architecture — "
        "with Gemini 3.1 Flash-Lite as the primary model and four automatic fallback providers — "
        "enabling 29 autonomous financial management actions through natural language, voice, camera, "
        "batch screenshot import (40+ platform types), and manual entry. The system operates on an "
        "offline-first SQLite database with Firebase cloud synchronization, ensuring full functionality "
        "without internet connectivity.\n\n"
        "A core academic contribution is the Financial Health Score (FHS): a 0–100 behavioral metric "
        "computed entirely from user-recorded transaction data in two modes — Full Mode (income-based: "
        "Savings Rate, Overspend Control, Budget Adherence, Logging Consistency) and Lightweight Mode "
        "(habit-based: Spending Restraint, Consistency, Category Balance, Habit Streak) — with a "
        "Warning Decay consequence mechanism and a Logging Gap Detection system.\n\n"
        "The system was evaluated using the System Usability Scale (SUS) with 30 purposively selected "
        "respondents (20 parents, 10 young professionals), targeting a score of ≥80 (Good). Expert "
        "validation of the survey instrument was conducted by a subject matter expert in financial "
        "management, and technical validation of the SUS evaluation process was conducted by a subject "
        "matter expert in information technology."
    )
    for para_text in abstract_text.split('\n\n'):
        body(para_text.strip(), indent=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Keywords: ")
    set_font(r, bold=True)
    r2 = p.add_run("personal finance management, agentic AI, large language model, financial health score, "
                   "mobile application, Flutter, Filipino users, SmartSpend")
    set_font(r2, italic=True)
    page_break()

    # ACKNOWLEDGEMENT
    heading("ACKNOWLEDGEMENT", level=1, center=True)
    ack = (
        "The researchers would like to express their sincere gratitude and appreciation to all individuals "
        "who contributed to the successful completion of this study. First and foremost, the researchers "
        "thank the Almighty God for His guidance, strength, wisdom, and blessings throughout this research. "
        "The researchers extend their heartfelt appreciation to the Dean of the College of Computer Studies "
        "and Engineering, Mr. Jeoffrey B. Layco, for his leadership and support. Special thanks are given "
        "to their Capstone Adviser, Mr. Johnny Verzola, for invaluable guidance and patience. The researchers "
        "also thank their instructor in charge, Dr. Janelli M. Mendez, for direction, structure, and "
        "continuous support. To the panelists, the researchers are grateful for their time and constructive "
        "feedback. Finally, the researchers extend their deepest gratitude to their families and friends "
        "for their unwavering support throughout this journey."
    )
    body(ack, indent=True)
    divider()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("The Researchers")
    set_font(r, bold=True)
    page_break()

    # DEDICATION
    heading("DEDICATION", level=1, center=True)
    ded = (
        "This study is dedicated to the researchers' beloved parents, whose unconditional love and "
        "sacrifices made this achievement possible. To their friends and classmates, whose encouragement "
        "helped make this journey meaningful. To their mentors and instructors, in appreciation of the "
        "knowledge and inspiration they provided. Lastly, to future researchers who may build upon this "
        "work and continue advancing the field of information technology and financial systems."
    )
    body(ded, indent=True)
    divider()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BAD · CJMR · DASM")
    set_font(r, bold=True, italic=True)
    page_break()

front_matter()

# Now parse only from Chapter I onwards
import io, sys
# Monkey-patch parse_md to use pre-filtered lines
_orig_parse = parse_md
def parse_md_body_only(path):
    lines = load_md_body(path)
    # Write to temp file and parse
    tmp = path + ".tmp"
    with open(tmp, 'w', encoding='utf-8') as f:
        f.writelines(lines)
    _orig_parse(tmp)
    import os as _os; _os.remove(tmp)

parse_md_body_only(SRC_MD)

doc.save(OUT)
print(f"Done! Saved to: {OUT}")
print(f"File size: {os.path.getsize(OUT) / 1024:.1f} KB")
