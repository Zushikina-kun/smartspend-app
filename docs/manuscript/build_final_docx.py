"""
build_final_docx.py  —  v4  CLEAN REBUILD
==========================================
Instead of mutating the source doc (fragile), we:
1. Open the source doc (SMARTSPEND_CAPSTONE_WORKING.docx)
2. Open a BLANK doc and copy settings/styles from source
3. Append sections in EXACT correct order:
   [TITLE] → [APPROVAL] → [ABSTRACT] → [ACK] → [DED] →
   [TOC] → [LOF] → [LOT] →
   [CH1 original + updated] → [CH2 original + updated] →
   [CH3 new] → [CH4 new] → [REFERENCES new] →
   [APPENDICES] → [APP-A new] → [APP-B Survey] → [APP-C Consent] → [APP-D SUS] →
   [CV all 3]

Source element ranges (from structure audit):
  Title page:     0 – 28
  Ack:           30 – 51
  Dedication:    52 – 60
  ToC:           61 – 98  (includes old Refs/Appendices TOC lines — we rebuild this)
  Ch1 body:      99 – 178  (includes images, tables)
  Ch2 body:     179 – 244
  Old Refs:     245 – 289  (SKIP — replaced by new references)
  APPENDICES:   290 – 290
  Old App-A Survey: 291 – 374
  Old App-B Consent: 375 – 401
  Old App-C SUS: 402 – 421
  CV + photos:  422 – 512

Outputs: SMARTSPEND_FINAL_V4.docx
"""

import copy, os
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

BASE = os.path.dirname(__file__)
SRC  = os.path.join(BASE, "SMARTSPEND_CAPSTONE_WORKING.docx")
OUT  = os.path.join(BASE, "SMARTSPEND_FINAL_V4.docx")
NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Load source ────────────────────────────────────────────────────────────────
src = Document(SRC)
src_body = list(src.element.body)   # list of all XML children

def src_els(start, end):
    """Return deep copies of source body elements [start..end] inclusive."""
    return [copy.deepcopy(src_body[i]) for i in range(start, end + 1)]

def get_text(el):
    return "".join(t.text or "" for t in el.findall(".//{%s}t" % NS)).strip()

# ── Create output doc from scratch, copy page setup from source ────────────────
out = Document()
# Copy page setup
src_sec = src.sections[0]
out_sec = out.sections[0]
out_sec.page_width    = src_sec.page_width
out_sec.page_height   = src_sec.page_height
out_sec.top_margin    = src_sec.top_margin
out_sec.bottom_margin = src_sec.bottom_margin
out_sec.left_margin   = src_sec.left_margin
out_sec.right_margin  = src_sec.right_margin
out_sec.header_distance = src_sec.header_distance
out_sec.footer_distance = src_sec.footer_distance

# Remove the default empty paragraph Word adds to a new Document
for child in list(out.element.body):
    out.element.body.remove(child)

body = out.element.body

def append(el):
    body.append(copy.deepcopy(el) if el.getparent() is not None else el)

def append_many(elements):
    for el in elements:
        body.append(el)

def pb():
    """Page break element."""
    return etree.fromstring(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:br w:type="page"/></w:r></w:p>'
    )

FONT = "Tahoma"

def _p(text, bold=False, italic=False, align="both", fi=720, ls=480, size=12):
    sz  = int(size * 2)
    b   = "<w:b/><w:bCs/>" if bold   else ""
    i   = "<w:i/><w:iCs/>" if italic else ""
    fi_ = f'<w:ind w:firstLine="{fi}"/>' if fi else ""
    ls_ = f'<w:spacing w:line="{ls}" w:lineRule="auto"/>' if ls else ""
    t   = (str(text).replace("&","&amp;").replace("<","&lt;")
                    .replace(">","&gt;").replace('"',"&quot;"))
    return etree.fromstring(
        f'<w:p xmlns:w="{NS}"><w:pPr><w:jc w:val="{align}"/>{ls_}{fi_}</w:pPr>'
        f'<w:r><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>{b}{i}</w:rPr>'
        f'<w:t xml:space="preserve">{t}</w:t></w:r></w:p>'
    )

def _blank():   return _p("", ls=480, fi=0)
def _hdr(t):    return _p(t, bold=True,  align="center", fi=0, ls=276)
def _subhdr(t): return _p(t, bold=True,  align="both",   fi=0, ls=480)
def _body(t, italic=False):  return _p(t, italic=italic, align="both", fi=720, ls=480)
def _body0(t, bold=False):   return _p(t, bold=bold, align="both", fi=0, ls=480)
def _right(t):  return _p(t, bold=True, align="right", fi=0, ls=480)

def _table(rows, col_pcts=None):
    ncols = max(len(r) for r in rows)
    tbl   = OxmlElement("w:tbl")
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle"); tblStyle.set(qn("w:val"), "TableGrid")
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), "9350"); tblW.set(qn("w:type"), "dxa")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000")
        tblBorders.append(b)
    tblPr.append(tblStyle); tblPr.append(tblW); tblPr.append(tblBorders)
    tbl.append(tblPr)
    tblGrid = OxmlElement("w:tblGrid")
    pcts = col_pcts or [100//ncols]*ncols
    for pct in pcts:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(int(9350*pct/100)))
        tblGrid.append(col)
    tbl.append(tblGrid)
    for ri, row in enumerate(rows):
        tr = OxmlElement("w:tr")
        is_hdr = (ri == 0)
        for ci in range(ncols):
            ct = row[ci] if ci < len(row) else ""
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(9350//ncols)); tcW.set(qn("w:type"),"dxa"); tcPr.append(tcW)
            if is_hdr:
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"D9D9D9"); tcPr.append(shd)
            tc.append(tcPr)
            t_esc = (str(ct).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;"))
            bx = "<w:b/><w:bCs/>" if is_hdr else ""
            tc.append(etree.fromstring(
                f'<w:p xmlns:w="{NS}"><w:pPr><w:spacing w:line="240" w:lineRule="auto"/></w:pPr>'
                f'<w:r><w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
                f'<w:sz w:val="20"/><w:szCs w:val="20"/>{bx}</w:rPr>'
                f'<w:t xml:space="preserve">{t_esc}</w:t></w:r></w:p>'
            ))
            tr.append(tc)
        tbl.append(tr)
    return tbl

# ── Patch source element in-place (text-only, no structural change) ────────────
def patch_text(el, old_text, new_text):
    for wt in el.findall(".//{%s}t" % NS):
        if wt.text and old_text in wt.text:
            wt.text = wt.text.replace(old_text, new_text)

# Apply patches to source elements before copying
# Patch title date
for el in src_body:
    if "April 2026" in get_text(el):
        patch_text(el, "April 2026", "August 2026")
        break
# Patch objective statement
for el in src_body:
    t = get_text(el)
    if "parents aged 35 to 55 and young professionals aged 21 to 35 in La Union, Philippines." in t:
        patch_text(el,
            "parents aged 35 to 55 and young professionals aged 21 to 35 in La Union, Philippines.",
            "parents aged 35 to 55 as the primary target population, and young professionals aged 21 to 35 as a secondary demographic, in La Union, Philippines.")
# Patch API key statement
for el in src_body:
    t = get_text(el)
    if "LLM API key is embedded within the application package" in t:
        patch_text(el,
            "The LLM API key is embedded within the application package as part of the free-tier academic deployment. To mitigate potential misuse, the system enforces a daily interaction limit of 60 AI requests per user, which resets automatically. A backend proxy server for secure API key management is planned as a post-capstone enhancement to align with production-level security practices.",
            "The LLM API key is not embedded within the application package. Instead, it is fetched securely at runtime via Firebase Remote Config, ensuring the key is never stored in the APK binary or exposed in the source code repository. To mitigate potential misuse, the system enforces a daily interaction limit of 60 AI requests per user, which resets automatically. A backend proxy server for fully server-side API key management is planned as a post-capstone enhancement.")
        break
# Patch backup version
for el in src_body:
    for wt in el.findall(".//{%s}t" % NS):
        if wt.text and "version 8 format" in wt.text:
            wt.text = wt.text.replace("version 8 format", "version 9 format")
# Patch badge count
for el in src_body:
    for wt in el.findall(".//{%s}t" % NS):
        if wt.text and "16 achievement badges" in wt.text:
            wt.text = wt.text.replace("16 achievement badges", "23 achievement badges")

# Find Table 1.2 in source and replace it
for i, el in enumerate(src_body):
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag == "tbl":
        cells = el.findall(".//{%s}tc" % NS)
        if cells and "Feature" in get_text(cells[0]) and "Tarsi" in "".join(get_text(c) for c in cells[:4]):
            src_body[i] = _table([
                ["Feature","Tarsi","YNAB","Monarch","Copilot","BudgetPH","Alkansya AI","GCash Pera Coach","SmartSpend"],
                ["Offline mode","Yes","No","No","No","Yes","No","No","Yes"],
                ["LLM chat assistant","No","No","No","No","Insights only","Yes (basic)","Yes (literacy Q&A)","Yes — 29 agentic actions"],
                ["Financial Health Score","No","No","No","No","Yes (simpler)","No","No","Yes (0–100, dual-mode)"],
                ["OCR receipt scanning","Yes","No","No","No","No","No","No","Yes"],
                ["Voice input (en-PH)","No","No","No","No","No","No","No","Yes"],
                ["Batch screenshot import (40+)","No","No","No","No","No","No","No","Yes"],
                ["Multi-period spending limits","No","No","No","No","Yes","No","No","Yes (daily/wk/mo/yr)"],
                ["Spending behavior analysis","No","No","No","No","No","No","No","Yes"],
                ["Bank synchronization","No","Yes","Yes","Yes","CSV import","No","GCash balance","No"],
                ["Filipino-English AI (Taglish)","No","No","No","No","No","Yes","Yes (PH languages)","Yes — full Taglish"],
                ["Free tier (Android)","Yes","No","No","No","Yes (PWA)","Limited","Yes (GCash req.)","Yes — always free"],
                ["Paluwagan tracker","No","No","No","No","Yes","No","No","No"],
                ["15th & 30th payday cycle","No","No","No","No","Yes","No","No","No"],
                ["Gamification (badges/quests)","No","No","No","No","Yes (XP/levels)","No","No","Yes (23 badges, 10 quests)"],
                ["Logging gap detection","No","No","No","No","No","No","No","Yes"],
                ["SSS/PhilHealth/Pag-IBIG","No","No","No","No","Yes (records)","No","No","Yes (AI compute)"],
                ["Insurance tracker","No","No","No","No","No","No","No","Yes"],
                ["Round-up savings","No","No","No","No","No","No","No","Yes"],
            ], col_pcts=[20,7,6,7,7,8,9,11,15])
            break

# Find Table 2.1 in source and replace it
for i, el in enumerate(src_body):
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag == "tbl":
        cells = el.findall(".//{%s}tc" % NS)
        if cells and "Category" in get_text(cells[0]) and len(cells) >= 2 and "Number" in get_text(cells[1]):
            src_body[i] = _table([
                ["Category","Number of Respondents"],
                ["Parents (Ages 35–55)","20"],
                ["Young Professionals (Ages 21–35)","10"],
                ["Total Respondents","30"],
                ["",""],
                ["Validators / Experts","Number of Validators"],
                ["Content Validator (Survey — Financial Management Expert)","1"],
                ["Technical Validator (System & SUS Process — IT Expert)","1"],
                ["Total Validators","2"],
            ], col_pcts=[75, 25])
            break

# ══════════════════════════════════════════════════════════════════════════════
# SECTIONS — import content blocks
# ══════════════════════════════════════════════════════════════════════════════
# Import content blocks from update_working_docx.py
import importlib.util, sys
spec = importlib.util.spec_from_file_location("content_blocks",
       os.path.join(BASE, "update_working_docx.py"))
# We can't easily import because update_working_docx runs on import.
# Instead, redefine the blocks we need inline here.
# (They reference _p, _blank, _body, _table, etc. already defined above)

def approval_sheet():
    return [
        pb(),
        _hdr("APPROVAL SHEET"), _blank(),
        _body("This is to certify that we have supervised the preparation of the Capstone Project and read the manuscript prepared by DIRECTO, BRIX A., RUBIS, CYRILLE JOHN M., and MADAYAG, DJAUNATHAN ALBERT S. entitled SMARTSPEND: AN AI-ASSISTED MOBILE FINANCIAL TRACKING AND ADVISORY APPLICATION FOR PERSONAL FINANCIAL MANAGEMENT and that the said capstone project has been submitted for final examination by the Oral Examination Committee."),
        _blank(),
        _p("_______________________________", bold=False, align="center", fi=0, ls=240),
        _p("Johnny F. Verzola, MTS",         bold=True,  align="center", fi=0, ls=240),
        _p("Capstone Project Adviser",        bold=False, align="center", fi=0, ls=240),
        _blank(),
        _body("As members of the Oral Examination Committee, we certify that we have examined this capstone project presented before the committee and hereby recommend that it be accepted in partial fulfillment of the capstone requirements for the degree in Bachelor of Science in Information Technology."),
        _blank(),
        _p("_______________________________",                   bold=False, align="center", fi=0, ls=240),
        _p("Ellen F. Mangaoang, MIT",                           bold=True,  align="center", fi=0, ls=240),
        _p("Chairperson",                                       bold=False, align="center", fi=0, ls=240),
        _p("_____________________     _____________________",  bold=False, align="center", fi=0, ls=240),
        _p("Jopher F. Reyes, MIT       Gelo Ryann M. Carbonell",bold=True,  align="center", fi=0, ls=240),
        _p("Member                            Member",          bold=False, align="center", fi=0, ls=240),
        _blank(),
        _body("This capstone project is hereby approved and accepted by the College of Computer Studies and Engineering in partial fulfillment of the requirements for the degree in Bachelor of Science in Information Technology."),
        _blank(),
        _p("_______________________________", bold=False, align="center", fi=0, ls=240),
        _p("Jeoffrey B. Layco, MIS",          bold=True,  align="center", fi=0, ls=240),
        _p("Dean, CCSE",                       bold=False, align="center", fi=0, ls=240),
    ]

def abstract_section():
    return [
        pb(), _hdr("CAPSTONE PROJECT ABSTRACT"), _blank(),
        _p("Title:  SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory Application for Personal Financial Management", bold=False, align="both", fi=0, ls=480),
        _p("Researchers:", bold=False, align="both", fi=0, ls=480),
        _p("        Directo, Brix A.",             bold=True, align="both", fi=0, ls=480),
        _p("        Rubis, Cyrille John M.",        bold=True, align="both", fi=0, ls=480),
        _p("        Madayag, Djaunathan Albert S.", bold=True, align="both", fi=0, ls=480),
        _p("Type of Document:  CAPSTONE PROJECT",  bold=False, align="both", fi=0, ls=480),
        _p("Type of Publication:  Unpublished",     bold=False, align="both", fi=0, ls=480),
        _p("Accrediting Institution:  Lorma Colleges, CLI Bldg., San Juan Campus, La Union", bold=False, align="both", fi=0, ls=480),
        _p("Teacher-in-Charge:  Janelli M. Mendez, DIT", bold=False, align="both", fi=0, ls=480),
        _blank(),
        _hdr("ABSTRACT"),
        _body("Financial mismanagement remains a critical and documented challenge among Filipino households, compounded by limited access to accessible, localized, and intelligent financial tools. This study designed, developed, and evaluated SmartSpend — an AI-assisted mobile financial tracking and advisory application for Android, built for parents aged 35–55 and young professionals aged 21–35 in La Union, Philippines.", italic=True),
        _body("SmartSpend integrates a multi-provider agentic large language model (LLM) architecture — with Gemini 3.1 Flash-Lite as the primary model and four automatic fallback providers — enabling 29 autonomous financial management actions through natural language, voice, camera, batch screenshot import (40+ platform types), and manual entry. The system operates on an offline-first SQLite database with Firebase cloud synchronization.", italic=True),
        _body("A core academic contribution is the Financial Health Score (FHS): a 0–100 behavioral metric computed in two modes — Full Mode (Savings Rate, Overspend Control, Budget Adherence, Logging Consistency) and Lightweight Mode (Spending Restraint, Consistency, Category Balance, Habit Streak) — with Warning Decay and Logging Gap Detection mechanisms.", italic=True),
        _body("The system was evaluated using the System Usability Scale (SUS) with 30 purposively selected respondents (20 parents, 10 young professionals), targeting a score of ≥80 (Good). Expert validation was conducted by subject matter experts in financial management and information technology.", italic=True),
        _p("Keywords:  personal finance management, agentic AI, large language model, financial health score, mobile application, Flutter, Filipino users, SmartSpend", bold=True, italic=True, align="both", fi=0, ls=480),
    ]

def list_of_figures():
    return [
        pb(), _hdr("LIST OF FIGURES"), _blank(),
        _p("Figure 1.1.  Financial Literacy Rates by Demographic Group (BSP, 2021; Inquiro, 2024)\t\t3", align="both", fi=0, ls=276),
        _p("Figure 1.2.  Conceptual Framework — SmartSpend Mobile Application (IPO Model)\t\t\t11", align="both", fi=0, ls=276),
        _p("Figure 2.1.  System Usability Scale (SUS) Score Interpretation\t\t\t\t\t34", align="both", fi=0, ls=276),
        _p("Figure 2.2.  Agile Kanban Workflow for SmartSpend Development\t\t\t\t\t35", align="both", fi=0, ls=276),
    ]

def list_of_tables():
    return [
        pb(), _hdr("LIST OF TABLES"), _blank(),
        _p("Table 1.1.  Three-Level Gap Analysis of SmartSpend\t\t\t\t\t\t4",  align="both", fi=0, ls=276),
        _p("Table 1.2.  Feature Comparison of SmartSpend with Existing Financial Applications\t\t6",  align="both", fi=0, ls=276),
        _p("Table 2.1.  Distribution of Respondents and Validators\t\t\t\t\t\t30", align="both", fi=0, ls=276),
        _p("Table 2.2.  Comparative Evaluation of LLM APIs for SmartSpend Integration\t\t\t44", align="both", fi=0, ls=276),
        _p("Table 2.3.  Agile Kanban Workflow Phases and Deliverables\t\t\t\t\t38", align="both", fi=0, ls=276),
    ]

def toc_section():
    """Rebuild TOC completely — no stray tables or misplaced content."""
    return [
        pb(), _hdr("TABLE OF CONTENTS"), _blank(),
        _p("",                                                                     bold=False, align="both", fi=0, ls=276),
        _p("TITLE PAGE\t\t\t\t\t\t\t\t\t\t\t\t\t\ti",                          bold=True,  align="both", fi=0, ls=276),
        _p("APPROVAL SHEET\t\t\t\t\t\t\t\t\t\t\t\t\tii",                        bold=True,  align="both", fi=0, ls=276),
        _p("ABSTRACT\t\t\t\t\t\t\t\t\t\t\t\t\t\tiii",                          bold=True,  align="both", fi=0, ls=276),
        _p("ACKNOWLEDGEMENT\t\t\t\t\t\t\t\t\t\t\t\tiv",                         bold=True,  align="both", fi=0, ls=276),
        _p("DEDICATION\t\t\t\t\t\t\t\t\t\t\t\t\t\tv",                           bold=True,  align="both", fi=0, ls=276),
        _p("TABLE OF CONTENTS\t\t\t\t\t\t\t\t\t\t\t\tvi",                       bold=True,  align="both", fi=0, ls=276),
        _p("LIST OF FIGURES\t\t\t\t\t\t\t\t\t\t\t\t\tvii",                      bold=True,  align="both", fi=0, ls=276),
        _p("LIST OF TABLES\t\t\t\t\t\t\t\t\t\t\t\t\tviii",                      bold=True,  align="both", fi=0, ls=276),
        _blank(),
        _p("CHAPTERS",                                                             bold=True,  align="both", fi=0, ls=276),
        _p("I\tINTRODUCTION\t\t\t\t\t\t\t\t\t\t\t1",                            bold=True,  align="both", fi=0, ls=276),
        _p("\tProject Context\t\t\t\t\t\t\t\t\t\t\t1",                           bold=False, align="both", fi=720, ls=276),
        _p("\t\tFigure 1.1. Financial Literacy Rates (BSP, 2021; Inquiro, 2024)\t3",bold=False,align="both", fi=1440, ls=276),
        _p("\t\tTable 1.1. Three-Level Gap Analysis of SmartSpend\t\t\t4",        bold=False, align="both", fi=1440, ls=276),
        _p("\t\tTable 1.2. Feature Comparison of SmartSpend with Existing Apps\t6",bold=False, align="both", fi=1440, ls=276),
        _p("\tConceptual Framework\t\t\t\t\t\t\t\t\t\t11",                       bold=False, align="both", fi=720, ls=276),
        _p("\t\tFigure 1.2. Conceptual Framework (IPO Model)\t\t\t\t11",          bold=False, align="both", fi=1440, ls=276),
        _p("\tStatement of Objectives\t\t\t\t\t\t\t\t\t\t12",                   bold=False, align="both", fi=720, ls=276),
        _p("\tScope and Limitations of the Study\t\t\t\t\t\t\t\t13",             bold=False, align="both", fi=720, ls=276),
        _p("\tPurpose and Description\t\t\t\t\t\t\t\t\t\t15",                   bold=False, align="both", fi=720, ls=276),
        _p("\t\tBeneficiaries of the Study\t\t\t\t\t\t\t\t17",                  bold=False, align="both", fi=1440, ls=276),
        _p("\tTechnical Background\t\t\t\t\t\t\t\t\t\t\t19",                    bold=False, align="both", fi=720, ls=276),
        _blank(),
        _p("II\tDESIGN AND METHODOLOGY\t\t\t\t\t\t\t\t\t26",                    bold=True,  align="both", fi=0, ls=276),
        _p("\tResearch Design\t\t\t\t\t\t\t\t\t\t\t26",                         bold=False, align="both", fi=720, ls=276),
        _p("\tPopulation and Locale\t\t\t\t\t\t\t\t\t\t28",                     bold=False, align="both", fi=720, ls=276),
        _p("\t\tTable 2.1. Distribution of Respondents and Validators\t\t\t30",  bold=False, align="both", fi=1440, ls=276),
        _p("\tEthical Considerations\t\t\t\t\t\t\t\t\t\t31",                    bold=False, align="both", fi=720, ls=276),
        _p("\tData Gathering Tools and Procedure\t\t\t\t\t\t\t\t32",             bold=False, align="both", fi=720, ls=276),
        _p("\t\tFigure 2.1. SUS Score Interpretation\t\t\t\t\t\t34",             bold=False, align="both", fi=1440, ls=276),
        _p("\tSoftware Methodology\t\t\t\t\t\t\t\t\t\t35",                      bold=False, align="both", fi=720, ls=276),
        _p("\t\tFigure 2.2. Agile Kanban Workflow\t\t\t\t\t\t\t35",              bold=False, align="both", fi=1440, ls=276),
        _p("\t\tTable 2.3. Kanban Phases and Deliverables\t\t\t\t\t38",          bold=False, align="both", fi=1440, ls=276),
        _blank(),
        _p("III\tRESULTS AND DISCUSSION\t\t\t\t\t\t\t\t\t40",                   bold=True,  align="both", fi=0, ls=276),
        _p("\tObjective 1 — Assessment of Financial Management Practices\t\t40",  bold=False, align="both", fi=720, ls=276),
        _p("\tObjective 2 — System Development and LLM Benchmarking\t\t\t42",    bold=False, align="both", fi=720, ls=276),
        _p("\t\tTable 2.2. Comparative Evaluation of LLM APIs\t\t\t\t44",        bold=False, align="both", fi=1440, ls=276),
        _p("\t\tFinancial Health Score — Full Computation\t\t\t\t\t48",           bold=False, align="both", fi=1440, ls=276),
        _p("\t\tTable 2.3. Kanban Phases and Deliverables\t\t\t\t\t38",          bold=False, align="both", fi=1440, ls=276),
        _p("\tObjective 3 — System Usability Evaluation (SUS)\t\t\t\t53",         bold=False, align="both", fi=720, ls=276),
        _blank(),
        _p("IV\tCONCLUSIONS AND RECOMMENDATIONS\t\t\t\t\t\t\t56",               bold=True,  align="both", fi=0, ls=276),
        _p("\tConclusions\t\t\t\t\t\t\t\t\t\t\t56",                             bold=False, align="both", fi=720, ls=276),
        _p("\tRecommendations\t\t\t\t\t\t\t\t\t\t\t57",                         bold=False, align="both", fi=720, ls=276),
        _blank(),
        _p("REFERENCES\t\t\t\t\t\t\t\t\t\t\t\t\t59",                           bold=True,  align="both", fi=0, ls=276),
        _blank(),
        _p("APPENDICES\t\t\t\t\t\t\t\t\t\t\t\t\t66",                           bold=True,  align="both", fi=0, ls=276),
        _p("\tAppendix A — Validation Certificates\t\t\t\t\t\t\t66",             bold=False, align="both", fi=720, ls=276),
        _p("\tAppendix B — Survey Questionnaire\t\t\t\t\t\t\t68",                bold=False, align="both", fi=720, ls=276),
        _p("\tAppendix C — Consent Form\t\t\t\t\t\t\t\t\t74",                   bold=False, align="both", fi=720, ls=276),
        _p("\tAppendix D — SUS Questionnaire\t\t\t\t\t\t\t\t76",                bold=False, align="both", fi=720, ls=276),
        _blank(),
        _p("CURRICULUM VITAE\t\t\t\t\t\t\t\t\t\t\t\t78",                      bold=True,  align="both", fi=0, ls=276),
    ]

def appendix_a():
    return [
        pb(), _hdr("APPENDIX A"), _hdr("VALIDATION CERTIFICATES"), _blank(),
        _subhdr("CONTENT VALIDATION CERTIFICATE — SURVEY QUESTIONNAIRE"),
        _body("This survey questionnaire has been reviewed for content validity and is deemed appropriate and relevant to the financial management experiences of the target population."),
        _blank(),
        _body0("Educational Background  :  ________________________________"),
        _body0("(e.g., BS Commerce, BS Accountancy, BS Business Administration, or equivalent)"),
        _blank(),
        _body0("Occupation  :  ________________________________"),
        _body0("(e.g., Business Owner, Financial Officer, Accountant, Financial Adviser, etc.)"),
        _blank(),
        _body0("Years of Experience  :  _____ years in financial management and/or business operations"),
        _blank(),
        _body0("Signature  :  _______________________________"),
        _body0("Name (optional)  :  _______________________________"),
        _body0("Date  :  _______________________________"),
        _blank(), _blank(),
        _subhdr("TECHNICAL VALIDATION CERTIFICATE — SYSTEM AND SUS EVALUATION"),
        _body("The SmartSpend system and its usability evaluation process have been reviewed by a subject matter expert in Information Technology to ensure technical soundness and proper SUS administration."),
        _blank(),
        _body0("Educational Background  :  ________________________________"),
        _body0("(e.g., BS Information Technology, BS Computer Science, or equivalent)"),
        _blank(),
        _body0("Occupation  :  ________________________________"),
        _body0("(e.g., Software Developer, IT Instructor, Systems Analyst, IT Professional, etc.)"),
        _blank(),
        _body0("Years of Experience  :  _____ years in IT / software development and/or usability evaluation"),
        _blank(),
        _body0("Signature  :  _______________________________"),
        _body0("Name (optional)  :  _______________________________"),
        _body0("Date  :  _______________________________"),
    ]

def chapter_three():
    els = [
        pb(), _hdr("CHAPTER III"), _hdr("RESULTS AND DISCUSSION"),
        _body("This chapter presents the results of the study based on the three stated objectives. It discusses the outcomes of each objective in relation to the development and evaluation of the SmartSpend mobile application."),
        _subhdr("Objective 1 — Assessment of Financial Management Practices"),
        _p("[NOTE: Complete after survey data collection (Week 7). Insert frequency tables and interview themes here.]", italic=True, bold=False, align="both", fi=720, ls=480),
        _body("The first objective was to assess the existing financial management practices, common budgeting challenges, and expense tracking behaviors of parents aged 35 to 55 and young professionals aged 21 to 35 in San Fernando City, La Union. Data was gathered through a validated structured survey questionnaire and supplementary interviews."),
        _body("A total of thirty (30) respondents participated — 20 parents aged 35 to 55 and 10 young professionals aged 21 to 35. [Insert Table 3.1 Respondent Profile here after data collection.] Survey results revealed that [insert findings on expense tracking methods, budgeting frequency, and financial challenges]. These findings are consistent with BSP (2021) data indicating that a large proportion of Filipino adults do not maintain formal written budgets."),
        _body("The assessment findings confirmed the presence of the financial management challenges identified in the literature review — manual effort burden, irregular tracking behavior, and lack of proactive feedback — and validated the need for an AI-assisted tool tailored to the Filipino context. These insights directly guided SmartSpend's core design: multi-modal AI input, the Financial Health Score, and the Warning Decay mechanism."),
        _subhdr("Objective 2 — System Development and LLM Benchmarking"),
        _body("The second objective was to design and develop the SmartSpend mobile application, including the selection of an appropriate Large Language Model API through comparative technical evaluation."),
        _subhdr("Comparative Analysis of Large Language Model APIs"),
        _body("The selection of an appropriate LLM API is a critical design decision because it directly influences the accuracy, latency, and cost of natural language expense parsing and conversational assistance. For a mobile financial assistant requiring Filipino-English capability and free-tier deployment, the evaluation criteria were weighted as follows: Filipino-English accuracy (25%), speed and latency (20%), tool use and JSON reliability (20%), free tier availability (15%), context window (10%), and financial reasoning quality (10%). Table 2.2 presents the comparative evaluation results."),
        _p("Table 2.2. Comparative Evaluation of LLM APIs for SmartSpend Integration", bold=True, align="center", fi=0, ls=276),
    ]
    els.append(_table([
        ["LLM / Model","Provider","Context Window","Speed (t/s)","Filipino-English","Tool Use / JSON","Free Tier","Selected?"],
        ["Gemini 3.1 Flash-Lite","Google","1,000,000","~400–600","★★★★★","★★★★★","1,000 req/day","✅ PRIMARY"],
        ["Gemini 3.5 Flash","Google","1,000,000","~200–400","★★★★★","★★★★★","250 req/day","✅ Fallback 1"],
        ["LLaMA 3.3 70B","Groq LPU","128,000","~315","★★★★☆","★★★★★","~14,400 req/day","✅ Fallback 2"],
        ["LLaMA 3.1 8B","Groq LPU","8,192","~800","★★★★☆","★★★★☆","~14,400 req/day","✅ Fallback 3"],
        ["LLaMA 3.1 70B","Cerebras WSE","128,000","~1,800","★★★★☆","★★★★☆","1M tokens/day","✅ Fallback 4"],
        ["GPT-5.6 Terra","OpenAI","1,050,000","~80–120","★★★★★","★★★★★","Paid only","❌ Cost"],
        ["Claude Fable 5","Anthropic","200,000","~70–100","★★★★★","★★★★★","Paid only","❌ Cost"],
        ["Gemini 3.7 Flash","Google","1,048,576","~300–500","★★★★★","★★★★★","Paid ($0.75/1M)","❌ No free tier"],
        ["Grok 4.6","xAI","500,000","~100–200","★★★★☆","★★★★★","Paid ($2/1M)","❌ Cost"],
        ["DeepSeek V4 Flash","DeepSeek","1,000,000","~200","★★★☆☆","★★★☆☆","$0.14/1M (5M trial)","❌ Weaker Filipino"],
        ["Qwen 3 32B","Alibaba/OpenRouter","128,000","~150–300","★★★★☆","★★★★★","Free preview","❌ Less tested"],
        ["Fin-R1 (7B)","Self-hosted","128,000","Varies","★★★☆☆","★★☆☆☆","Self-host","❌ No hosted API"],
        ["Mistral 7B","Mistral AI","32,000","~600","★★★☆☆","★★★☆☆","Self-host","❌ Poor Filipino"],
        ["GPT-4o Mini","OpenAI","128,000","~120","★★★★☆","★★★★★","No free tier","❌ Cost"],
        ["Gemma 2 9B","Google","8,192","~500","★★★☆☆","★★★☆☆","Local only","❌ No hosted API"],
    ], col_pcts=[18,12,12,9,10,10,14,15]))
    els += [
        _body("Gemini 3.1 Flash-Lite was selected as the primary model because it offers the highest free-tier request quota (1,000 req/day), the best Filipino-English multilingual performance among free-tier models, a 1-million token context window for full user financial context injection, and native function calling support essential for the 29 agentic action types (Li et al., 2024; Google, 2024f). GPT-5.6 and Claude Fable 5 are paid-only — cost-prohibitive for academic deployment with 30 respondents at 60 messages/day each."),
        _body("SmartSpend uses dynamic full-context injection rather than RAG. A typical user has 20–50 expenses, 5–10 budgets, and 3–5 goals (~1,000–5,000 tokens), fitting within any evaluated model's context window. RAG adds unnecessary vector search overhead for this small per-user dataset (Davenport & Mittal, 2022)."),
        _subhdr("Financial Health Score — Full Computation"),
        _body("The Financial Health Score (FHS) is SmartSpend's core academic contribution — a 0-to-100 behavioral metric computed from user-recorded transaction data. Its design is informed by the Financial Health Network FinHealth Score (Financial Health Network, 2021, 2026), the UNSGSA Financial Health Measurement Framework (UNSGSA, 2021), and the CFPB Financial Well-Being Scale (Consumer Financial Protection Bureau, 2017). Unlike these survey-based frameworks, SmartSpend's FHS is a behavioral computation from transaction data — no surveys or external data feeds required."),
        _body("Full Mode — Income Tracking Enabled (4 components × 25 pts = 100 maximum):"),
        _body0("     Component 1 — Savings Rate (25 pts):  Score = 25 × min(1.0, savingsRate / 0.20). The 20% target comes from the 50/30/20 budgeting rule (Warren & Tyagi, 2005)."),
        _body0("     Component 2 — Overspend Control (25 pts):  Score = 25 × (1 − overDays / activeDays). Derived from the FinHealth Score Spend pillar (Financial Health Network, 2021)."),
        _body0("     Component 3 — Budget Adherence (25 pts):  Score = 25 × (onBudgetCategories / totalBudgetCategories). No budgets set = full 25 pts. Based on zero-based budgeting theory (Ramsey, 2003)."),
        _body0("     Component 4 — Logging Consistency (25 pts):  Score = 25 × (loggedDays / activeDays). Consistent tracking reduces discretionary spending by 10–20% (Thaler & Sunstein, 2008)."),
        _body("Lightweight Mode — Income Tracking Disabled (for students, freelancers, informal workers):  Spending Restraint (25 pts) vs user-set limit;  Logging Consistency (25 pts);  Category Balance (25 pts) no single category >40%;  Habit Streak (25 pts) consecutive days, full credit at 14 days (Duhigg, 2012)."),
        _body("Score Adjustments: Warning Decay (−5 pts/day, max −15) when budget warnings are ignored — applying loss aversion theory (Kahneman & Tversky, 1979; Thaler & Sunstein, 2008). Gap Adjustment (+2 or −3 pts/day) for confirmed no-spend or unlogged-spend days — applying behavioral honesty mechanisms (Ariely, 2008). Final score clamped 0–100."),
        _p("Table 2.3. Agile Kanban Workflow Phases and Deliverables", bold=True, align="center", fi=0, ls=276),
    ]
    els.append(_table([
        ["Phase","Key Tasks","Deliverable"],
        ["Backlog","Define all features; needs survey; literature review on PH financial gaps","Prioritized feature list; literature review"],
        ["Requirements","Translate findings into specs; validate questionnaire; LLM API benchmarking","Validated questionnaire; LLM benchmarking matrix (Table 2.2)"],
        ["Design","SQLite schema (20 tables); FHS formula; UI wireframes; data flow diagrams","System architecture; database schema; FHS documentation"],
        ["Development","Build expense tracking; integrate Gemini 3.1 Flash-Lite; add OCR/voice/barcode/batch screenshots; FHS engine; Firebase sync; gamification","Functional app; 29 agentic actions operational"],
        ["Testing","LLM parsing accuracy test; SUS with 30 respondents; interviews; bug log","SUS scores; parsing observations; bug documentation"],
        ["Deployment","Build release APKs (arm64/armeabi/x86_64); prepare Demo Mode; publish GitHub Releases","Release APKs v2.9.7; project documentation"],
        ["Done/Review","Analyze SUS scores; review feedback; identify improvements; document recommendations","Final evaluation report; post-capstone roadmap"],
    ], col_pcts=[14, 52, 34]))
    els += [
        _subhdr("System Development Results — SmartSpend v2.9.7"),
        _body("SmartSpend v2.9.7 was developed across seven Kanban phases. Key specifications: Platform: Android (Flutter/Dart); Version: 2.9.7; SQLite schema: Version 11, 20 tables; APK size: 44.7 MB (arm64-v8a, release, obfuscated); AI providers: 5 (Gemini 3.1 Flash-Lite → Gemini 3.5 Flash → Groq LLaMA 3.3 70B → LLaMA 3.1 8B → Cerebras LLaMA 3.1); Agentic action types: 29; Input modalities: 6; Screenshot platforms detected: 40+; Achievement badges: 23; Daily quests: 10 rotating; Currencies: 57. All major features confirmed working on Poco X6 Pro running Android 16."),
        _subhdr("Objective 3 — System Usability Evaluation (SUS)"),
        _p("[NOTE: Complete after SUS administration with 30 respondents (Week 7). Insert SUS computation table, per-respondent scores, final average score, interpretation per Bangor et al. (2009), and qualitative feedback here. Target: ≥80 (Good).]", italic=True, bold=False, align="both", fi=720, ls=480),
        _body("The third objective was to evaluate the usability of the SmartSpend application using the System Usability Scale (SUS). The SUS was administered to thirty (30) respondents — 20 parents and 10 young professionals — following a guided live demonstration using Demo Mode with pre-loaded Filipino sample data."),
        _body("SUS scores were computed using the standard formula: for odd-numbered items, contribution = item score minus 1; for even-numbered items, contribution = 5 minus item score; sum multiplied by 2.5 to yield a 0–100 score (Brooke, 1996)."),
        _p("Overall SUS Score:  [INSERT SCORE]     Grade: [A/B/C]     Adjective: [per Bangor et al., 2009]     Target: ≥80 (Good)", bold=True, italic=True, align="center", fi=0, ls=480),
        _p("[Insert SUS computation table here after Week 7 data collection.]", italic=True, align="center", fi=0, ls=480),
        _body("[Insert qualitative feedback summary here — expected themes: ease of AI chat, utility of FHS score, appreciation for Lite Mode toggle, suggestions for future features.]"),
    ]
    return els

def chapter_four():
    return [
        pb(), _hdr("CHAPTER IV"), _hdr("CONCLUSIONS AND RECOMMENDATIONS"),
        _body("This chapter presents the findings of the study and provides recommendations based on the results and insights gained throughout the research."),
        _subhdr("Conclusions"),
        _body("For the first objective — assessment of financial management practices: The survey and interview data confirmed the presence of financial management challenges identified in the literature: the manual effort burden of traditional expense tracking, irregular budgeting behavior, and the absence of visible consequences for ignoring financial warnings. These findings validated the design rationale for SmartSpend's core features — multi-modal AI input, the Financial Health Score, and the Warning Decay mechanism."),
        _body("For the second objective — system development and LLM benchmarking: SmartSpend v2.9.7 was successfully developed as a fully functional Android application. The comparative benchmarking of 15 LLM API providers confirmed Gemini 3.1 Flash-Lite as the optimal primary model — highest free-tier quota, best Filipino-English performance, and native function calling at zero cost."),
        _body("For the third objective — usability evaluation: [Insert conclusion based on actual SUS score after Week 7 data collection. Target outcome: 'SmartSpend achieved a SUS score of [X] — [Adjective] per Bangor et al. (2009) — meeting/exceeding the ≥80 target, indicating acceptable usability for the target population.']"),
        _body("Overall, SmartSpend demonstrates that a free, offline-capable, Filipino-first AI financial management system can be built entirely on free-tier services — representing a meaningful contribution to financial technology research in the Philippine context by combining agentic AI architecture, behavioral finance principles, gamification mechanics, and culturally localized design in a single mobile application."),
        _subhdr("Recommendations"),
        _body("Based on the findings, development experience, and usability evaluation, the following recommendations are proposed:"),
        _body0("1.   Paluwagan tracker — A rotating savings group tracker should be the highest-priority post-capstone feature, representing a uniquely Filipino informal savings behavior not supported by any other app reviewed. The existing debt and recurring transaction infrastructure provides a suitable architectural base."),
        _body0("2.   15th and 30th payday cycle awareness — Implement payday-cycle-aware budgeting resets aligned with the Philippine standard of semi-monthly salary payments, further strengthening SmartSpend's Filipino-first positioning."),
        _body0("3.   Backend API proxy — Move LLM API key management to a server-side proxy (e.g., Firebase Cloud Function) to eliminate device-side key exposure entirely beyond Firebase Remote Config protection."),
        _body0("4.   Play Store submission — After implementing the backend proxy and a privacy policy, submit SmartSpend to the Google Play Store for wider distribution to Filipino users."),
        _body0("5.   SQLite encryption — Implement SQLCipher-based encryption for the local database in a future schema migration to enhance data security for sensitive financial records."),
        _body0("6.   Couple and family wallet sharing — Allow multiple users to contribute to and view a shared household wallet, directly addressing parents managing household finances with a spouse."),
        _body0("7.   OFW remittance tracking — Add inbound international remittance tracking as a distinct income category, addressing the significant overseas Filipino worker demographic."),
        _body("For future research, longitudinal studies measuring the actual impact of SmartSpend on financial behavior over 3–6 months — savings rate, budget adherence, FHS trend — would provide stronger empirical evidence for the effectiveness of the behavioral intervention mechanisms designed in this study."),
    ]

def references_block():
    refs = [
        "Arcila, A., et al. (2026). FrontierFinance: A challenging benchmark for measuring frontier intelligence of finance agents. arXiv:2608.11683. https://arxiv.org/abs/2608.11683",
        "Ariely, D. (2008). Predictably irrational: The hidden forces that shape our decisions. HarperCollins.",
        "Bangko Sentral ng Pilipinas. (2021). 2021 Financial Inclusion Survey. BSP. https://www.bsp.gov.ph/Inclusive-Finance/Financial-Inclusion-Surveys/2021-FIS-Report.pdf",
        "Bangko Sentral ng Pilipinas. (2025). Consumer Finance and Inclusion Survey (CFIS) 2025. BSP.",
        "Bangor, A., Kortum, P., & Miller, J. (2009). Determining what individual SUS scores mean: Adding an adjective rating scale. Journal of Usability Studies, 4(3), 114-123.",
        "Bitrián, P., Buil, I., & Catalán, S. (2021). Making finance fun: The gamification of personal financial management apps. International Journal of Bank Marketing, 39(7), 1310-1332. https://doi.org/10.1108/IJBM-09-2020-0491",
        "Bloomberg. (2026). How the Philippines' first fintech unicorn is minting financial inclusion. https://sponsored.bloomberg.com/article/mynt/how-the-philippines-first-fintech-unicorn-is-minting-financial-inclusion",
        "Bociek, J. (2023). mobile_scanner: A universal barcode and QR code scanner for Flutter. https://pub.dev/packages/mobile_scanner",
        "Brooke, J. (1996). SUS: A quick and dirty usability scale. In P. W. Jordan et al. (Eds.), Usability evaluation in industry (pp. 189-194). Taylor & Francis.",
        "Cambridge Judge Business School. (2025). From automation to autonomy: The agentic AI era of financial services. https://www.jbs.cam.ac.uk/2025/from-automation-to-autonomy-the-agentic-ai-era-of-financial-services/",
        "Consumer Financial Protection Bureau. (2017). Financial well-being scale: Scale development technical report. CFPB. https://files.consumerfinance.gov/f/documents/201705_cfpb_financial-well-being-scale-technical-report.pdf",
        "Creswell, J. W., & Plano Clark, V. L. (2011). Designing and conducting mixed methods research. Sage Publications.",
        "Davenport, T. H., & Mittal, N. (2022). All-in on AI: How smart companies win big with artificial intelligence. Harvard Business Review Press.",
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
        "Deloitte. (2026). Agentic AI boosts wealth management. https://www.deloitte.com/us/en/insights/industry/financial-services/financial-services-industry-predictions/2026/agentic-ai-wealth-management-productivity.html",
        "Duhigg, C. (2012). The power of habit: Why we do what we do in life and business. Random House.",
        "Dwivedi, Y. K., et al. (2021). Artificial intelligence (AI): Multidisciplinary perspectives on emerging challenges, opportunities, and agenda for research, practice and policy. International Journal of Information Management, 57, 101994. https://doi.org/10.1016/j.ijinfomgt.2019.08.002",
        "Ernst & Young. (2026a). Nearly half of global consumers now use AI to guide savings and investment decisions. EY. https://www.ey.com/en_gl/newsroom/2026/04/nearly-half-of-global-consumers-now-use-ai-to-guide-savings-and-investment-decisions",
        "Ernst & Young. (2026b). EY survey: Autonomous AI is no longer theoretical. EY. https://www.ey.com/en_nl/newsroom/2026/03/ey-survey-autonomous-ai-is-no-longer-theoretical-as-adoption-grows-despite-ongoing-trust-concerns",
        "Financial Health Network. (2021). FinHealth Score Toolkit. https://finhealthnetwork.org/tools/financial-health-score/",
        "Financial Health Network. (2026). From insight to impact: The next phase of financial health measurement. https://finhealthnetwork.org/research/from-insight-to-impact-the-next-phase-of-financial-health-measurement/",
        "Flores, C. A. R. (2025). Financial freedom of Filipinos in personal finance management. Pantao: The International Journal of the Humanities and Social Sciences, 4(1). https://pantaojournal.com/2025/01/27/v4-i1-7/",
        "Flutter 4 Fun. (2022). fl_chart: A highly customizable Flutter chart library. https://pub.dev/packages/fl_chart",
        "GCash / Mynt. (2026). GCash launches country's first AI financial coach embedded in e-wallet [Press release]. PR Newswire. https://www.prnewswire.com/apac/news-releases/ph-fintech-gcash-launches-countrys-first-ai-financial-coach-embedded-in-e-wallet-to-strengthen-financial-literacy-302718569.html",
        "Google. (2024a). Firebase Authentication documentation. https://firebase.google.com/docs/auth",
        "Google. (2024b). Firebase Crashlytics documentation. https://firebase.google.com/docs/crashlytics",
        "Google. (2024c). Firebase Firestore documentation. https://firebase.google.com/docs/firestore",
        "Google. (2024e). Google ML Kit documentation. https://developers.google.com/ml-kit",
        "Google. (2024f). Flutter documentation. https://flutter.dev/docs",
        "Hean, O., Saha, U., & Saha, B. (2025). Can AI help with your personal finances? Applied Economics. https://doi.org/10.1080/00036846.2025.2450384",
        "IBM. (2025). Agentic AI in financial services: Navigating innovation. https://www.ibm.com/think/insights/agentic-ai-financial-services-ethical-adoption",
        "Inquiro. (2024). Financial literacy in the Philippines: Key statistics. https://inquiro.ph/financial-literacy-in-the-philippines-2024-key-statistics/",
        "Juniper Research. (2026). Gamification in banking: How game mechanics drive financial behavior change. [Research report].",
        "Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. Econometrica, 47(2), 263-292.",
        "Li, Z., et al. (2024). A survey of large language models for financial applications. arXiv:2406.11903. https://arxiv.org/abs/2406.11903",
        "Liu, X., et al. (2023). FinGPT: Open-source financial large language models. arXiv:2306.06031. https://arxiv.org/abs/2306.06031",
        "Liu, Z., et al. (2025). Fin-R1: A large language model for financial reasoning through reinforcement learning. arXiv:2503.16252. https://arxiv.org/abs/2503.16252",
        "Meyll, T., et al. (2025). Spendception: The psychological impact of digital payments on consumer purchase behavior and impulse buying. Behavioral Sciences, 15(3), 387. https://doi.org/10.3390/bs15030387",
        "Nielsen, J. (2006). Progressive disclosure. Nielsen Norman Group. https://www.nngroup.com/articles/progressive-disclosure/",
        "NielsenIQ. (2026). The new financial reality: How Filipino consumers are spending, saving, and banking in 2026. https://nielseniq.com/global/en/insights/report/2026/the-new-financial-reality-how-filipino-consumers-are-spending-saving-and-banking-in-2026/",
        "Philippine Statistics Authority. (2021). Family Income and Expenditure Survey (FIES) 2021. PSA. https://www.psa.gov.ph",
        "Philippine Statistics Authority. (2025). Philippine Digital Economy Satellite Account (PDESA) 2025. PSA. https://psa.gov.ph",
        "Plaid. (2026). State of intelligent finance report — Spring 2026. https://plaid.com/blog/state-of-intelligent-finance-report-spring-2026/",
        "Ramsey, D. (2003). Financial peace revisited. Viking.",
        "Roux, A. (2019). sqflite: SQLite plugin for Flutter. https://pub.dev/packages/sqflite",
        "Sloane, L. (2022). speech_to_text: A Flutter plugin for on-device speech recognition. https://pub.dev/packages/speech_to_text",
        "Social Weather Stations. (2026). SWS financial inclusion survey: Philippines financial inclusion rises to 58%. Cited in CoinGeek. https://coingeek.com/10-point-surge-pushes-philippines-financial-inclusion-to-58/",
        "Springer. (2026). Digital nudges and financial inclusion: A study on behavioral interventions. Lecture Notes in Networks and Systems. https://link.springer.com/content/pdf/10.1007/978-3-032-00343-0_14.pdf",
        "Statista. (2024). Number of mobile app downloads worldwide. https://www.statista.com/statistics/271644/worldwide-free-and-paid-mobile-app-store-downloads/",
        "Stefanov, T., Stefanova, M., & Varbanova, S. (2024). Personal finance management application. TEM Journal, 13(3), 2066-2075. https://doi.org/10.18421/TEM133-34",
        "Strivecloud. (2026). Fintech app gamification: Data shows 22% boost in saving habits. https://strivecloud.io/blog/mobile-app-gamification-fintech",
        "Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257-285.",
        "Tarsi – Budget Tracker. (2026). Tarsi – Budget Tracker [Mobile application]. App Store. https://apps.apple.com/ph/app/tarsi-budget-tracker/",
        "The Flutter Authors. (2013). local_auth: Flutter plugin for biometric authentication. https://pub.dev/packages/local_auth",
        "Thaler, R. H., & Sunstein, C. R. (2008). Nudge: Improving decisions about health, wealth, and happiness. Yale University Press.",
        "UNSGSA. (2021). Measuring financial health: A framework for practitioners. https://www.unsgsa.org",
        "Wajid, F., et al. (2025). Gamification: Revolutionizing financial planning systems. World Journal of Advanced Engineering Technology and Sciences. https://www.wjaets.com/sites/default/files/fulltext_pdf/WJAETS-2025-0158.pdf",
        "Warren, E., & Tyagi, A. W. (2005). All your worth: The ultimate lifetime money plan. Free Press.",
        "World Bank. (2022). Global Findex Database 2021. https://www.worldbank.org/en/publication/globalfindex",
        "World Economic Forum. (2024). How agentic AI will transform financial services. https://www.weforum.org/stories/2024/12/agentic-ai-financial-services-autonomy-efficiency-and-inclusion/",
        "Yang, H., et al. (2023). FinGPT: Democratizing internet-scale data for financial large language models. arXiv:2307.10485. https://arxiv.org/abs/2307.10485",
        "Yomio. (2026). YNAB alternatives: Which budget app actually works in 2026? https://yomio.app/en/blog/ynab-alternatives",
        "ZXing Project. (2023). ZXing barcode scanning library. https://github.com/zxing/zxing",
    ]
    els = [pb(), _hdr("REFERENCES"), _blank()]
    for r in refs:
        els.append(_body0(r))
    return els

# ══════════════════════════════════════════════════════════════════════════════
# BUILD THE OUTPUT IN CORRECT ORDER
# ══════════════════════════════════════════════════════════════════════════════
print("Building SMARTSPEND_FINAL_V4.docx in correct section order...")

# 1. TITLE PAGE (source 0–28)
print("  [1] Title page")
append_many(src_els(0, 28))

# 2. APPROVAL SHEET (new)
print("  [2] Approval Sheet")
append_many(approval_sheet())

# 3. ABSTRACT (new)
print("  [3] Abstract")
append_many(abstract_section())

# 4. ACKNOWLEDGEMENT (source 30–51)
print("  [4] Acknowledgement")
append_many(src_els(30, 51))

# 5. DEDICATION (source 52–60)
print("  [5] Dedication")
append_many(src_els(52, 60))

# 6. TABLE OF CONTENTS (fresh rebuild — no stray tables)
print("  [6] Table of Contents")
append_many(toc_section())

# 7. LIST OF FIGURES (new)
print("  [7] List of Figures")
append_many(list_of_figures())

# 8. LIST OF TABLES (new)
print("  [8] List of Tables")
append_many(list_of_tables())

# 9. CHAPTER I (source 99–178, includes images and tables)
print("  [9] Chapter I")
append_many(src_els(99, 178))

# 10. CHAPTER II (source 179–244)
print("  [10] Chapter II")
append_many(src_els(179, 244))

# 11. CHAPTER III (new)
print("  [11] Chapter III")
append_many(chapter_three())

# 12. CHAPTER IV (new)
print("  [12] Chapter IV")
append_many(chapter_four())

# 13. REFERENCES (new — 63 entries)
print("  [13] References")
append_many(references_block())

# 14. APPENDICES header + Appendix A (new) + Appendix B/C/D from source
print("  [14] Appendices")
append_many([pb(), _hdr("APPENDICES")])
append_many(appendix_a())

# Appendix B — Survey Questionnaire (source 291–374 approx)
# Find the range in source
app_b_start = next((i for i, el in enumerate(src_body)
                    if "APPENDIX B" in get_text(el).upper() and len(get_text(el)) < 20), -1)
app_c_start = next((i for i, el in enumerate(src_body)
                    if "APPENDIX C" in get_text(el).upper() and len(get_text(el)) < 20), -1)
cv_start    = next((i for i, el in enumerate(src_body)
                    if "CURRICULUM VITAE" in get_text(el).upper() and len(get_text(el)) < 25), -1)

# Relabel Appendix B → B (Survey), Appendix C → C (Consent), and add new Appendix D (SUS)
if app_b_start >= 0:
    end_b = app_c_start - 1 if app_c_start > app_b_start else cv_start - 1
    print(f"  [14] Appendix B (Survey): src[{app_b_start}..{end_b}]")
    # Copy Appendix B block, but prepend page break
    append_many([pb()])
    append_many(src_els(app_b_start, end_b))

if app_c_start >= 0:
    end_c = next((i for i in range(app_c_start + 1, len(src_body))
                  if "APPENDIX" in get_text(src_body[i]).upper()
                  and len(get_text(src_body[i])) < 20 and i > app_c_start + 2), cv_start - 1)
    if end_c < 0 or end_c <= app_c_start:
        end_c = cv_start - 1 if cv_start > app_c_start else len(src_body) - 1
    print(f"  [14] Appendix C (Consent): src[{app_c_start}..{end_c}]")
    append_many([pb()])
    append_many(src_els(app_c_start, end_c))

# Appendix D — SUS (was Appendix C in original, now renamed D)
# Find the SUS questionnaire block
sus_start = next((i for i in range(app_c_start + 1 if app_c_start >= 0 else 0, len(src_body))
                  if "System Usability Scale" in get_text(src_body[i]) and len(get_text(src_body[i])) < 40), -1)
if sus_start < 0:
    # Try finding by "APPENDIX C" which in old doc was SUS
    for i, el in enumerate(src_body):
        if "APPENDIX C" in get_text(el).upper() and len(get_text(el)) < 20:
            sus_start = i
            break
if sus_start >= 0:
    sus_end = cv_start - 1 if cv_start > sus_start else len(src_body) - 1
    print(f"  [14] Appendix D (SUS): src[{sus_start}..{sus_end}]")
    append_many([pb()])
    append_many([_hdr("APPENDIX D"), _hdr("SYSTEM USABILITY SCALE (SUS) QUESTIONNAIRE"), _blank()])
    # Skip the original appendix header, copy the SUS content
    sus_content_start = sus_start + 1
    for i in range(sus_content_start, sus_end + 1):
        el = src_body[i]
        t  = get_text(el)
        # Skip old "APPENDIX C" headers — we already added APPENDIX D header
        if "APPENDIX C" in t.upper() and len(t) < 20:
            continue
        if "System Usability Scale (SUS) Questionnaire" in t and len(t) < 45:
            continue
        append_many(src_els(i, i))

# 15. CURRICULUM VITAE (source cv_start to end)
if cv_start >= 0:
    print(f"  [15] Curriculum Vitae: src[{cv_start}..{len(src_body)-1}]")
    append_many([pb()])
    append_many(src_els(cv_start, len(src_body) - 2))  # -2 to skip the final sectPr

print(f"\nSaving to {OUT}...")
out.save(OUT)
size = os.path.getsize(OUT)
print(f"Done! Size: {size/1024:.1f} KB")

# ── Quick verification ─────────────────────────────────────────────────────────
print("\n=== Verification ===")
v = Document(OUT)
vb = list(v.element.body)
sections = [
    "APPROVAL SHEET","CAPSTONE PROJECT ABSTRACT","ACKNOWLEDGEMENT","DEDICATION",
    "TABLE OF CONTENTS","LIST OF FIGURES","LIST OF TABLES",
    "Chapter I","CHAPTER II","CHAPTER III","CHAPTER IV",
    "REFERENCES","APPENDIX A","APPENDIX B","APPENDIX C","APPENDIX D",
    "CURRICULUM VITAE"
]
for s in sections:
    found = any(s.lower() in get_text(el).strip().lower()
                and len(get_text(el).strip()) < 80 for el in vb)
    print(f"  {'OK     ' if found else 'MISSING'}  {s}")

in_refs, ref_count = False, 0
for el in vb:
    t = get_text(el).strip()
    if t == "REFERENCES": in_refs = True; continue
    if in_refs:
        if t.upper().startswith("APPENDIX") or t.upper() == "APPENDICES": break
        if t: ref_count += 1
print(f"\n  References: {ref_count}")
print(f"  Tables: {len(v.tables)}")
for i, tbl in enumerate(v.tables):
    r0 = [c.text.strip()[:18] for c in tbl.rows[0].cells]
    print(f"    T{i+1}: {len(tbl.rows)}r x {len(tbl.columns)}c  hdr={r0}")
print(f"  Images: {len(v.inline_shapes)}")
for i, sh in enumerate(v.inline_shapes):
    try: print(f"    IMG{i+1}: {sh.width.inches:.2f}\" x {sh.height.inches:.2f}\"")
    except: print(f"    IMG{i+1}: unknown size")

words = sum(len(get_text(el).split()) for el in vb)
print(f"\n  Word count: ~{words:,} (~{words//250} pages est.)")
