"""Quick check: verify new content in UPDATED doc has ls=480 (double space) matching templates."""
from docx import Document
from docx.oxml.ns import qn

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def get_ls(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None: return None
    sp = pPr.find(qn("w:spacing"))
    if sp is None: return None
    return sp.get(qn("w:line"))

def get_font_size(para):
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                return int(sz.get(qn("w:val"))) / 2
    return None

doc = Document("SMARTSPEND_UPDATED_MANUSCRIPT.docx")

print("=== LINE SPACING CHECK — new sections (should all be 480 for body) ===")
print(f"{'STYLE':20}  {'LS':6}  {'SZ':5}  TEXT[:60]")
print("-"*90)

in_ch3 = False
shown = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if "CHAPTER III" in t.upper() and len(t) < 20:
        in_ch3 = True
    if in_ch3 and t:
        ls  = get_ls(p) or "inh"
        sz  = get_font_size(p) or "inh"
        ok  = "✅" if ls == "480" else ("⚠️" if ls == "inh" else "❌")
        print(f"{ok} {p.style.name:18}  {ls:6}  {sz!s:5}  {t[:60]}")
        shown += 1
        if shown >= 30:
            break

# Summary: count paragraphs by line spacing in new content
print("\n=== SUMMARY — all paragraphs in doc by line spacing ===")
ls_counts = {}
for p in doc.paragraphs:
    ls = get_ls(p) or "inherited"
    ls_counts[ls] = ls_counts.get(ls, 0) + 1
for ls, cnt in sorted(ls_counts.items(), key=lambda x: -x[1]):
    label = {"240":"single","276":"1.15x","360":"1.5x","480":"double","inherited":"inherited"}.get(ls, ls)
    print(f"  {cnt:5d}  ls={ls} ({label})")

# Compare specific para from template vs updated
print("\n=== TEMPLATE vs UPDATED — first body para comparison ===")
tmpl = Document("templates/TEMPLATE_LORMA_ACCESS_PLUS.docx")
tmpl_body_paras = [p for p in tmpl.paragraphs
                   if p.text.strip() and len(p.text.strip()) > 40
                   and get_ls(p) == "480"]
upd_ch3_paras   = []
in_c = False
for p in doc.paragraphs:
    if "CHAPTER III" in p.text.upper() and len(p.text.strip()) < 20:
        in_c = True; continue
    if in_c and p.text.strip() and len(p.text.strip()) > 40:
        upd_ch3_paras.append(p)
        if len(upd_ch3_paras) >= 3: break

if tmpl_body_paras:
    tp = tmpl_body_paras[0]
    print(f"Template body: ls={get_ls(tp)} sz={get_font_size(tp)} | {tp.text[:50]}")
for p in upd_ch3_paras:
    print(f"Updated Ch3:   ls={get_ls(p)} sz={get_font_size(p)} | {p.text[:50]}")
