from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('SMARTSPEND_UPDATED_MANUSCRIPT.docx')

# 1. Page setup
sec = doc.sections[0]
print("=== PAGE SETUP ===")
print(f"  {sec.page_width.inches:.2f}\" x {sec.page_height.inches:.2f}\"")
print(f"  Margins: top={sec.top_margin.inches:.2f}\" bot={sec.bottom_margin.inches:.2f}\" left={sec.left_margin.inches:.2f}\" right={sec.right_margin.inches:.2f}\"")

# 2. Section presence
print("\n=== SECTION CHECK ===")
needed = [
    "APPROVAL SHEET", "ABSTRACT", "ACKNOWLEDGEMENT", "DEDICATION",
    "TABLE OF CONTENTS", "LIST OF FIGURES", "LIST OF TABLES",
    "Chapter I", "CHAPTER II", "CHAPTER III", "CHAPTER IV",
    "REFERENCES", "APPENDIX A", "APPENDIX B", "APPENDIX C",
    "CURRICULUM VITAE"
]
for s in needed:
    found = any(s.lower() in p.text.lower() and len(p.text.strip()) < 80
                for p in doc.paragraphs)
    mark = "OK     " if found else "MISSING"
    print(f"  {mark}  {s}")

# 3. Tables
print(f"\n=== TABLES ({len(doc.tables)}) ===")
for i, t in enumerate(doc.tables):
    r0 = [c.text.strip()[:18] for c in t.rows[0].cells]
    print(f"  T{i+1}: {len(t.rows)}r x {len(t.columns)}c  hdr={r0}")

# 4. Images
print(f"\n=== IMAGES ({len(doc.inline_shapes)}) ===")
for i, sh in enumerate(doc.inline_shapes):
    try:
        print(f"  IMG{i+1}: {sh.width.inches:.2f}\" x {sh.height.inches:.2f}\"")
    except:
        print(f"  IMG{i+1}: size unknown")

# 5. References count
print("\n=== REFERENCES COUNT ===")
in_refs, count = False, 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "REFERENCES":
        in_refs = True; continue
    if in_refs:
        if t.upper().startswith("APPENDIX") or t.upper().startswith("APPENDICES"):
            break
        if t:
            count += 1
print(f"  {count} reference entries found")

# 6. Font consistency check — sample first 20 non-empty paras
print("\n=== FONT SPOT CHECK (first 20 non-empty paras) ===")
shown = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    fn, fsz = "inh", "inh"
    if p.runs:
        r = p.runs[0]
        fn  = r.font.name or "inh"
        fsz = f"{r.font.size.pt:.0f}pt" if r.font.size else "inh"
    sty = p.style.name
    print(f"  [{sty}] {fn}/{fsz}  \"{t[:55]}\"")
    shown += 1
    if shown >= 20:
        break

# 7. Styles used
print("\n=== STYLES USED ===")
styles = {}
for p in doc.paragraphs:
    sn = p.style.name
    styles[sn] = styles.get(sn, 0) + 1
for s, c in sorted(styles.items(), key=lambda x: -x[1]):
    print(f"  {c:5d}x  {s}")

# 8. Approx page count estimate
wc = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"\n=== WORD COUNT ESTIMATE: ~{wc:,} words ===")
print(f"    (~{wc//250} pages at 250 words/page estimate)")
