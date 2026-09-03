"""
build_compliance_matrix.py
Generates SmartSpend_Compliance_Matrix_PreFinal.docx
Pre-filled with project info, ready to print for Pre-Final Defense.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.dirname(__file__)
OUT    = os.path.join(BASE, '..', 'output', 'SmartSpend_Compliance_Matrix_PreFinal.docx')

MAROON = RGBColor(0x5C, 0x0E, 0x24)
DARK   = RGBColor(0x22, 0x22, 0x22)
GREY   = RGBColor(0x6B, 0x6B, 0x6B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
FONT   = 'Tahoma'

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_borders(table, color='AAAAAA', sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    b = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)
        b.append(el)
    tblPr.append(b)

def _set_tbl_width(table, twips):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for e in tblPr.findall(qn('w:tblW')): tblPr.remove(e)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), str(twips)); tw.set(qn('w:type'), 'dxa')
    tblPr.append(tw)

def _p(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
       before=2, after=2, italic=False):
    para = cell.add_paragraph() if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    # Use first paragraph if empty
    if len(cell.paragraphs) >= 2 and not cell.paragraphs[0].text:
        para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    return para

doc = Document()
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
    setattr(sec, attr, Inches(0.75))

doc.styles['Normal'].font.name = FONT
doc.styles['Normal'].font.size = Pt(9)
doc.styles['Normal'].paragraph_format.space_before = Pt(0)
doc.styles['Normal'].paragraph_format.space_after  = Pt(2)

TW = 10224  # total width in twips

# ── HEADER ────────────────────────────────────────────────────────────────────
hdr = doc.add_table(rows=1, cols=1)
_set_tbl_width(hdr, TW)
hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrC = hdr.cell(0,0)
_set_cell_bg(hdrC, '5C0E24')
hp = hdrC.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hp.paragraph_format.space_before = Pt(6)
hp.paragraph_format.space_after  = Pt(6)
r = hp.add_run('LORMA COLLEGES\nCOLLEGE OF COMPUTER STUDIES AND ENGINEERING\nCity of San Fernando, La Union')
r.font.name = FONT; r.font.size = Pt(12); r.bold = True; r.font.color.rgb = WHITE

doc.add_paragraph()

# ── TITLE ─────────────────────────────────────────────────────────────────────
t = doc.add_paragraph('COMPLIANCE MATRIX')
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(4)
t.paragraph_format.space_after  = Pt(2)
tr = t.runs[0]; tr.font.name = FONT; tr.font.size = Pt(14); tr.bold = True; tr.font.color.rgb = MAROON

doc.add_paragraph()

# ── PROJECT INFO TABLE ────────────────────────────────────────────────────────
info = doc.add_table(rows=4, cols=2)
_set_tbl_width(info, TW)
_set_borders(info, 'BBBBBB', 4)
info.alignment = WD_TABLE_ALIGNMENT.LEFT
label_w, val_w = 2400, 7824
rows_data = [
    ('Title of Project:',
     'SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory\n'
     'Application for Personal Financial Management'),
    ('Proponent(s):',
     'Directo, Brix A.\nRubis, Cyrille John M.\nMadayag, Djaunathan Albert S.'),
    ('Type of Defense:',
     '☑ Pre-Final     ☐ Final'),
    ('Date of Defense:', '___________________________'),
]
for i, (lbl, val) in enumerate(rows_data):
    lc = info.cell(i, 0)
    vc = info.cell(i, 1)
    for c, w in ((lc, label_w),(vc, val_w)):
        tcPr = c._tc.get_or_add_tcPr()
        tw = OxmlElement('w:tcW'); tw.set(qn('w:w'), str(w)); tw.set(qn('w:type'),'dxa'); tcPr.append(tw)
    lc.paragraphs[0].clear()
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(3); lp.paragraph_format.space_after = Pt(3)
    lr = lp.add_run(lbl); lr.font.name=FONT; lr.font.size=Pt(9); lr.bold=True; lr.font.color.rgb=DARK
    vc.paragraphs[0].clear()
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(3); vp.paragraph_format.space_after = Pt(3)
    vr = vp.add_run(val); vr.font.name=FONT; vr.font.size=Pt(9); vr.font.color.rgb=DARK

doc.add_paragraph()

# ── COMPLIANCE TABLE BUILDER ──────────────────────────────────────────────────
def _compliance_section(label, num_rows):
    # Section header row
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6); sp.paragraph_format.space_after = Pt(2)
    sr = sp.add_run(label)
    sr.font.name=FONT; sr.font.size=Pt(11); sr.bold=True; sr.font.color.rgb=MAROON

    cols = 4
    tbl = doc.add_table(rows=num_rows+1, cols=cols)
    _set_tbl_width(tbl, TW)
    _set_borders(tbl, 'AAAAAA', 4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [400, 3800, 3200, 2824]
    headers = ['No.', 'Recommendations BY', 'Action Taken', 'Remarks']

    for ci in range(cols):
        hc = tbl.cell(0, ci)
        tcPr = hc._tc.get_or_add_tcPr()
        tw = OxmlElement('w:tcW'); tw.set(qn('w:w'), str(col_widths[ci])); tw.set(qn('w:type'),'dxa'); tcPr.append(tw)
        _set_cell_bg(hc, 'D9D9D9')
        hc.paragraphs[0].clear()
        p = hc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(headers[ci])
        r.font.name=FONT; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=DARK

    for ri in range(1, num_rows+1):
        for ci in range(cols):
            c = tbl.cell(ri, ci)
            tcPr = c._tc.get_or_add_tcPr()
            tw = OxmlElement('w:tcW'); tw.set(qn('w:w'), str(col_widths[ci])); tw.set(qn('w:type'),'dxa'); tcPr.append(tw)
            c.paragraphs[0].clear()
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(14)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(ri))
                r.font.name=FONT; r.font.size=Pt(9); r.font.color.rgb=GREY

    doc.add_paragraph()

_compliance_section('MACHINE (Application)', 6)
_compliance_section('PAPER (Manuscript)', 10)

# ── NOTHING FOLLOWS ──────────────────────────────────────────────────────────
nf = doc.add_paragraph('— — — NOTHING FOLLOWS — — —')
nf.alignment = WD_ALIGN_PARAGRAPH.CENTER
nf.paragraph_format.space_before = Pt(4)
nfr = nf.runs[0]; nfr.font.name=FONT; nfr.font.size=Pt(9); nfr.font.color.rgb=GREY; nfr.italic=True

doc.add_paragraph()

# ── CERTIFICATION + SIGNATURES ────────────────────────────────────────────────
cert_p = doc.add_paragraph(
    'This is to certify that the aforementioned recommendations were made and agreed upon '
    'during the defense by the researcher(s)/proponent(s) and the Oral Examination Committee '
    '(OrEC) composed of:')
cert_p.paragraph_format.space_before = Pt(6)
cert_p.paragraph_format.space_after  = Pt(10)
cr = cert_p.runs[0]; cr.font.name=FONT; cr.font.size=Pt(9); cr.font.color.rgb=DARK

# Signature table
sig = doc.add_table(rows=3, cols=3)
_set_tbl_width(sig, TW)
sig.alignment = WD_TABLE_ALIGNMENT.LEFT
# Remove borders
tbl2 = sig._tbl
tblPr2 = tbl2.tblPr
b2 = OxmlElement('w:tblBorders')
for side in ('top','left','bottom','right','insideH','insideV'):
    el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
    el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto'); b2.append(el)
tblPr2.append(b2)

col_w = [3408, 400, 6416]
signatories = [
    ('SHEKIRO R. RAPOSAS, MIS', 'Teacher In-charge'),
    ('JOHNNY F. VERZOLA, MTS', 'Adviser'),  # updated to actual adviser
]
panelists = [
    ('ELLEN F. MANGAOANG, MIT', 'Chairperson'),
    ('JOPHER F. REYES, MIT', 'Member'),
    ('GELO RYANN M. CARBONELL', 'Member'),
]

# Left column: teacher-in-charge + adviser stacked
# Right column: panel members
def _sig_block(para, name, role, size=9):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(0)
    r1 = para.add_run(name)
    r1.font.name=FONT; r1.font.size=Pt(size); r1.bold=True; r1.font.color.rgb=DARK
    return para

for row_i in range(3):
    for ci in range(3):
        c = sig.cell(row_i, ci)
        tcPr = c._tc.get_or_add_tcPr()
        tw = OxmlElement('w:tcW'); tw.set(qn('w:w'), str(col_w[ci])); tw.set(qn('w:type'),'dxa'); tcPr.append(tw)
        c.paragraphs[0].clear()

# Row 0: labels
left_top = sig.cell(0,0); right_top = sig.cell(0,2)
for name, role in signatories:
    p = left_top.add_paragraph() if left_top.paragraphs[0].text else left_top.paragraphs[0]
    _sig_block(p, name, role)
    rp = left_top.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(role)
    rr.font.name=FONT; rr.font.size=Pt(8); rr.font.color.rgb=GREY
    left_top.add_paragraph()  # spacing

for name, role in panelists:
    p = right_top.add_paragraph() if right_top.paragraphs[0].text else right_top.paragraphs[0]
    _sig_block(p, name, role)
    rp = right_top.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(role)
    rr.font.name=FONT; rr.font.size=Pt(8); rr.font.color.rgb=GREY
    right_top.add_paragraph()

# Signature line row
for ci, label in ((0,'Signature over Printed Name — Researcher(s)'),(2,'Oral Examination Committee')):
    c = sig.cell(2, ci)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run('_' * 40)
    r.font.name=FONT; r.font.size=Pt(9)
    p2 = c.add_paragraph(label)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]; r2.font.name=FONT; r2.font.size=Pt(8); r2.font.color.rgb=GREY

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f'Saved: {OUT}')
