"""
build_sus_sheet.py  —  SUS Computation Sheet for SmartSpend (30 respondents)
Djaunathan fills in the 10 raw scores per respondent; adjusted values and
SUS scores are computed manually using the formula printed at the top.
Run:  python build_sus_sheet.py
Out:  docs/manuscript/output/SUS_Computation_Sheet.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT  = Path(__file__).parent / '..' / 'output' / 'SUS_Computation_Sheet.docx'
FONT = 'Tahoma'
MAROON = RGBColor(0x5C, 0x0E, 0x24)
DARK   = RGBColor(0x22, 0x22, 0x22)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLUE   = RGBColor(0x2A, 0x4A, 0x7F)
GREEN  = RGBColor(0x1A, 0x6B, 0x3A)

doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(13.0)
sec.page_height   = Inches(8.5)
sec.top_margin    = Inches(0.45)
sec.bottom_margin = Inches(0.45)
sec.left_margin   = Inches(0.45)
sec.right_margin  = Inches(0.45)
doc.styles['Normal'].font.name = FONT
doc.styles['Normal'].font.size = Pt(9)


def shade(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)


def border(cell, color='999999', sz=4):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        x = OxmlElement(f'w:{side}')
        x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), str(sz))
        x.set(qn('w:space'), '0'); x.set(qn('w:color'), color)
        b.append(x)
    tcPr.append(b)


def cp(cell, text, bold=False, size=8, align='center', fg=None, italic=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.font.name = FONT; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if fg: r.font.color.rgb = fg


def heading(text, size=12, color=None, center=True, before=8, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = True
    r.font.color.rgb = color or MAROON


def body(text, size=9, italic=False, before=1, after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.italic = italic


# ── TITLE ──────────────────────────────────────────────────────────────────────
heading('SmartSpend — System Usability Scale (SUS) Computation Sheet', 13)
heading('Capstone Project Evaluation  |  Lorma Colleges CCSE BSIT  |  AY 2026–2027, 1st Semester', 9,
        color=DARK, before=0, after=2)
heading('Project Manager & QA Lead: Djaunathan Albert S. Madayag  |  '
        'Study Population: 30 Respondents (20 Parents 35–55, 10 Young Professionals 21–35)', 8.5,
        color=DARK, before=0, after=2)
heading('Formula:  Odd items (Q1,Q3,Q5,Q7,Q9): Adjusted = Raw − 1  |  '
        'Even items (Q2,Q4,Q6,Q8,Q10): Adjusted = 5 − Raw  |  '
        'SUS Score = Sum of 10 Adjusted Values × 2.5  |  Target: ≥ 80', 8.5,
        color=RGBColor(0x44, 0x00, 0x11), before=0, after=4)

# ── SUS ITEMS REFERENCE ────────────────────────────────────────────────────────
heading('SUS Questionnaire Items (respondent rates each 1–5: 1=Strongly Disagree, 5=Strongly Agree)',
        8.5, color=BLUE, center=False, before=2, after=2)

sus_items = [
    ('Q1',  'I think that I would like to use this system frequently.',           'ODD  → Adj = Raw − 1'),
    ('Q2',  'I found the system unnecessarily complex.',                           'EVEN → Adj = 5 − Raw'),
    ('Q3',  'I thought the system was easy to use.',                              'ODD  → Adj = Raw − 1'),
    ('Q4',  'I think that I would need support from a technical person.',         'EVEN → Adj = 5 − Raw'),
    ('Q5',  'I found the various functions in this system well integrated.',      'ODD  → Adj = Raw − 1'),
    ('Q6',  'I thought there was too much inconsistency in this system.',         'EVEN → Adj = 5 − Raw'),
    ('Q7',  'I would imagine that most people would learn this system quickly.',  'ODD  → Adj = Raw − 1'),
    ('Q8',  'I found the system very cumbersome / awkward to use.',               'EVEN → Adj = 5 − Raw'),
    ('Q9',  'I felt very confident using the system.',                            'ODD  → Adj = Raw − 1'),
    ('Q10', 'I needed to learn a lot of things before I could get going.',        'EVEN → Adj = 5 − Raw'),
]
ref_tbl = doc.add_table(rows=2, cols=5)
ref_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for idx, (q, text, formula) in enumerate(sus_items):
    ri, ci = divmod(idx, 5)
    cell = ref_tbl.rows[ri].cells[ci]
    shade(cell, 'F5EEF0' if idx % 2 == 0 else 'FAF6F7')
    border(cell, 'CCAABB', 3)
    cp(cell, f'{q}: {text}\n{formula}', size=7.5, align='left',
       fg=MAROON if 'ODD' in formula else BLUE)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)

# ── MAIN COMPUTATION TABLE ─────────────────────────────────────────────────────
heading('Table 1. SUS Raw Scores, Adjusted Values, and Final Scores — 30 Respondents',
        9, color=DARK, center=False, before=2, after=2)

NCOLS = 24  # No | Group | Q1raw..Q10raw | Q1adj..Q10adj | Sum | Score
NROWS = 2 + 30 + 1 + 1  # 2 headers + 30 data + separator + avg
tbl = doc.add_table(rows=NROWS, cols=NCOLS)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

CW = [Inches(0.28), Inches(0.72)] + [Inches(0.31)]*10 + [Inches(0.31)]*10 + [Inches(0.35), Inches(0.58)]
for ri in range(NROWS):
    for ci, w in enumerate(CW):
        tbl.rows[ri].cells[ci].width = w

# Header row 0
H0_TEXT   = ['No.', 'Group'] + [f'Q{i}' for i in range(1,11)] + [f'Q{i}*' for i in range(1,11)] + ['Σ', 'SUS\nScore']
H0_SHADES = ['5C0E24','5C0E24'] + ['2A4A7F']*10 + ['1A6B3A']*10 + ['5C0E24','5C0E24']
for ci, (txt, shd) in enumerate(zip(H0_TEXT, H0_SHADES)):
    cell = tbl.rows[0].cells[ci]
    shade(cell, shd); border(cell, '555555', 6)
    cp(cell, txt, bold=True, size=8, fg=WHITE)

# Sub-header row 1
H1_TEXT   = ['', 'Raw scores (1–5) →'] + ['Raw']*10 + ['Adj']*10 + ['Sum\n(max 40)', 'Score\n(max 100)']
H1_SHADES = ['DDDDDD','DDDDDD'] + ['DDEEFF']*10 + ['DDFFEE']*10 + ['DDDDDD','DDDDDD']
for ci, (txt, shd) in enumerate(zip(H1_TEXT, H1_SHADES)):
    cell = tbl.rows[1].cells[ci]
    shade(cell, shd); border(cell, '999999', 3)
    cp(cell, txt, bold=False, size=7.5, fg=DARK)

# Data rows
groups = ['Parent (35–55)']*20 + ['Young Prof. (21–35)']*10
for i in range(30):
    ri = i + 2
    row_shade  = 'FFFFFF' if i % 2 == 0 else 'F8F8F8'
    grp_shade  = 'FFF0F3' if i < 20 else 'F0FFF4'
    raw_shade  = 'FFFDE7'   # yellow — fill in
    adj_shade  = 'E8F5E9'   # green — compute
    sum_shade  = 'E3F2FD'   # blue
    scr_shade  = 'EDE7F6'   # purple

    shade(tbl.rows[ri].cells[0], row_shade); border(tbl.rows[ri].cells[0])
    cp(tbl.rows[ri].cells[0], str(i+1), bold=True, size=8.5)

    shade(tbl.rows[ri].cells[1], grp_shade); border(tbl.rows[ri].cells[1])
    cp(tbl.rows[ri].cells[1], groups[i], size=7.5, align='left',
       fg=MAROON if i < 20 else GREEN)

    for ci in range(2, 12):
        shade(tbl.rows[ri].cells[ci], raw_shade); border(tbl.rows[ri].cells[ci], 'CCBBAA')
        cp(tbl.rows[ri].cells[ci], '_____', size=8, fg=RGBColor(0xCC,0xCC,0xCC))

    for ci in range(12, 22):
        shade(tbl.rows[ri].cells[ci], adj_shade); border(tbl.rows[ri].cells[ci], 'AACCAA')
        cp(tbl.rows[ri].cells[ci], '_____', size=8, fg=RGBColor(0xBB,0xCC,0xBB))

    shade(tbl.rows[ri].cells[22], sum_shade); border(tbl.rows[ri].cells[22])
    cp(tbl.rows[ri].cells[22], '____', size=8.5)

    shade(tbl.rows[ri].cells[23], scr_shade); border(tbl.rows[ri].cells[23])
    cp(tbl.rows[ri].cells[23], '______', size=9, bold=True, fg=MAROON)

# Separator row 32
for ci in range(NCOLS):
    shade(tbl.rows[32].cells[ci], 'CCCCCC'); border(tbl.rows[32].cells[ci], '888888', 6)
    cp(tbl.rows[32].cells[ci], '')

# Average row 33
shade(tbl.rows[33].cells[0], '5C0E24'); border(tbl.rows[33].cells[0])
cp(tbl.rows[33].cells[0], '—', bold=True, size=8, fg=WHITE)
shade(tbl.rows[33].cells[1], '5C0E24'); border(tbl.rows[33].cells[1])
cp(tbl.rows[33].cells[1], 'OVERALL AVERAGE (n=30)', bold=True, size=8.5, fg=WHITE, align='left')
for ci in range(2, 23):
    shade(tbl.rows[33].cells[ci], 'EEE8F5'); border(tbl.rows[33].cells[ci])
    cp(tbl.rows[33].cells[ci], '')
shade(tbl.rows[33].cells[23], 'EDE7F6'); border(tbl.rows[33].cells[23])
cp(tbl.rows[33].cells[23], '[  ____  ]', bold=True, size=10, fg=MAROON)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)

# ── SCORE INTERPRETATION ───────────────────────────────────────────────────────
heading('Table 2. SUS Score Interpretation Guide (Bangor et al., 2009; Brooke, 1996)',
        9, color=DARK, center=False, before=2, after=2)

interp_rows = [
    ['SUS Score Range', 'Grade', 'Adjective\n(Bangor et al., 2009)', 'Acceptability\n(Brooke, 1996)', 'SmartSpend Target'],
    ['≥ 90.0',     'A+', 'Best Imaginable',  'Acceptable',     ''],
    ['85.0 – 89.9','A',  'Excellent',        'Acceptable',     ''],
    ['80.0 – 84.9','B',  'Good',             'Acceptable',     '← MINIMUM TARGET'],
    ['70.0 – 79.9','C',  'OK',               'Marginal',       ''],
    ['51.0 – 69.9','D',  'Poor',             'Marginal',       ''],
    ['< 51.0',     'F',  'Awful',            'Not Acceptable', ''],
]
row_shades = ['5C0E24','E8F5E9','E8F5E9','C8E6C9','FFF9C4','FFCCBC','FFCCBC']
iw = [Inches(1.4), Inches(0.5), Inches(1.3), Inches(1.4), Inches(1.8)]
itbl = doc.add_table(rows=7, cols=5)
itbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for ri, (rdata, rshade) in enumerate(zip(interp_rows, row_shades)):
    for ci, txt in enumerate(rdata):
        cell = itbl.rows[ri].cells[ci]
        cell.width = iw[ci]
        shade(cell, rshade); border(cell, 'AAAAAA', 3)
        is_hdr = ri == 0
        is_target = ri == 3 and ci == 4
        cp(cell, txt, bold=is_hdr or is_target, size=9,
           fg=WHITE if is_hdr else (MAROON if is_target else DARK))

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)

# ── INSTRUCTIONS ───────────────────────────────────────────────────────────────
heading('Instructions for Djaunathan (QA Lead):', 9, color=MAROON, center=False, before=4, after=2)
steps = [
    '1.  Print this sheet or open in Word/Excel. After each respondent demo session, collect the filled SUS questionnaire.',
    '2.  Enter their 10 raw scores (1–5) in the yellow cells (Q1–Q10 columns).',
    '3.  Compute adjusted values — ODD items (Q1,Q3,Q5,Q7,Q9): subtract 1. EVEN items (Q2,Q4,Q6,Q8,Q10): subtract from 5.',
    '4.  Enter adjusted values in the green cells (Q1*–Q10* columns).',
    '5.  Sum all 10 adjusted values and enter in Σ column (maximum = 40).',
    '6.  Multiply Σ × 2.5 = SUS Score. Enter in the purple SUS Score column (maximum = 100).',
    '7.  After all 30 respondents are done: average the 30 SUS scores for the Overall Average.',
    '8.  Compare the Overall Average to Table 2. Target is ≥ 80.0 (Grade B, "Good", Acceptable).',
    '9.  Bring completed sheets to Brix so scores can be inserted into Chapter III of the manuscript.',
    '',
    'References:',
    'Brooke, J. (1996). SUS: A quick and dirty usability scale. In Usability evaluation in industry (pp. 189–194). Taylor & Francis.',
    'Bangor, A., Kortum, P., & Miller, J. (2009). Determining what individual SUS scores mean. Journal of Usability Studies, 4(3), 114–123.',
]
for line in steps:
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(line)
    r2.font.name = FONT; r2.font.size = Pt(8.5)
    r2.font.italic = line.startswith('Reference') or line.startswith('Brooke') or line.startswith('Bangor')
    r2.font.bold   = line.startswith('Reference')

doc.save(str(OUT))
print(f'  SUS sheet saved: {OUT.name}  ({OUT.stat().st_size//1024} KB)')
