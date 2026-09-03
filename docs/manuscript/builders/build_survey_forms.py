"""
build_survey_forms.py
Generates two DOCX files with the exact Google Forms setup instructions:
  1. SmartSpend_Survey1_FinancialPractices_Form.docx
  2. SmartSpend_Survey2_SUS_Form.docx

Each document contains:
  - The exact questions to type into Google Forms
  - Question type for each item
  - Options/scale for each item
  - Copy-paste ready text

Run:  python build_survey_forms.py
Out:  docs/manuscript/output/
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / '..' / 'output'
FONT   = 'Tahoma'
MAROON = RGBColor(0x5C, 0x0E, 0x24)
BLUE   = RGBColor(0x2A, 0x4A, 0x7F)
GREEN  = RGBColor(0x1A, 0x6B, 0x3A)
DARK   = RGBColor(0x22, 0x22, 0x22)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(16)
    r.font.bold = True; r.font.color.rgb = MAROON


def h2(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(13)
    r.font.bold = True; r.font.color.rgb = color or BLUE


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(11)
    r.font.bold = True; r.font.color.rgb = DARK


def body(doc, text, italic=False, before=2, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(11)
    r.font.italic = italic


def question_block(doc, number, question_text, q_type, options=None, required=True):
    """Render a single Google Forms question block."""
    # Question number + text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'Q{number}. ')
    r1.font.name = FONT; r1.font.size = Pt(11); r1.font.bold = True
    r1.font.color.rgb = MAROON
    r2 = p.add_run(question_text)
    r2.font.name = FONT; r2.font.size = Pt(11); r2.font.bold = True

    # Type tag
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(1)
    p2.paragraph_format.left_indent  = Inches(0.25)
    r3 = p2.add_run(f'  Type: {q_type}')
    r3.font.name = FONT; r3.font.size = Pt(10); r3.font.italic = True
    r3.font.color.rgb = BLUE

    req_tag = '  [Required ✓]' if required else '  [Optional]'
    r4 = p2.add_run(req_tag)
    r4.font.name = FONT; r4.font.size = Pt(10)
    r4.font.color.rgb = GREEN if required else RGBColor(0x99, 0x99, 0x99)

    # Options
    if options:
        for opt in options:
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after  = Pt(0)
            p3.paragraph_format.left_indent  = Inches(0.5)
            r5 = p3.add_run(f'○  {opt}')
            r5.font.name = FONT; r5.font.size = Pt(10.5)


def tip_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, 'FFF8DC')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f'💡  {text}')
    r.font.name = FONT; r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
# FORM 1 — FINANCIAL MANAGEMENT PRACTICES SURVEY
# ══════════════════════════════════════════════════════════════════════════════

doc1 = new_doc()

h1(doc1, 'SmartSpend — Survey 1 Setup Guide')
h1(doc1, 'Financial Management Practices Questionnaire')
body(doc1, 'Google Forms Setup Instructions for Djaunathan (QA Lead)', italic=True)
body(doc1, 'Copy each question exactly as written below into a new Google Form.', before=0)
body(doc1, 'Go to: forms.google.com → Blank form → Add questions one by one.', before=0)

tip_box(doc1, 'Form Title to use: "SmartSpend Research Survey — Financial Management Practices"\n'
              'Description: "This survey is part of a capstone research study at Lorma Colleges CCSE. '
              'Your responses will remain confidential and will only be used for academic purposes."')

# ── FORM HEADER SETTINGS ──────────────────────────────────────────────────────
h2(doc1, 'Google Forms Settings (before adding questions)')
for s in [
    'Click the ⚙️ Settings gear icon',
    'Under Responses: turn ON "Collect email addresses" → OFF (anonymous)',
    'Under Presentation: turn ON "Show progress bar"',
    'Under Presentation: Confirmation message → type: "Thank you for participating in our SmartSpend research study!"',
    'Click Save',
]:
    p = doc1.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(s); r.font.name = FONT; r.font.size = Pt(11)

# ── SECTION 1 ─────────────────────────────────────────────────────────────────
h2(doc1, 'SECTION 1: Respondent Profile', color=MAROON)
body(doc1, 'Add a Section header in Google Forms: click the = icon → "Section" → title: "Part I — Respondent Profile"')

question_block(doc1, 1, 'What is your age?', 'Short answer (text)', required=True)
question_block(doc1, 2, 'What is your gender?', 'Multiple choice', [
    'Male', 'Female', 'Prefer not to say'
])
question_block(doc1, 3, 'Which best describes your primary role?', 'Multiple choice', [
    'Parent / Household Financial Manager (Ages 35–55)',
    'Young Professional (Ages 21–35)',
])
tip_box(doc1, 'Q3 is your screening question. Make sure all 20 parents select the first option '
              'and all 10 young professionals select the second.')

question_block(doc1, 4, 'What is your current employment status?', 'Multiple choice', [
    'Employed (full-time)', 'Employed (part-time)', 'Self-employed / Business owner',
    'Freelancer', 'Student', 'Unemployed / Not currently working',
])
question_block(doc1, 5, 'What is your monthly income range?', 'Multiple choice', [
    'Below ₱10,000', '₱10,000 – ₱20,000', '₱20,000 – ₱40,000',
    'Above ₱40,000', 'Prefer not to say / Not applicable',
])

# ── SECTION 2 ─────────────────────────────────────────────────────────────────
h2(doc1, 'SECTION 2: Financial Management Practices', color=MAROON)
body(doc1, 'Add a new Section: title "Part II — Financial Management Practices"')

question_block(doc1, 6, 'How do you usually track your expenses?', 'Checkboxes (tick all that apply)', [
    'Written notebook / diary', 'Spreadsheet (Excel, Google Sheets)',
    'Mobile budgeting application', 'Mental tracking only (no written record)',
    'I do not track my expenses', 'Other (please specify)',
])
question_block(doc1, 7, 'How often do you monitor your budget?', 'Multiple choice', [
    'Daily', 'Weekly', 'Monthly', 'Rarely', 'Never',
])
question_block(doc1, 8, 'What is your primary source of income?', 'Multiple choice', [
    'Salary (full-time employment)', 'Part-time job',
    'Freelance / project-based work', 'Business / self-employment',
    'Allowance (from family)', 'Pension / government benefit',
    'Other',
])
question_block(doc1, 9, 'Do you regularly set aside savings from your income?', 'Multiple choice', [
    'Yes, every month', 'Sometimes (not consistently)',
    'Rarely', 'No',
])
question_block(doc1, 10,
    'Approximately what percentage of your monthly income do you currently save?',
    'Multiple choice', [
        'I don\'t save regularly',
        'Less than 10%', '10% – 20%', 'More than 20%', 'I\'m not sure',
    ])

# ── SECTION 3 ─────────────────────────────────────────────────────────────────
h2(doc1, 'SECTION 3: Budgeting Challenges', color=MAROON)
body(doc1, 'Add a new Section: title "Part III — Budgeting Challenges"')

question_block(doc1, 11,
    'What are your common difficulties in managing your expenses? (Check all that apply)',
    'Checkboxes', [
        'Overspending beyond my budget',
        'Forgetting to record my expenses',
        'Too much effort to track manually',
        'Difficulty categorizing my expenses',
        'No clear system or method for budgeting',
        'Irregular income makes budgeting difficult',
        'I don\'t have difficulties',
        'Other (please specify)',
    ])
question_block(doc1, 12,
    'How difficult is it for you to track your daily expenses?',
    'Linear scale (1 = Very Easy, 5 = Very Difficult)', [
        '1 — Very Easy', '2 — Easy', '3 — Moderate',
        '4 — Difficult', '5 — Very Difficult',
    ])
tip_box(doc1, 'In Google Forms: choose "Linear scale" → Min value: 1, label: "Very Easy" → Max value: 5, label: "Very Difficult"')

question_block(doc1, 13,
    'Do you find manual expense tracking inconvenient?',
    'Multiple choice', ['Yes', 'No', 'Sometimes'])
question_block(doc1, 14,
    'Do you currently have existing debts or loans you are paying?',
    'Multiple choice', ['Yes', 'No'])
question_block(doc1, 15,
    'Do you ever experience difficulty meeting your expenses before your next salary or allowance?',
    'Multiple choice', ['Yes', 'No', 'Sometimes'])
question_block(doc1, 16,
    'Have you ever exceeded your monthly budget and then ignored a spending warning or alert?',
    'Multiple choice', ['Yes', 'No', 'Not applicable'])

# ── SECTION 4 ─────────────────────────────────────────────────────────────────
h2(doc1, 'SECTION 4: Feature Needs Assessment', color=MAROON)
body(doc1, 'Add a new Section: title "Part IV — Feature Needs"')

question_block(doc1, 17,
    'Which of the following features would you find helpful in a financial tracking app? (Check all that apply)',
    'Checkboxes', [
        'Automatic receipt scanning (camera / photo)',
        'Voice input for logging expenses',
        'AI-powered financial advice and tips',
        'Automatic categorization of expenses',
        'Budget alerts and spending warnings',
        'Savings goal tracking',
        'Visual charts and analytics',
        'Offline functionality (works without internet)',
        'GCash / Maya / e-wallet integration',
        'SSS / PhilHealth / Pag-IBIG contribution tracking',
    ])
question_block(doc1, 18,
    'Would you use an AI-powered financial assistant application to help manage your finances?',
    'Multiple choice', ['Yes, definitely', 'Probably yes', 'Not sure', 'Probably not', 'No'])
question_block(doc1, 19,
    'What features or improvements would you suggest for a Filipino financial management app?',
    'Paragraph (long text)', required=False)

# ── SHARING INSTRUCTIONS ──────────────────────────────────────────────────────
h2(doc1, 'How to Share the Form', color=GREEN)
for s in [
    'Click Send (paper airplane icon) at the top right',
    'Click the link icon 🔗',
    'Check "Shorten URL"',
    'Copy the short link and share via group chat / QR code',
    'OR: Click the QR code icon to generate a QR code respondents can scan with their phone',
]:
    p = doc1.add_paragraph(style='List Number')
    r = p.add_run(s); r.font.name = FONT; r.font.size = Pt(11)

h2(doc1, 'Downloading Results for Analysis', color=GREEN)
for s in [
    'Go to the Responses tab in your Google Form',
    'Click the Google Sheets icon (green) to open responses in a spreadsheet',
    'Download as CSV: File → Download → Comma Separated Values (.csv)',
    'Give the CSV to Brix to run through the SUS calculator script',
]:
    p = doc1.add_paragraph(style='List Number')
    r = p.add_run(s); r.font.name = FONT; r.font.size = Pt(11)

out1 = OUT / 'SmartSpend_Survey1_Form_Setup.docx'
doc1.save(str(out1))
print(f'  Survey 1 form guide saved: {out1.name}')


# ══════════════════════════════════════════════════════════════════════════════
# FORM 2 — SUS QUESTIONNAIRE
# ══════════════════════════════════════════════════════════════════════════════

doc2 = new_doc()

h1(doc2, 'SmartSpend — Survey 2 Setup Guide')
h1(doc2, 'System Usability Scale (SUS) Questionnaire')
body(doc2, 'Google Forms Setup Instructions — Administer AFTER the SmartSpend demo', italic=True)
body(doc2, 'This is a separate form from Survey 1. Create a new Google Form.', before=0)

tip_box(doc2, 'Form Title: "SmartSpend Usability Evaluation — SUS Questionnaire"\n'
              'Description: "Please rate your experience with the SmartSpend app using the scale below.\n'
              '1 = Strongly Disagree    2 = Disagree    3 = Neutral    4 = Agree    5 = Strongly Agree\n'
              'There are no right or wrong answers. Answer based on your first impression of the app."')

h2(doc2, 'IMPORTANT: Google Forms Settings for SUS')
tip_box(doc2,
    'For ALL 10 questions: use "Linear scale" → Min: 1, Max: 5\n'
    'Min label: "Strongly Disagree"    Max label: "Strongly Agree"\n'
    'Mark all questions as Required.\n'
    'Do NOT show respondents the scoring formula — just collect the raw 1–5 ratings.')

h2(doc2, 'Add a Section header: "System Usability Scale (SUS)"', color=MAROON)
body(doc2, 'Description for the section: "For each statement below, select the number that best '
          'reflects your personal reaction to SmartSpend. Do not think too long — '
          'your immediate response is what matters."')

# 10 SUS items
sus_items = [
    ('I think that I would like to use this system frequently.',
     'ODD item — do NOT tell respondents about odd/even scoring'),
    ('I found the system unnecessarily complex.',
     'EVEN item'),
    ('I thought the system was easy to use.',
     'ODD item'),
    ('I think that I would need the support of a technical person to be able to use this system.',
     'EVEN item'),
    ('I found the various functions in this system were well integrated.',
     'ODD item'),
    ('I thought there was too much inconsistency in this system.',
     'EVEN item'),
    ('I would imagine that most people would learn to use this system very quickly.',
     'ODD item'),
    ('I found the system very cumbersome / awkward to use.',
     'EVEN item'),
    ('I felt very confident using the system.',
     'ODD item'),
    ('I needed to learn a lot of things before I could get going with this system.',
     'EVEN item'),
]

for i, (question, note) in enumerate(sus_items, 1):
    p = doc2.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'Q{i}.  ')
    r1.font.name = FONT; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = MAROON
    r2 = p.add_run(question)
    r2.font.name = FONT; r2.font.size = Pt(11); r2.font.bold = True

    p2 = doc2.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(1)
    r3 = p2.add_run('  Type: Linear scale (1 = Strongly Disagree → 5 = Strongly Agree)  ')
    r3.font.name = FONT; r3.font.size = Pt(10); r3.font.italic = True; r3.font.color.rgb = BLUE
    r4 = p2.add_run(f'[{note}]')
    r4.font.name = FONT; r4.font.size = Pt(9); r4.font.italic = True
    r4.font.color.rgb = RGBColor(0x99, 0x66, 0x00)

# Respondent info section
h2(doc2, 'Optional: Add at the END of the SUS form', color=GREEN)
body(doc2, 'Add a new Section after Q10 titled "Your Profile" with these 2 optional questions:')
question_block(doc2, 11, 'Which best describes you?', 'Multiple choice (optional)', [
    'Parent / Household Financial Manager (Ages 35–55)',
    'Young Professional (Ages 21–35)',
], required=False)
question_block(doc2, 12,
    'Any additional comments about the SmartSpend app?',
    'Paragraph (optional)', required=False)

# Sharing
h2(doc2, 'How to Administer (Step by Step)', color=GREEN)
steps = [
    'Open the SmartSpend app on your phone in Demo Mode (Login screen → Skip / Demo)',
    'Show the app to the respondent for 8–9 minutes following the demo script',
    'Let the respondent try the app briefly (voice logging, AI chat, analytics)',
    'After the demo, open the SUS Google Form link on YOUR phone or the respondent\'s phone',
    'Let the respondent fill in all 10 items independently (3–4 minutes)',
    'Submit the form',
    'Move to the next respondent',
]
for i, s in enumerate(steps, 1):
    p = doc2.add_paragraph(style='List Number')
    r = p.add_run(s); r.font.name = FONT; r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(2)

tip_box(doc2,
    'Target: 30 respondents total — 20 parents (35–55) + 10 young professionals (21–35).\n'
    'You can do both Survey 1 and Survey 2 in the same session.\n'
    'Survey 1 first (before demo) → Demo → Survey 2 (after demo).\n'
    'Each full session takes about 20–25 minutes per respondent.')

h2(doc2, 'Downloading Results', color=GREEN)
body(doc2, 'After all 30 respondents are done:')
for s in [
    'Responses tab → click Google Sheets icon to open in spreadsheet',
    'File → Download → CSV',
    'Give the CSV file to Brix',
    'Run: python compute_sus_scores.py  (in docs/manuscript/builders/)',
    'The script will auto-compute all 30 SUS scores and output the final table',
]:
    p = doc2.add_paragraph(style='List Number')
    r = p.add_run(s); r.font.name = FONT; r.font.size = Pt(11)

out2 = OUT / 'SmartSpend_Survey2_SUS_Form_Setup.docx'
doc2.save(str(out2))
print(f'  Survey 2 SUS form guide saved: {out2.name}')
