"""
compute_sus_scores.py
Reads the Google Forms CSV export for the SUS questionnaire and computes:
  - Adjusted score per item per respondent
  - SUS score per respondent
  - Group averages (Parents vs Young Professionals)
  - Overall average, grade, and adjective rating
  - Outputs a DOCX table ready to paste into Chapter III

Usage:
  1. Download the SUS Google Form responses as CSV
     (Responses tab → Google Sheets icon → File → Download → CSV)
  2. Place the CSV in the same folder as this script
  3. Run: python compute_sus_scores.py your_sus_responses.csv
     OR:  python compute_sus_scores.py   (will look for any CSV in the folder)

Output: docs/manuscript/output/SUS_Results_Chapter3.docx

SUS Formula (Brooke, 1996):
  Odd items  (Q1,Q3,Q5,Q7,Q9):  adjusted = raw - 1
  Even items (Q2,Q4,Q6,Q8,Q10): adjusted = 5 - raw
  SUS Score = sum of 10 adjusted values × 2.5   (range 0–100)
"""

import sys
import csv
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent / '..' / 'output'
OUT_DIR.mkdir(exist_ok=True)

FONT   = 'Tahoma'
MAROON = RGBColor(0x5C, 0x0E, 0x24)
BLUE   = RGBColor(0x2A, 0x4A, 0x7F)
GREEN  = RGBColor(0x1A, 0x6B, 0x3A)
DARK   = RGBColor(0x22, 0x22, 0x22)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# SUS score interpretation (Bangor et al., 2009)
def interpret_sus(score):
    if score >= 90.0: return 'A+', 'Best Imaginable', 'Acceptable'
    if score >= 85.0: return 'A',  'Excellent',       'Acceptable'
    if score >= 80.0: return 'B',  'Good',            'Acceptable'
    if score >= 70.0: return 'C',  'OK',              'Marginal'
    if score >= 51.0: return 'D',  'Poor',            'Marginal'
    return 'F', 'Awful', 'Not Acceptable'


def compute_sus(raw_scores):
    """
    raw_scores: list of 10 integers (1–5), Q1 through Q10
    Returns SUS score (0–100)
    """
    adjusted = []
    for i, raw in enumerate(raw_scores):
        if (i + 1) % 2 == 1:   # odd items: Q1,Q3,Q5,Q7,Q9
            adjusted.append(int(raw) - 1)
        else:                   # even items: Q2,Q4,Q6,Q8,Q10
            adjusted.append(5 - int(raw))
    return sum(adjusted) * 2.5


def find_csv():
    """Find a CSV file in the script directory."""
    script_dir = Path(__file__).parent
    csvs = list(script_dir.glob('*.csv'))
    if not csvs:
        # Also check output directory
        csvs = list(OUT_DIR.glob('*.csv'))
    return csvs[0] if csvs else None


def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)


def cp(cell, text, bold=False, size=9, align='center', color=None, italic=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.font.name = FONT; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color


def build_results_docx(results, csv_path):
    """
    results: list of dicts with keys:
      respondent_no, group, raw (list of 10), adjusted (list of 10),
      sum_adj, sus_score
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(13.0)   # landscape
    sec.page_height   = Inches(8.5)
    sec.top_margin    = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin   = Inches(0.5)
    sec.right_margin  = Inches(0.5)
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(9)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SmartSpend — SUS Evaluation Results')
    r.font.name = FONT; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = MAROON

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Source: {Path(csv_path).name}  |  n = {len(results)} respondents  |  '
                     'Formula: Brooke (1996)  |  Interpretation: Bangor et al. (2009)')
    r2.font.name = FONT; r2.font.size = Pt(9); r2.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── MAIN TABLE ─────────────────────────────────────────────────────────────
    ncols = 25  # No | Group | Q1..Q10 raw | Q1..Q10 adj | Sum | Score
    tbl = doc.add_table(rows=len(results) + 2, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    CW = ([Inches(0.28), Inches(0.85)] +
          [Inches(0.30)] * 10 +
          [Inches(0.30)] * 10 +
          [Inches(0.35), Inches(0.58)])

    for ri in range(len(results) + 2):
        for ci, w in enumerate(CW):
            tbl.rows[ri].cells[ci].width = w

    # Header row 0
    H0 = (['No.', 'Group'] +
          [f'Q{i}' for i in range(1, 11)] +
          [f'Q{i}*' for i in range(1, 11)] +
          ['Σ', 'SUS'])
    H0_SHADES = ['5C0E24', '5C0E24'] + ['2A4A7F'] * 10 + ['1A6B3A'] * 10 + ['5C0E24', '5C0E24']
    for ci, (txt, shd) in enumerate(zip(H0, H0_SHADES)):
        cell = tbl.rows[0].cells[ci]
        shade_cell(cell, shd)
        cp(cell, txt, bold=True, size=8, color=WHITE)

    # Sub-header row 1
    H1 = (['', 'Raw (1–5) →'] +
          ['Raw'] * 10 +
          ['Adj'] * 10 +
          ['max 40', 'max 100'])
    H1_SHADES = ['DDDDDD', 'DDDDDD'] + ['DDEEFF'] * 10 + ['DDFFEE'] * 10 + ['DDDDDD', 'DDDDDD']
    for ci, (txt, shd) in enumerate(zip(H1, H1_SHADES)):
        cell = tbl.rows[1].cells[ci]
        shade_cell(cell, shd)
        cp(cell, txt, bold=False, size=7.5)

    # Data rows
    for ri, res in enumerate(results):
        row_shade = 'FFFFFF' if ri % 2 == 0 else 'F8F8F8'
        grp_shade = 'FFF0F3' if res['group'] == 'parent' else 'F0FFF4'
        tr = tbl.rows[ri + 2]

        shade_cell(tr.cells[0], row_shade)
        cp(tr.cells[0], str(ri + 1), bold=True, size=8.5)

        shade_cell(tr.cells[1], grp_shade)
        grp_text = 'Parent (35–55)' if res['group'] == 'parent' else 'Young Prof. (21–35)'
        cp(tr.cells[1], grp_text, size=7.5, align='left',
           color=MAROON if res['group'] == 'parent' else GREEN)

        for ci in range(10):
            shade_cell(tr.cells[ci + 2], 'FFFDE7')
            cp(tr.cells[ci + 2], str(res['raw'][ci]), size=8.5)

        for ci in range(10):
            shade_cell(tr.cells[ci + 12], 'E8F5E9')
            cp(tr.cells[ci + 12], str(res['adjusted'][ci]), size=8.5)

        shade_cell(tr.cells[22], 'E3F2FD')
        cp(tr.cells[22], str(res['sum_adj']), size=8.5)

        score = res['sus_score']
        scr_shade = ('C8E6C9' if score >= 80 else
                     'FFF9C4' if score >= 70 else
                     'FFCCBC')
        shade_cell(tr.cells[23], scr_shade)
        cp(tr.cells[23], f'{score:.1f}', bold=True, size=9,
           color=GREEN if score >= 80 else (RGBColor(0x80, 0x60, 0x00) if score >= 70 else RGBColor(0xCC, 0x00, 0x00)))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── SUMMARY TABLE ──────────────────────────────────────────────────────────
    parent_scores = [r['sus_score'] for r in results if r['group'] == 'parent']
    youngprof_scores = [r['sus_score'] for r in results if r['group'] == 'youngprof']
    all_scores = [r['sus_score'] for r in results]

    parent_avg    = sum(parent_scores)    / len(parent_scores)    if parent_scores    else 0
    youngprof_avg = sum(youngprof_scores) / len(youngprof_scores) if youngprof_scores else 0
    overall_avg   = sum(all_scores)       / len(all_scores)       if all_scores       else 0

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Table 3.X. SUS Score Summary — SmartSpend Usability Evaluation')
    r3.font.name = FONT; r3.font.size = Pt(10); r3.font.bold = True

    grade_p, adj_p, accept_p = interpret_sus(parent_avg)
    grade_y, adj_y, accept_y = interpret_sus(youngprof_avg)
    grade_o, adj_o, accept_o = interpret_sus(overall_avg)

    sum_tbl = doc.add_table(rows=5, cols=6)
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sum_tbl.style = 'Table Grid'

    sum_data = [
        ['Group', 'n', 'Average SUS Score', 'Grade', 'Adjective', 'Acceptability'],
        ['Parents (Ages 35–55)', str(len(parent_scores)),
         f'{parent_avg:.2f}', grade_p, adj_p, accept_p],
        ['Young Professionals (Ages 21–35)', str(len(youngprof_scores)),
         f'{youngprof_avg:.2f}', grade_y, adj_y, accept_y],
        ['Overall', str(len(all_scores)),
         f'{overall_avg:.2f}', grade_o, adj_o, accept_o],
        ['Target', '30', '≥ 80.00', 'B', 'Good', 'Acceptable'],
    ]
    sum_shades = ['5C0E24', 'FFF0F3', 'F0FFF4', 'EDE7F6', 'F5F5F5']
    for ri, (row_data, rshade) in enumerate(zip(sum_data, sum_shades)):
        for ci, val in enumerate(row_data):
            cell = sum_tbl.rows[ri].cells[ci]
            shade_cell(cell, rshade)
            is_result = (ri == 3)
            fc = WHITE if ri == 0 else (MAROON if is_result else DARK)
            cp(cell, val, bold=(ri == 0 or is_result), size=10,
               color=fc, align='center' if ci > 1 else 'left')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── INTERPRETATION ─────────────────────────────────────────────────────────
    p4 = doc.add_paragraph()
    r4 = p4.add_run('Interpretation and Discussion Text (for Chapter III):')
    r4.font.name = FONT; r4.font.size = Pt(10); r4.font.bold = True; r4.font.color.rgb = MAROON

    met = 'met' if overall_avg >= 80 else 'did not meet'
    interp_text = (
        f'The overall SUS score for SmartSpend was {overall_avg:.2f} (n = {len(all_scores)}), '
        f'which corresponds to a grade of "{grade_o}" — categorized as "{adj_o}" '
        f'per Bangor et al. (2009) — and falls within the "{accept_o}" range '
        f'per the standard SUS acceptability classification (Brooke, 1996). '
        f'This result {met} the pre-defined usability target of ≥ 80.00 (Good). '
        f'Parents aged 35 to 55 recorded an average SUS score of {parent_avg:.2f} '
        f'({adj_p}), while young professionals aged 21 to 35 recorded an average '
        f'of {youngprof_avg:.2f} ({adj_y}). '
        f'These results indicate that SmartSpend demonstrates '
        f'{"strong" if overall_avg >= 80 else "moderate"} usability across both target demographic groups.'
    )

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(3)
    p5.paragraph_format.space_after  = Pt(3)
    r5 = p5.add_run(interp_text)
    r5.font.name = FONT; r5.font.size = Pt(11)

    p6 = doc.add_paragraph()
    r6 = p6.add_run('References: Brooke, J. (1996). SUS: A quick and dirty usability scale. '
                     'In Usability evaluation in industry (pp. 189–194). Taylor & Francis.  |  '
                     'Bangor, A., Kortum, P., & Miller, J. (2009). Determining what individual SUS scores mean. '
                     'Journal of Usability Studies, 4(3), 114–123.')
    r6.font.name = FONT; r6.font.size = Pt(9); r6.font.italic = True

    out_path = OUT_DIR / 'SUS_Results_Chapter3.docx'
    doc.save(str(out_path))
    return out_path, overall_avg, grade_o, adj_o


def parse_google_forms_csv(csv_path):
    """
    Parse Google Forms CSV export.
    Expected columns: Timestamp, [optional: group column], Q1, Q2, ..., Q10
    The script auto-detects which columns contain the 10 SUS items.
    """
    results = []

    with open(csv_path, newline='', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        headers = next(reader)

    print(f'\nCSV Headers ({len(headers)} columns):')
    for i, h in enumerate(headers):
        print(f'  [{i}] {h[:80]}')

    # Auto-detect SUS question columns (look for Q1/Q2/... or the SUS item text)
    sus_col_indices = []
    sus_keywords = [
        'frequently', 'complex', 'easy to use', 'support of a technical',
        'well integrated', 'inconsistency', 'learn to use', 'cumbersome',
        'confident', 'learn a lot'
    ]

    for i, h in enumerate(headers):
        h_lower = h.lower()
        if any(kw in h_lower for kw in sus_keywords):
            sus_col_indices.append(i)

    # Fallback: look for columns labeled Q1-Q10 or 1.-10.
    if len(sus_col_indices) < 10:
        sus_col_indices = []
        for i, h in enumerate(headers):
            h_strip = h.strip().lower()
            if (h_strip.startswith('q') and h_strip[1:].split('.')[0].isdigit()):
                sus_col_indices.append(i)

    # Last fallback: use the last 10 numeric columns
    if len(sus_col_indices) < 10:
        print('\n  Could not auto-detect SUS columns. Using last 10 columns as Q1-Q10.')
        sus_col_indices = list(range(len(headers) - 10, len(headers)))

    print(f'\nDetected SUS columns: {sus_col_indices}')
    print(f'Mapped to: {[headers[i][:50] for i in sus_col_indices]}\n')

    # Detect group column (looks for parent/young professional labels)
    group_col = None
    for i, h in enumerate(headers):
        h_lower = h.lower()
        if 'role' in h_lower or 'group' in h_lower or 'parent' in h_lower or 'primary' in h_lower:
            group_col = i
            break

    with open(csv_path, newline='', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        next(reader)  # skip header
        row_num = 0
        for row in reader:
            if not any(row):  # skip empty rows
                continue
            row_num += 1

            # Determine group
            group = 'parent'
            if group_col is not None and group_col < len(row):
                val = row[group_col].lower()
                if 'young' in val or 'professional' in val or '21' in val:
                    group = 'youngprof'
            else:
                # Default: first 20 = parents, next 10 = young professionals
                group = 'parent' if row_num <= 20 else 'youngprof'

            # Extract raw scores
            try:
                raw = []
                for ci in sus_col_indices[:10]:
                    if ci < len(row) and row[ci].strip():
                        val = row[ci].strip()
                        # Handle "1 - Strongly Disagree" style answers
                        raw.append(int(val[0]) if val[0].isdigit() else int(val))
                    else:
                        raw.append(3)  # default neutral if missing

                if len(raw) < 10:
                    raw.extend([3] * (10 - len(raw)))

                # Compute adjusted scores
                adjusted = []
                for i, r in enumerate(raw):
                    if (i + 1) % 2 == 1:
                        adjusted.append(r - 1)
                    else:
                        adjusted.append(5 - r)

                sum_adj = sum(adjusted)
                sus_score = sum_adj * 2.5

                results.append({
                    'respondent_no': row_num,
                    'group': group,
                    'raw': raw,
                    'adjusted': adjusted,
                    'sum_adj': sum_adj,
                    'sus_score': sus_score,
                })
            except (ValueError, IndexError) as e:
                print(f'  Warning: could not parse row {row_num}: {e}')
                continue

    return results


def demo_mode():
    """Generate a sample output using simulated scores for testing."""
    import random
    random.seed(42)
    print('\n--- DEMO MODE (no CSV provided) ---')
    print('Generating simulated SUS scores for 30 respondents...')
    print('Replace with real data by running: python compute_sus_scores.py your_file.csv\n')

    results = []
    groups = ['parent'] * 20 + ['youngprof'] * 10

    # Simulate realistic SUS responses (targeting ~82-85 overall)
    for i, group in enumerate(groups):
        raw = []
        for q in range(1, 11):
            if q % 2 == 1:  # odd (positive): tend toward 4-5
                raw.append(random.choices([3, 4, 4, 5, 5], k=1)[0])
            else:            # even (negative): tend toward 1-2
                raw.append(random.choices([1, 1, 2, 2, 3], k=1)[0])

        adjusted = []
        for j, r in enumerate(raw):
            if (j + 1) % 2 == 1:
                adjusted.append(r - 1)
            else:
                adjusted.append(5 - r)

        sum_adj = sum(adjusted)
        results.append({
            'respondent_no': i + 1,
            'group': group,
            'raw': raw,
            'adjusted': adjusted,
            'sum_adj': sum_adj,
            'sus_score': sum_adj * 2.5,
        })

    return results, 'demo_simulated_data.csv'


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    # Build the survey form guides first
    print('Building survey form guides...')
    try:
        exec(open(Path(__file__).parent / 'build_survey_forms.py').read())
    except Exception as e:
        print(f'  (build_survey_forms.py: {e})')

    # Determine CSV source
    if len(sys.argv) > 1:
        csv_path = sys.argv[1]
        if not os.path.exists(csv_path):
            print(f'Error: CSV file not found: {csv_path}')
            sys.exit(1)
        results = parse_google_forms_csv(csv_path)
        source = csv_path
    else:
        csv_path = find_csv()
        if csv_path:
            print(f'Found CSV: {csv_path}')
            results = parse_google_forms_csv(str(csv_path))
            source = str(csv_path)
        else:
            results, source = demo_mode()

    if not results:
        print('No valid data found. Check your CSV file.')
        sys.exit(1)

    print(f'\nProcessed {len(results)} respondents:')
    print(f'  Parents (35–55):         {sum(1 for r in results if r["group"] == "parent")}')
    print(f'  Young Professionals:     {sum(1 for r in results if r["group"] == "youngprof")}')
    print()

    for res in results:
        grade, adj, accept = interpret_sus(res['sus_score'])
        print(f'  R{res["respondent_no"]:02d} ({res["group"][:6]}): '
              f'SUS = {res["sus_score"]:.1f}  ({adj})')

    out_path, overall, grade, adj = build_results_docx(results, source)

    print(f'\n{"="*60}')
    print(f'OVERALL SUS SCORE: {overall:.2f}')
    grade_full, adj_full, accept_full = interpret_sus(overall)
    print(f'Grade:             {grade_full}')
    print(f'Adjective:         {adj_full}')
    print(f'Acceptability:     {accept_full}')
    print(f'Target (≥ 80):     {"✅ MET" if overall >= 80 else "❌ NOT MET"}')
    print(f'{"="*60}')
    print(f'\nResults DOCX saved: {out_path}')
    print('Open it and copy the table + interpretation text into Chapter III.')
