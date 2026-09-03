"""
build_validator_letters.py  —  Two validator briefing letters for SmartSpend
  Letter 1: Content Validator (survey questionnaire)
  Letter 2: Technical Validator (system evaluation + SUS)
Run:  python build_validator_letters.py
Out:  docs/manuscript/output/Validator_Briefing_Letters.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT  = Path(__file__).parent / '..' / 'output' / 'Validator_Briefing_Letters.docx'
FONT = 'Tahoma'
MAROON = RGBColor(0x5C, 0x0E, 0x24)
DARK   = RGBColor(0x22, 0x22, 0x22)
BLUE   = RGBColor(0x2A, 0x4A, 0x7F)

doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11.0)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(1.25)
sec.right_margin  = Inches(1.25)
doc.styles['Normal'].font.name = FONT
doc.styles['Normal'].font.size = Pt(11)


def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)


def para(text, bold=False, italic=False, size=11, align='left',
         color=None, before=0, after=6, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.alignment = {'left':   WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right':  WD_ALIGN_PARAGRAPH.RIGHT,
                   'both':   WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p


def rule():
    """Horizontal divider paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('─' * 80)
    r.font.name = FONT; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def section_header(text):
    para(text, bold=True, size=11, color=MAROON, before=10, after=4)


def school_header():
    para('LORMA COLLEGES', bold=True, size=14, align='center', color=MAROON, before=0, after=2)
    para('College of Computer Studies and Engineering', size=11, align='center', after=2)
    para('Bachelor of Science in Information Technology — 4th Year, 1st Semester, AY 2026–2027',
         size=10, align='center', after=6, italic=True)
    rule()


def signature_block(name_line='________________________________',
                    title_line='Name and Signature',
                    date_line='Date: _______________'):
    para('')
    para(name_line, bold=True, size=11, before=24, after=2)
    para(title_line, size=10, italic=True, after=2)
    para(date_line, size=10, after=2)


# ══════════════════════════════════════════════════════════════════════════════
# LETTER 1 — CONTENT VALIDATOR (Survey Questionnaire)
# ══════════════════════════════════════════════════════════════════════════════

school_header()

para('CONTENT VALIDATION REQUEST LETTER', bold=True, size=13, align='center',
     color=MAROON, before=8, after=4)
para('Survey Questionnaire — Financial Management Practices',
     size=11, align='center', italic=True, after=12)

para('September 2026', size=11, align='right', after=12)

para('Dear Respected Validator,', bold=True, size=11, after=8)

para(
    'We, the members of Lucid Frame — BSIT 4th Year research group of Lorma Colleges, '
    'College of Computer Studies and Engineering — are currently conducting a capstone '
    'project titled:',
    size=11, align='both', after=6, indent=True
)

para(
    '"SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory Application '
    'for Personal Financial Management"',
    bold=True, size=11, align='center', color=BLUE, before=4, after=8
)

para(
    'In partial fulfillment of the requirements for the degree of Bachelor of Science '
    'in Information Technology, we are conducting a research study to assess the '
    'existing financial management practices, common budgeting challenges, and expense '
    'tracking behaviors of our target population: parents aged 35–55 (primary) and '
    'young professionals aged 21–35 (secondary) in San Fernando City, La Union.',
    size=11, align='both', after=8, indent=True
)

para(
    'In this regard, we would like to request your valuable expertise in validating '
    'the enclosed survey questionnaire. Your role as a Content Validator is to assess '
    'whether the survey items:',
    size=11, align='both', after=6, indent=True
)

bullets = [
    'Are clearly worded, unambiguous, and understandable to the target population',
    'Accurately measure the construct they are intended to assess (financial management practices, budgeting behaviors, and expense tracking habits)',
    'Are appropriate in scope, relevance, and cultural sensitivity for Filipino respondents',
    'Are free from leading, double-barreled, or biased questions',
    'Collectively cover the key thematic areas needed to achieve Research Objective 1',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(b)
    r.font.name = FONT; r.font.size = Pt(11)

para('')

para(
    'Your expert opinion and recommendations will be used to refine the questionnaire '
    'before it is administered to the 30 respondents. You are free to add written '
    'comments directly on the questionnaire form.',
    size=11, align='both', after=8, indent=True
)

section_header('Qualification of Content Validator')
para(
    'We are looking for a validator with a background in any of the following: '
    'BS Commerce, BS Accountancy, BS Business Administration, Financial Management, '
    'Economics, or a related field, with relevant experience in financial management, '
    'business operations, financial planning, or financial literacy advocacy.',
    size=11, align='both', after=8, indent=True
)

section_header('Validation Certificate (Please Complete)')
para('Please fill in and sign the Validation Certificate attached to the survey questionnaire '
     '(Appendix A of the manuscript). The certificate requires your educational background, '
     'occupation, years of experience, and signature.',
     size=11, align='both', after=8, indent=True)
para('Note: Your name on the certificate is optional. Your professional credentials '
     'are sufficient to establish your qualifications for the validation process.',
     size=11, align='both', italic=True, color=RGBColor(0x55,0x55,0x55), after=12, indent=True)

para('Respectfully yours,', size=11, before=8, after=4)
para('Lucid Frame Research Group', bold=True, size=11, after=2)
para('Brix A. Directo  |  Cyrille John M. Rubis  |  Djaunathan Albert S. Madayag', size=10, after=2)
para('BSIT 4th Year — Lorma Colleges CCSE  |  AY 2026–2027, 1st Semester', size=10, italic=True, after=2)
para('Noted by: Janelli M. Mendez, DIT  (Capstone Adviser)', size=10, after=12)

signature_block(
    name_line='________________________________',
    title_line='Validator Signature over Printed Name',
    date_line='Date: _______________'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# LETTER 2 — TECHNICAL VALIDATOR (System + SUS)
# ══════════════════════════════════════════════════════════════════════════════

school_header()

para('TECHNICAL VALIDATION REQUEST LETTER', bold=True, size=13, align='center',
     color=MAROON, before=8, after=4)
para('System Evaluation and SUS Usability Assessment',
     size=11, align='center', italic=True, after=12)

para('September 2026', size=11, align='right', after=12)

para('Dear Respected Validator,', bold=True, size=11, after=8)

para(
    'We are writing to respectfully request your technical expertise in validating '
    'the SmartSpend mobile application as part of our capstone project:',
    size=11, align='both', after=6, indent=True
)

para(
    '"SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory Application '
    'for Personal Financial Management"',
    bold=True, size=11, align='center', color=BLUE, before=4, after=8
)

para(
    'SmartSpend is an Android application built with Flutter/Dart that integrates '
    'a multi-provider agentic Large Language Model (LLM) architecture to enable '
    '31 autonomous financial management actions through natural language, voice, '
    'camera, and batch screenshot import. The system features an offline-first '
    'SQLite database, Firebase cloud synchronization, and a proprietary '
    'Financial Health Score (FHS) algorithm.',
    size=11, align='both', after=8, indent=True
)

para(
    'As our Technical Validator, we are requesting your assessment of the following:',
    size=11, align='both', after=6, indent=True
)

tech_bullets = [
    'Technical soundness of the system architecture (Flutter, SQLite v11, Firebase, multi-provider LLM)',
    'Appropriateness and correctness of the Financial Health Score (FHS) formula and its academic basis',
    'Soundness of the System Usability Scale (SUS) administration procedure, including the 10-item questionnaire, scoring formula (Brooke, 1996), and interpretation scale (Bangor et al., 2009)',
    'Suitability of the SUS as an evaluation instrument for this type of mobile application',
    'Adequacy of the purposive sampling strategy (30 respondents: 20 parents 35–55, 10 young professionals 21–35)',
    'Overall technical completeness and readiness for pre-final defense',
]
for b in tech_bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(b)
    r.font.name = FONT; r.font.size = Pt(11)

para('')

section_header('System Access for Evaluation')
para(
    'A demonstration of the SmartSpend application will be arranged at your convenience. '
    'Alternatively, an APK installation file can be provided for independent evaluation. '
    'A Demo Mode is available that pre-populates the app with sample Filipino student '
    'financial data — no account registration required.',
    size=11, align='both', after=8, indent=True
)
para('GitHub Repository: https://github.com/Zushikina-kun/smartspend-app',
     size=10, italic=True, color=BLUE, after=8, indent=True)

section_header('Qualification of Technical Validator')
para(
    'We are looking for a validator with a background in BS Information Technology, '
    'BS Computer Science, Software Engineering, or a related IT field, with experience '
    'in software development, systems analysis, mobile application development, or '
    'usability evaluation.',
    size=11, align='both', after=8, indent=True
)

section_header('Validation Certificate (Please Complete)')
para(
    'Please fill in and sign the Technical Validation Certificate attached '
    '(Appendix A of the manuscript). The certificate requires your educational '
    'background, occupation, years of experience in IT/software development '
    'and/or usability evaluation, and signature.',
    size=11, align='both', after=8, indent=True
)
para('Note: Your name on the certificate is optional. Your professional credentials '
     'are sufficient to establish your qualifications.',
     size=11, align='both', italic=True, color=RGBColor(0x55,0x55,0x55), after=12, indent=True)

para('Respectfully yours,', size=11, before=8, after=4)
para('Lucid Frame Research Group', bold=True, size=11, after=2)
para('Brix A. Directo  |  Cyrille John M. Rubis  |  Djaunathan Albert S. Madayag', size=10, after=2)
para('BSIT 4th Year — Lorma Colleges CCSE  |  AY 2026–2027, 1st Semester', size=10, italic=True, after=2)
para('Noted by: Janelli M. Mendez, DIT  (Capstone Adviser)', size=10, after=12)

signature_block(
    name_line='________________________________',
    title_line='Validator Signature over Printed Name',
    date_line='Date: _______________'
)

doc.save(str(OUT))
print(f'  Validator letters saved: {OUT.name}  ({OUT.stat().st_size//1024} KB)')
