"""
build_technical_reference.py
Generates SmartSpend_Technical_Reference.docx — a standalone, panel-ready
document covering:
  1. Financial Health Score (FHS) — full formula, academic basis, comparison
  2. Financial Management Score (FMS) — formula, rationale, defense Q&A
  3. LLM Comparative Benchmarking — all 19 models, 4 task types, selection rationale
  4. App Feature Comparison — full matrix, 22 apps, SmartSpend gaps & leads
  5. Academic Reference List (APA 7th)

Run:  python build_technical_reference.py
Output: docs/manuscript/output/SmartSpend_Technical_Reference.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = Path(__file__).parent
OUTDIR = BASE / '..' / 'output'
OUTDIR.mkdir(exist_ok=True)
OUT    = OUTDIR / 'SmartSpend_Technical_Reference.docx'

# ── Colour constants ───────────────────────────────────────────────────────────
MAROON  = RGBColor(0x5C, 0x0E, 0x24)
GOLD    = RGBColor(0xC9, 0xA8, 0x4C)
DARK    = RGBColor(0x22, 0x22, 0x22)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LGREY   = RGBColor(0xD9, 0xD9, 0xD9)
MGREY   = RGBColor(0xF2, 0xF2, 0xF2)
BLUE    = RGBColor(0x2A, 0x6E, 0xAF)
GREEN   = RGBColor(0x2E, 0x7D, 0x32)

FONT = 'Tahoma'

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins — 1-inch all sides
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
for attr in ('top_margin','bottom_margin','left_margin','right_margin'):
    setattr(sec, attr, Inches(1.0))

# Normal style
ns = doc.styles['Normal']
ns.font.name = FONT
ns.font.size = Pt(11)
ns.paragraph_format.space_before = Pt(0)
ns.paragraph_format.space_after  = Pt(4)

# ── Helpers ────────────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _cell_borders(cell, color='AAAAAA', sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=None):
    """Add a heading paragraph with Lorma styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    run.font.size = Pt({1: 14, 2: 12, 3: 11}.get(level, 11))
    run.font.color.rgb = color or (MAROON if level == 1 else (BLUE if level == 2 else DARK))
    if level == 1:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(text, italic=False, indent=False, before=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name   = FONT
    run.font.size   = Pt(10.5)
    run.font.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)
    return p

def add_formula(text):
    """Monospace-styled formula block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.4)
    p.paragraph_format.space_before  = Pt(3)
    p.paragraph_format.space_after   = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    return p

def add_table(rows, col_widths=None, header_shade='5C0E24', row_shade='F2F2F2',
              alt_shade='FFFFFF', font_size=9):
    """
    rows: list of lists (first row = header).
    col_widths: list of Inches values; if None, equal distribution across 6.5".
    """
    ncols = max(len(r) for r in rows)
    if col_widths is None:
        w = 6.5 / ncols
        col_widths = [Inches(w)] * ncols

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    for ri, row in enumerate(rows):
        is_hdr = (ri == 0)
        tr = tbl.rows[ri]
        for ci in range(ncols):
            cell = tr.cells[ci]
            if ci < len(col_widths):
                cell.width = col_widths[ci]
            txt  = row[ci] if ci < len(row) else ''

            # Background
            if is_hdr:
                _shade_cell(cell, header_shade)
            elif ri % 2 == 0:
                _shade_cell(cell, alt_shade)
            else:
                _shade_cell(cell, row_shade)

            # Text
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(txt))
            run.font.name  = FONT
            run.font.size  = Pt(font_size)
            run.font.bold  = is_hdr
            if is_hdr:
                run.font.color.rgb = WHITE

    return tbl

def add_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.name  = FONT
    run.font.size  = Pt(9.5)
    run.font.bold  = True
    run.font.color.rgb = DARK

def add_defense_box(question, answer):
    """Shaded box for panel defense Q&A."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    q_cell = tbl.rows[0].cells[0]
    _shade_cell(q_cell, 'EAD9E0')
    qp = q_cell.paragraphs[0]
    qp.paragraph_format.space_before = Pt(3)
    qp.paragraph_format.space_after  = Pt(3)
    qr = qp.add_run(f'❓ Panel Question:  {question}')
    qr.font.name = FONT; qr.font.size = Pt(10); qr.font.bold = True
    qr.font.color.rgb = MAROON

    a_cell = tbl.rows[1].cells[0]
    _shade_cell(a_cell, 'F9F4F6')
    ap = a_cell.paragraphs[0]
    ap.paragraph_format.space_before = Pt(3)
    ap.paragraph_format.space_after  = Pt(3)
    ap.paragraph_format.left_indent  = Inches(0.15)
    ar = ap.add_run(f'✅ Answer:  {answer}')
    ar.font.name = FONT; ar.font.size = Pt(10)
    ar.font.color.rgb = DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_apa(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent    = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_before   = Pt(2)
    p.paragraph_format.space_after    = Pt(2)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10)

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
r = p.add_run('SmartSpend')
r.font.name = FONT; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = MAROON

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Technical Reference Document')
r.font.name = FONT; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Financial Health Score Formula · FMS Formula · LLM Comparative Study · App Feature Comparison')
r.font.name = FONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = BLUE

doc.add_paragraph()
for line in [
    'Version: 2.9.10',
    'Group: Lucid Frame (Directo, Rubis, Madayag)',
    'Lorma Colleges — CCSE, BSIT 4th Year',
    'Academic Year: 2026–2027, 1st Semester',
    'Prepared for: Pre-Final Defense & Panel Evaluation',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = FONT; r.font.size = Pt(10.5)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — FINANCIAL HEALTH SCORE (FHS)
# ══════════════════════════════════════════════════════════════════════════════

add_heading('SECTION 1 — FINANCIAL HEALTH SCORE (FHS)', 1)

add_body(
    'The Financial Health Score (FHS) is SmartSpend\'s core academic contribution — '
    'a 0-to-100 behavioral metric computed in real-time from the user\'s own recorded '
    'transaction data. Unlike survey-based instruments (CFPB, 2017), SmartSpend\'s FHS '
    'requires no self-report: every component is computed directly from SQLite records. '
    'It operates in two modes to accommodate the diverse income structures of Filipino users.'
)

# 1.1 Academic frameworks
add_heading('1.1  Existing Frameworks This Study Builds On', 2)

add_table([
    ['Framework', 'Published By', 'Year', 'What It Measures', 'How SmartSpend Uses It'],
    ['FinHealth Score®',
     'Financial Health Network',
     '2021 / 2026 update',
     '4 pillars: Spend, Save, Borrow, Plan/Protect\n8 measurable indicators',
     'Savings Rate ↔ Save pillar\nOverspend Control ↔ Spend pillar\nStructure of 4 equal-weight components'],
    ['UNSGSA Financial Health Framework',
     'UN Secretary-General\'s Special Advocate',
     '2021',
     'Ability to meet obligations, feel secure, and make life-affirming choices',
     'Four-component structure reflects UNSGSA\'s emphasis on spending, savings, and consistency behaviors'],
    ['CFPB Financial Well-Being Scale',
     'US Consumer Financial Protection Bureau',
     '2017',
     '10-question survey → 0–100 well-being score',
     'Provides validated 0–100 scale and definition of financial well-being; SmartSpend computes from behavioral data instead of survey'],
    ['FinHealth Score® 2026 Update',
     'Financial Health Network',
     '2026',
     'Renamed pillar 4 to "Plan and Protect"; emphasises separating tracking behavior from health outcomes',
     'Basis for adding FMS as a separate metric (Part 14 of RESEARCH_BASIS.md)'],
], col_widths=[Inches(1.3), Inches(1.5), Inches(0.8), Inches(1.7), Inches(1.7)], font_size=9)
add_caption('Table 1.1. Academic Frameworks Informing the SmartSpend Financial Health Score')

# 1.2 Full Mode Formula
add_heading('1.2  FHS Full Mode — Income Tracking Enabled', 2)
add_body(
    'Full Mode is activated when the user has set a monthly income and enabled '
    'income-wallet tracking. It computes four components each worth 25 points '
    '(maximum: 100). The final score is clamped to the range [0, 100].'
)

add_heading('Component 1 — Savings Rate (25 pts)', 3)
add_formula('savingsRate = (monthlyIncome − totalSpentThisMonth) / monthlyIncome')
add_formula('comp1 = 25 × min(1.0, savingsRate / 0.20)')
add_body(
    'Interpretation: Full 25 pts when saving ≥ 20% of income; scales proportionally '
    'below that. If income is not set, partial credit is awarded (20 pts if spent < ₱5,000; '
    '15 pts if < ₱10,000; 10 pts otherwise) with an "Unmeasured" flag shown in the UI.',
    before=2
)
add_body(
    'Academic Basis: The 20% savings target is drawn directly from the 50/30/20 budgeting '
    'rule (Warren & Tyagi, 2005): 50% Needs, 30% Wants, 20% Savings/Debt. This rule is '
    'one of the most widely cited personal finance frameworks in academic and popular '
    'finance literature.'
)

add_heading('Component 2 — Overspend Control (25 pts)', 3)
add_formula('dailyBudget = monthlyIncome / daysInMonth')
add_formula('overDays = count of logged days where dailySpend > dailyBudget')
add_formula('comp2 = 25 × (1 − overDays / activeDays)')
add_body(
    'Interpretation: Full 25 pts when no logged day exceeded the daily budget; '
    'decreases proportionally as more days overshoot. If income is not set, '
    '20 pts partial credit is awarded.',
    before=2
)
add_body(
    'Academic Basis: Day-level spending discipline directly implements the FinHealth '
    'Score® Spend pillar — "spending less than income" as a measurable behavior '
    '(Financial Health Network, 2021). The day-level granularity is adapted from '
    'MindsBudget\'s transaction-based approach, which the research review identified '
    'as the most directly comparable system to SmartSpend.'
)

add_heading('Component 3 — Budget Adherence (25 pts)', 3)
add_formula('comp3 = 25 × (onBudgetCategories / totalConfiguredBudgets)')
add_body(
    'Interpretation: Full 25 pts when all configured category budgets are on track; '
    'scales proportionally with each exceeded budget. If no budgets are configured, '
    'full 25 pts is awarded (users are not penalized for not setting budgets — '
    'the system encourages but does not require them).',
    before=2
)
add_body(
    'Percentage-based budgets (e.g., "25% of income for Food") are resolved to their '
    'actual peso amount before comparison using the stored monthly income value.'
)
add_body(
    'Academic Basis: Category-level budgeting is a core feature of zero-based '
    'budgeting theory (Ramsey, 2003). YNAB\'s research (cited in Yomio, 2026) '
    'found that users who set specific category budgets overspend 32% less than '
    'those who track without budgets.'
)

add_heading('Component 4 — Logging Consistency (25 pts)', 3)
add_formula('loggedDays = unique dates in expenses table for current calendar month')
add_formula('activeDays = daysPassed (unless user started after the 7th — then span from first entry)')
add_formula('activeDays = min(activeDays, loggedDays × 2)   [retroactive bulk-entry fairness cap]')
add_formula('comp4 = 25 × min(1.0, loggedDays / activeDays)')
add_body(
    'Interpretation: Full 25 pts when the user has logged at least once per active '
    'day. The formula has two fairness adjustments:',
    before=2
)
add_bullet('Mid-month grace: if the user\'s first entry was after the 7th, span is counted from that entry (prevents penalizing late starters).')
add_bullet('Bulk-entry cap: activeDays is capped at loggedDays × 2 so catching up on missed days with bulk entries doesn\'t unfairly tank the score.')
add_body(
    'Scoped to current month only — historical expenses from prior months are '
    'excluded to prevent very old entries from distorting the current-month score.'
)
add_body(
    'Academic Basis: Consistent financial tracking has been shown to reduce '
    'discretionary spending by 10–20% through increased financial awareness '
    '(Mindfulsuite, 2026). The behavioral tracking habit itself drives better '
    'financial decisions (Thaler & Sunstein, 2008). Note: based on the FHN '
    '2026 update recommendation, Logging Consistency was also separated into '
    'the standalone Financial Management Score (Section 2).'
)

add_heading('Score Adjustments', 3)
add_body('Two post-computation adjustments are applied on top of the raw 4-component sum:')

add_body('Warning Decay  (−5 pts/day, maximum −15 pts total):', before=6)
add_formula('penalty = min(warningDecayDays, 3) × 5')
add_formula('adjustedScore = rawScore − penalty')
add_body(
    'Triggered when the user continues spending in an over-budget category after '
    'a budget warning has been issued. The decay counter increments once per day '
    '(maximum 3 days) and resets when all budgets return to within limits.',
    before=2
)
add_body(
    'Academic Basis: Loss aversion theory (Kahneman & Tversky, 1979) establishes '
    'that losses feel approximately twice as painful as equivalent gains feel '
    'pleasurable. The Warning Decay mechanism makes the consequence of ignoring '
    'budget warnings tangible and numerically visible — applying nudge theory '
    '(Thaler & Sunstein, 2008).'
)

add_body('Gap Adjustment  (−3 pts/day, maximum −15; or +2 pts/day, maximum +10):', before=6)
add_formula('penalty = min(gapPenaltyDays × 3, 15)')
add_formula('bonus   = min(gapCleanDays × 2, 10)')
add_formula('adjustedScore = rawScore − penalty + bonus')
add_body(
    'When the app detects a multi-day logging gap, a startup prompt asks the user: '
    '"Did you have expenses during [date range]?" If the user confirms they spent '
    'but forgot to log, penalty days are recorded (−3 pts/day). If they confirm '
    'those were genuine no-spend days, bonus days are recorded (+2 pts/day).',
    before=2
)
add_body(
    'Academic Basis: This mechanism implements accurate self-monitoring, grounded '
    'in behavioral finance research on self-reporting accuracy (Ariely, 2008).'
)

# 1.3 Lightweight Mode
page_break()
add_heading('1.3  FHS Lightweight Mode — Income Tracking Disabled', 2)
add_body(
    'Lightweight Mode is activated when the user disables income-wallet tracking — '
    'designed for students, freelancers, and informal workers who do not have a '
    'predictable monthly income. The four components are redesigned to use only '
    'spending behavior data.'
)
add_body(
    'Academic Basis for dual-mode design: The Financial Health Network (2021) '
    'explicitly acknowledges that financial health metrics must adapt to diverse '
    'income structures. The BSP (2021) confirms that a significant portion of '
    'Filipino adults are in informal employment or have irregular income.'
)

add_table([
    ['Component', 'Formula', 'Full Score Condition', 'Academic Basis'],
    ['Spending Restraint\n(25 pts)',
     'If limit set:\ncomp1 = 25 × max(0, (1 − ((ratio − 0.8) / 1.2)))\nwhere ratio = spent / limit\n\nIf no limit:\ncomp1 = 10–25 pts based on\nWant/Need ratio (≤30% wants = 25 pts)',
     'Spending ≤ 80% of user-set limit',
     'Zero-based budgeting (Ramsey, 2003)\nWant/Need ratio from 50/30/20 rule (Warren & Tyagi, 2005)'],
    ['Logging Consistency\n(25 pts)',
     'Same formula as Full Mode',
     'Logging every active day',
     'Tracking reduces discretionary spending\n10–20% (Mindfulsuite, 2026)'],
    ['Category Balance\n(25 pts)',
     'comp3 = 25 × max(0, (1 − ((topRatio − 0.4) / 0.6)))\nwhere topRatio = topCategory / totalSpent\n\nIf budgets set: replaced by Budget Adherence',
     'No single category > 40% of total spending',
     'Spending diversification principle\nBudget Adherence (Ramsey, 2003) if budgets exist'],
    ['Habit Streak\n(25 pts)',
     'streak = consecutive days ending today\n        with ≥1 logged expense\ncomp4 = 25 × min(1.0, streak / 14)',
     '14+ consecutive logging days',
     'Habit formation theory — 21-day habit rule\nadapted to 14-day milestone (Duhigg, 2012)'],
], col_widths=[Inches(1.2), Inches(2.2), Inches(1.5), Inches(1.6)], font_size=8.5)
add_caption('Table 1.2. FHS Lightweight Mode — Components, Formulas, and Academic Basis')

# 1.4 Score labels
add_heading('1.4  FHS Score Classification and Labels', 2)

add_table([
    ['Score Range', 'Label', 'Color (UI)', 'Interpretation'],
    ['90–100', '🏆 Excellent', 'Deep Green', 'Top-tier financial health; all components performing optimally'],
    ['75–89',  '✅ Good',      'Green',      'Strong financial habits; minor improvements possible'],
    ['60–74',  '📈 Fair',      'Teal/Teal',  'On track but notable gaps in one or more components'],
    ['45–59',  '⚠️ Needs Work','Orange',     'Multiple components below optimal; action required'],
    ['0–44',   '🚨 Poor',      'Red',        'Significant financial management challenges present'],
], col_widths=[Inches(1.0), Inches(1.2), Inches(1.2), Inches(3.1)], font_size=10)
add_caption('Table 1.3. FHS Score Classification (5-Tier Scale, v2.9.5+)')

# 1.5 FHS comparison
add_heading('1.5  SmartSpend FHS vs Other Financial Health Scoring Systems', 2)
add_body(
    'A review of 9 existing financial health scoring systems (conducted August 2026) '
    'was used to validate SmartSpend\'s design decisions and identify gaps.'
)

add_table([
    ['System', 'Scale', 'Core Components', 'Computation Method', 'Key Comparison Point'],
    ['Financial Health Network\nFinHealth Score®', '0–100',
     'Spend, Save, Borrow,\nPlan/Protect (8 indicators)', 'Survey + behavioral indicators\n(requires external data feeds)',
     'Primary academic reference;\nSmartSpend uses same pillar structure\nbut computes from SQLite (no bank API)'],
    ['CFPB Financial Well-Being Scale', '0–100',
     'Security, freedom, on-track,\nabsorb shocks', '10-question self-report survey',
     'Validated 0–100 range;\nSmartSpend replaces survey with\nbehavioral transaction data'],
    ['MindsBudget', '0–100',
     'Spending Rate (30%), Discipline (25%),\nMonth Stability (20%), Emergency Runway (25%)',
     'Transaction-based computation',
     'Most directly comparable to SmartSpend;\nboth use transaction data without bank sync'],
    ['Rateweb', '0–100',
     'Savings (22%), Debt (18%), Emergency (18%),\nNet Worth (15%), Goals (12%), Others (25%)',
     'Transaction + manual data input',
     'More components; shows SmartSpend\'s\n4-component equal-weight model is defensible'],
    ['Elenvo', '0–100',
     '6 dimensions: Retirement, Emergency,\nDebt, Cash Flow, Protection, Tax',
     'Manual data entry per dimension',
     'Explicit "unmeasured" handling;\nSmartSpend adopted this for income-absent state'],
    ['Wingman Money (AU, 2026)', '0–100',
     'Daily behaviors + long-term wellbeing\nindicators based on FHN framework',
     'Transaction-based, mobile app',
     'Closest commercial equivalent;\nvalidates behavioral FHS is viable.\nSmartSpend adds dual-mode and offline-first'],
    ['BudgetPH / KindlyF', '0–100',
     'Budget adherence + streak tracking\n(simpler formula, fewer components)',
     'Transaction-based, Filipino app',
     'Only PH competitor with FHS;\nSmartSpend has more academic rigor\nand dual-mode support'],
    ['Cleo', 'Score (0–100)',
     'Savings behavior + bill payment\n(formula not published)',
     'Bank sync + behavior hybrid',
     'Opaque formula; not academically\nattributable; requires bank connectivity'],
    ['SmartSpend FHS', '0–100',
     'Full Mode: Savings Rate, Overspend Control,\nBudget Adherence, Logging Consistency\nLightweight: Restraint, Consistency,\nBalance, Habit Streak',
     'Fully deterministic; computed\nfrom SQLite; no surveys; no bank API;\nworks fully offline',
     'Only dual-mode behavioral FHS;\nformula fully documented and traceable\nto academic frameworks (FinHN, UNSGSA, CFPB)'],
], col_widths=[Inches(1.5), Inches(0.5), Inches(1.7), Inches(1.3), Inches(1.5)], font_size=8)
add_caption('Table 1.4. Comparative Review of Financial Health Scoring Systems (August 2026)')

# Defense Q&A for FHS
add_heading('1.6  Panel Defense Q&A — FHS', 2)

add_defense_box(
    'Why does SmartSpend have its own Financial Health Score? Isn\'t there already a standard?',
    'Existing standards like the CFPB Financial Well-Being Scale (2017) and the FinHealth Score® (Financial Health Network, 2021) are designed for institutional measurement — they require survey responses or bank data feeds. SmartSpend\'s FHS computes the same concept from the user\'s own recorded transaction data, entirely offline. This is the key academic contribution: a behavioral, real-time, survey-free financial health metric for mobile users without bank connectivity.'
)
add_defense_box(
    'Why 4 equal-weight components? Why not weight them differently?',
    'The equal 25%/25%/25%/25% weighting is intentional and defensible. Each component addresses a distinct dimension of financial health: income management (Savings Rate), daily control (Overspend Control), category discipline (Budget Adherence), and data quality (Logging Consistency). The FinToolSuite review confirms that an equal-weight model is the simplest academically defensible structure when no longitudinal data exists to calibrate differential weights for a specific population.'
)
add_defense_box(
    'Why is Logging Consistency still in the FHS?',
    'In the absence of bank API connectivity — which is not available for most Philippine banks (BSP Open Finance only launched pilot in July 2025) — user-entered expense data is the sole input for all FHS components. Logging Consistency therefore serves a dual role: it measures tracking behavior AND is a prerequisite for the other three components to compute accurately. This is explicitly acknowledged as a design constraint in the Scope and Limitations. In v2.9.6, logging behavior was also separated into the standalone Financial Management Score, giving panels the option to view it either way.'
)
add_defense_box(
    'What\'s the difference between the FHS and the CFPB scale?',
    'The CFPB Financial Well-Being Scale (2017) is a validated 10-question survey that measures perceived financial security and freedom. It\'s designed for population studies and self-assessment. SmartSpend\'s FHS is computed — not surveyed. Every number comes from actual recorded transactions. The two tools are complementary: the CFPB scale measures how the user feels about their finances; SmartSpend\'s FHS measures what their financial data actually shows.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — FINANCIAL MANAGEMENT SCORE (FMS)
# ══════════════════════════════════════════════════════════════════════════════

page_break()
add_heading('SECTION 2 — FINANCIAL MANAGEMENT SCORE (FMS)', 1)

add_body(
    'SmartSpend v2.9.6 introduced the Financial Management Score (FMS) as a companion '
    'metric to the FHS. While the FHS measures financial outcomes (savings rate, overspend '
    'control, budget adherence), the FMS measures financial management behavior — how '
    'consistently and diligently the user manages their financial data in the app.'
)
add_body(
    'Research recommendation basis: The FHN 2026 update emphasizes that financial health '
    '(outcomes) should be measured separately from financial management behavior (practices '
    'that produce those outcomes). Rateweb (2026) and Elenvo AI (2026) both note that '
    'mixed metrics reduce interpretability and actionability for end users.'
)

add_heading('2.1  FMS Formula — 4 Components × 25 pts = 100 Maximum', 2)

add_table([
    ['Component', 'What It Measures', 'Formula', 'Full Score Condition'],
    ['1. Logging Consistency\n(25 pts)',
     'How regularly expenses are\nrecorded this month',
     'comp1 = 25 × min(1.0, loggedDays / activeDays)\n(same formula as FHS Full Mode Component 4)',
     'Logging every active day (daily)'],
    ['2. Budget Setup\n(25 pts)',
     'Whether category budgets or\nspending limits are configured',
     'comp2 = 25 × min(1.0, budgetCount / 5)\n(full at 5 or more budgets/limits)',
     '≥ 5 budgets configured'],
    ['3. Goal Tracking\n(25 pts)',
     'Whether savings goals exist\nand have active contributions',
     'comp3 = 25 pts (2+ active goals with contributions)\n      = 15 pts (1+ active goal, no contributions)\n      = 5 pts (goals exist but no target amounts)',
     '2+ goals with progress contributions'],
    ['4. Data Completeness\n(25 pts)',
     'Income and wallet data entered\n(income mode) or spending\nlimit set (lightweight mode)',
     'Income mode: 12 pts (income set) + 13 pts (1+ wallet balance)\nLightweight mode: 25 pts (limit set), 10 pts (no limit)',
     'Income set + at least 1 wallet balance entered'],
], col_widths=[Inches(1.3), Inches(1.5), Inches(2.4), Inches(1.3)], font_size=9)
add_caption('Table 2.1. Financial Management Score (FMS) — Components and Formulas')

add_heading('2.2  FMS Score Classification', 2)

add_table([
    ['Score Range', 'Label', 'Color (UI)', 'Meaning'],
    ['85–100', '🏅 Expert Tracker',      'Purple', 'Diligent, comprehensive financial tracking habits'],
    ['70–84',  '✅ Active Manager',       'Blue',   'Good habits with minor data gaps'],
    ['50–69',  '📋 Getting Started',      'Teal',   'Some features configured; needs more consistency'],
    ['< 50',   '⚠️ Set Up Your Profile', 'Grey',   'Core features (budgets, goals, income) not yet configured'],
], col_widths=[Inches(1.0), Inches(1.5), Inches(0.9), Inches(3.1)], font_size=10)
add_caption('Table 2.2. FMS Score Classification (4-Tier Scale)')

add_heading('2.3  How FHS and FMS Complement Each Other', 2)

add_table([
    ['', 'FHS (Financial Health Score)', 'FMS (Financial Management Score)'],
    ['Measures', 'Financial outcomes this month', 'Financial tracking behavior this month'],
    ['High score means', 'You are financially healthy', 'You are diligently managing your data'],
    ['Low score means', 'Financial outcomes need improvement', 'You need to track more consistently or configure more features'],
    ['Can be high while other is low?', 'Yes — excellent tracker with poor spending control', 'Yes — great spender who rarely opens the app'],
    ['Actionable tip', 'Address the weakest FHS component', 'Log more consistently, fill in income/wallet, add budgets'],
    ['Key academic distinction', 'Financial health outcome\n(FinHealth Network, 2021)', 'Financial management behavior\n(Rateweb, 2026; Elenvo AI, 2026)'],
], col_widths=[Inches(1.8), Inches(2.35), Inches(2.35)], font_size=10)
add_caption('Table 2.3. FHS vs FMS — How the Two Scores Complement Each Other')

add_defense_box(
    'If the FHS already has Logging Consistency, why is there a separate FMS?',
    'The FHS retains Logging Consistency because, without bank API access, logged data is the only input source for Savings Rate, Overspend Control, and Budget Adherence. However, per the FHN 2026 recommendation, we also surface it separately in the FMS alongside three additional management behaviors (Budget Setup, Goal Tracking, Data Completeness) that aren\'t in the FHS. This gives the user two complementary views: what their finances look like (FHS) and how well they\'re using the tool to track them (FMS).'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — LLM COMPARATIVE BENCHMARKING
# ══════════════════════════════════════════════════════════════════════════════

page_break()
add_heading('SECTION 3 — LLM COMPARATIVE BENCHMARKING', 1)

add_body(
    'The selection of an appropriate Large Language Model API is a critical design decision '
    'for SmartSpend because the LLM directly influences the accuracy and latency of natural '
    'language expense parsing, the reliability of 31 agentic actions, and the quality of '
    'Filipino-English conversational financial advice. This section documents the '
    'comparative evaluation of 19 LLM APIs and the rationale for the selected 5-provider '
    'failover architecture.'
)

# 3.1 Why LLMs in personal finance
add_heading('3.1  Research Basis: Why Use an LLM for a Finance App?', 2)

add_table([
    ['Research Finding', 'Source', 'Implication for SmartSpend'],
    ['LLMs can reduce manual effort in financial data entry and improve categorization accuracy for multilingual inputs',
     'Li et al. (2024); Hean et al. (2025)',
     'Justifies multi-modal AI input (voice, OCR, natural language) replacing manual form entry'],
    ['Conversational AI enables financial guidance at scale without the cost of human advisers',
     'Hean et al. (2025)',
     'Justifies AI chat as the primary interaction paradigm for expense logging and financial advice'],
    ['49% of global consumers used AI for savings/investment decisions in the 6 months prior to April 2026',
     'EY Global AI Survey (2026a)',
     'Establishes AI-assisted finance as mainstream consumer behavior, not experimental'],
    ['18% of global consumers used AI specifically for budgeting and household finance management',
     'EY Global AI Survey (2026a)',
     'Direct validation of SmartSpend\'s primary use case'],
    ['Tool harness architecture affects LLM performance on financial tasks more than base model choice',
     'FrontierFinance Benchmark (Arcila et al., 2026)',
     'Validates SmartSpend\'s context-injection agentic architecture over naive API calls'],
    ['Finance-specialized LLMs (FinGPT, Fin-R1) are stronger on financial NLP tasks (classification, sentiment) but weaker than frontier general models on reasoning and multilingual use',
     'Li et al. (2024); Liu et al. (2025); FrontierFinance (2026)',
     'Confirms that Gemini 3.1 Flash-Lite (general multilingual) outperforms finance-specific models for SmartSpend\'s mixed Filipino-English workload'],
], col_widths=[Inches(2.3), Inches(1.5), Inches(2.7)], font_size=9)
add_caption('Table 3.1. Research Basis for LLM Use in SmartSpend')

# 3.2 Selection criteria
add_heading('3.2  LLM Selection Criteria and Weights', 2)
add_body(
    'SmartSpend\'s selection criteria were weighted to reflect the constraints of an '
    'academic mobile deployment: zero budget, Filipino-English users, real-time mobile '
    'response times, and the need for reliable structured JSON output for agentic actions.'
)

add_table([
    ['Criterion', 'Weight', 'Rationale'],
    ['Filipino-English Accuracy', '25%',
     'Primary language need: Taglish expense parsing (e.g., "nagbayad ako ng 150 sa Jollibee"). Google\'s training data coverage of Southeast Asian languages gives Gemini models an advantage.'],
    ['Speed / Latency', '20%',
     'Mobile UX requires < 3 seconds response time for real-time logging. Groq LPU (~315–800 t/s) and Cerebras WSE (~1,800 t/s) lead here; GPT-5.6 (~80–120 t/s) is too slow for snappy mobile interactions.'],
    ['Tool Use / JSON Reliability', '20%',
     '31 agentic actions require the LLM to return valid, parseable JSON every time. Models with native function calling (Gemini, GPT-5.6, Claude) lead; smaller open-source models have higher JSON error rates.'],
    ['Free Tier Availability', '15%',
     'Academic deployment serving 30 respondents at 60 messages/day requires ~1,800 daily requests. Gemini Flash-Lite (1,000/day) + Groq (14,400/day) + Cerebras (1M tokens/day) covers this at zero cost.'],
    ['Context Window', '10%',
     'SmartSpend injects ~2,000–5,000 tokens of financial context per message. All evaluated models handle this comfortably; a 1M+ window is overkill but never a bottleneck.'],
    ['Financial Reasoning Quality', '10%',
     'Advisory queries (SSS contributions, debt payoff strategy, savings timeline) require accurate financial reasoning. Gemini Pro and GPT-5.6 lead; Flash-Lite is adequate for most queries.'],
], col_widths=[Inches(1.6), Inches(0.6), Inches(4.3)], font_size=9.5)
add_caption('Table 3.2. LLM Selection Criteria and Weights for SmartSpend')

# 3.3 Full benchmarking table
page_break()
add_heading('3.3  Full LLM Benchmarking Table — All 19 Models Evaluated', 2)

add_table([
    ['Model', 'Provider', 'Context', 'Speed\n(t/s)', 'Filipino\nAccuracy', 'Tool Use\n/ JSON', 'Free Tier', 'Selected?'],
    ['Gemini 3.1 Flash-Lite', 'Google',   '1,000,000', '~400–600',   '★★★★★', '★★★★★', '✅ 1,000/day', '✅ PRIMARY'],
    ['Gemini 3.5 Flash',      'Google',   '1,000,000', '~200–400',   '★★★★★', '★★★★★', '✅ 250/day',   '✅ Fallback 1'],
    ['LLaMA 3.3 70B',         'Groq LPU', '128,000',  '~315',       '★★★★☆', '★★★★★', '✅ 14,400/day','✅ Fallback 2'],
    ['LLaMA 3.1 8B',          'Groq LPU', '8,192',    '~800',       '★★★★☆', '★★★★☆', '✅ 14,400/day','✅ Fallback 3'],
    ['LLaMA 3.1 70B',         'Cerebras', '128,000',  '~1,800',     '★★★★☆', '★★★★☆', '✅ 1M tokens', '✅ Fallback 4'],
    ['Qwen 3 32B',            'Alibaba/OpenRouter','128,000','~150–300','★★★★☆','★★★★★','✅ Free preview','❌ Less tested'],
    ['GPT-5.6 Luna',          'OpenAI',  '1,050,000', '~120',       '★★★★☆', '★★★★★', '❌ $0.20/$1.20/1M','❌ Cost'],
    ['GPT-5.6 Terra',         'OpenAI',  '1,050,000', '~80–120',    '★★★★★', '★★★★★', '❌ $2/$12/1M', '❌ Cost'],
    ['GPT-5.6 Sol',           'OpenAI',  '1,050,000', '~80–120',    '★★★★★', '★★★★★', '❌ $5/$30/1M', '❌ Cost'],
    ['Claude Fable 5',        'Anthropic','200,000',  '~70–100',    '★★★★★', '★★★★★', '❌ Paid only', '❌ Cost'],
    ['Gemini 3.7 Flash',      'Google',  '1,048,576', '~300–500',   '★★★★★', '★★★★★', '❌ $0.75/1M',  '❌ No free'],
    ['Gemini 3.1 Pro',        'Google',  '2,000,000', '~100–150',   '★★★★★', '★★★★★', '⚠️ 100 RPD',  '❌ Low quota'],
    ['Grok 4.6',              'xAI',     '500,000',   '~100–200',   '★★★★☆', '★★★★★', '❌ $2/1M',     '❌ Cost'],
    ['DeepSeek V4',           'DeepSeek','1,000,000', '~200',       '★★★☆☆', '★★★☆☆', '⚠️ 5M trial', '❌ Weak Filipino'],
    ['Fin-R1 (7B)',           'Self-host','128,000',  'Varies',     '★★★☆☆', '★★☆☆☆', '✅ Self-host', '❌ No API; no Filipino'],
    ['FinGPT',                'Self-host','Varies',   'Varies',     '★★☆☆☆', '★★★☆☆', '✅ Self-host', '❌ Market-focused; no Filipino'],
    ['Mistral 7B',            'Mistral', '32,000',    '~600',       '★★★☆☆', '★★★☆☆', '✅ Self-host', '❌ Poor Filipino'],
    ['GPT-4o Mini',           'OpenAI',  '128,000',   '~120',       '★★★★☆', '★★★★★', '❌ No free',   '❌ Cost'],
    ['Gemma 2 9B',            'Google',  '8,192',     '~500',       '★★★☆☆', '★★★☆☆', '✅ Local only','❌ No hosted API'],
], col_widths=[Inches(1.4), Inches(1.0), Inches(0.7), Inches(0.6), Inches(0.65), Inches(0.65), Inches(0.9), Inches(0.85)],
   font_size=8, header_shade='5C0E24')
add_caption('Table 3.3. Comparative Evaluation of LLM APIs for SmartSpend — All 19 Models (August 2026)\nSources: FrontierFinance (arXiv:2608.11683); AIMUltiple Finance LLM Benchmark 2026; micro1.ai REALM; Google AI pricing; Groq free tier documentation')

# 3.4 Task-specific performance
page_break()
add_heading('3.4  LLM Performance by Task Type', 2)
add_body(
    'SmartSpend uses LLMs for four distinct task types. Performance differs significantly '
    'across models for each task.'
)

add_heading('Task A — Filipino-English Expense Parsing', 3)
add_body('Example: "Nagbayad ako ng 150 pesos sa Jollibee kaninang lunch"', italic=True)
add_table([
    ['Model', 'Accuracy', 'Speed', 'Notes'],
    ['Gemini 3.1 Flash-Lite', 'Very Good', 'Very Fast (~400–600 t/s)', 'Best free-tier option; correctly handles Tagalog/Taglish merchant names and Filipino food items'],
    ['Gemini 3.5 Flash', 'Excellent', 'Fast (~200–400 t/s)', 'Better quality than Flash-Lite; 4× fewer free daily requests'],
    ['GPT-5.6', 'Excellent', 'Moderate (~80–120 t/s)', 'Best overall NLP but prohibitively costly for academic deployment'],
    ['Claude Fable 5', 'Excellent', 'Moderate (~70–100 t/s)', 'Strong instruction following; paid only'],
    ['Groq LLaMA 3.3 70B', 'Good', 'Fast (~315 t/s)', 'Handles Filipino reasonably; best free open-source option'],
    ['Groq LLaMA 3.1 8B', 'Adequate', 'Very Fast (~800 t/s)', 'Occasionally misclassifies Filipino food items; used for simple queries'],
    ['Mistral 7B / DeepSeek', 'Weak–Adequate', 'Varies', 'Poor Tagalog/Taglish handling; not suitable for SmartSpend primary role'],
    ['Fin-R1 / FinGPT', 'Weak', 'Varies', 'Trained on financial market data; poor conversational Filipino-English'],
], col_widths=[Inches(1.5), Inches(0.8), Inches(1.5), Inches(2.7)], font_size=9)

add_heading('Task B — Agentic Action Execution (JSON Reliability)', 3)
add_body('The LLM must return valid JSON matching one of 31 action type schemas every time.', italic=True)
add_table([
    ['Model', 'JSON Reliability', 'Instruction Following', 'Notes'],
    ['GPT-5.6', 'Near-perfect', 'Excellent', 'Best JSON reliability overall; paid only'],
    ['Claude Fable 5 / Sonnet 5', 'Near-perfect', 'Excellent', 'Highest SWE-bench scores; paid only'],
    ['Gemini 3.5 Flash', 'Excellent', 'Excellent', 'Best free-tier option for action execution'],
    ['Gemini 3.1 Flash-Lite', 'Very Good', 'Good', 'Occasionally misidentifies action type on ambiguous inputs; fallback parser catches ~99% of errors'],
    ['Groq LLaMA 3.3 70B', 'Good', 'Good', 'Reliable for well-defined prompts; used for fallback 2'],
    ['Groq LLaMA 3.1 8B', 'Adequate', 'Adequate', 'More occasional JSON format errors; mitigated by fallback parser'],
    ['Cerebras LLaMA 3.1', 'Adequate', 'Adequate', 'Speed (1,800+ t/s) is primary advantage; quality similar to LLaMA 3.1 8B'],
], col_widths=[Inches(1.5), Inches(1.0), Inches(1.2), Inches(2.8)], font_size=9)

add_heading('Task C — Philippine Financial Advisory', 3)
add_body('Example: "How do I apply for SSS Flexi Fund? How much PhilHealth should I pay at ₱30K salary?"', italic=True)
add_table([
    ['Model', 'PH Financial Knowledge', 'Notes'],
    ['Gemini 3.1 Pro', 'Very Good', 'Best PH financial knowledge — Google\'s training includes PH government sources; limited to 100 RPD free'],
    ['Gemini 3.5 Flash', 'Good–Very Good', 'Strong PH knowledge; 250 RPD free; used for financial advice routing tier'],
    ['Gemini 3.1 Flash-Lite', 'Good', 'Adequate for most advisory queries; correct on SSS/PhilHealth contribution computation'],
    ['GPT-5.6 Terra', 'Good', 'Strong general financial knowledge; PH-specific data may lag training cutoff'],
    ['Claude Fable 5', 'Good', 'Best for long-form financial reasoning and planning narratives'],
    ['Groq LLaMA 3.3 70B', 'Adequate', 'General financial advice; weaker on PH-specific government services'],
    ['DeepSeek V4', 'Weak–Adequate', 'Strong math reasoning; limited PH government/financial context'],
], col_widths=[Inches(1.5), Inches(1.2), Inches(3.8)], font_size=9)

# 3.5 Multi-provider architecture
page_break()
add_heading('3.5  SmartSpend 5-Provider Failover Architecture', 2)
add_body(
    'Rather than committing to a single provider, SmartSpend routes queries across five '
    'providers based on query complexity and available quota. This architecture provides '
    'zero cost for academic deployment, near-zero downtime, and task-matched quality.'
)

add_table([
    ['Priority', 'Provider', 'Model', 'Free Limit', 'Role'],
    ['1 (Primary)', 'Google AI',  'Gemini 3.1 Flash-Lite', '1,000 req/day', 'Best Filipino-English accuracy; 1M token context; used for standard queries'],
    ['2', 'Google AI',  'Gemini 3.5 Flash',     '250 req/day',    'Higher reasoning quality; financial advice and complex analysis'],
    ['3', 'Groq LPU',  'LLaMA 3.3 70B',        '~14,400 req/day','Best open-source reasoning; ~315 t/s; used for complex queries'],
    ['4', 'Groq LPU',  'LLaMA 3.1 8B',         '~14,400 req/day','Fastest simple queries; ~800 t/s; expense parsing fallback'],
    ['5', 'Cerebras WSE','LLaMA 3.1 70B',       '1M tokens/day',  'Highest raw throughput (~1,800 t/s); last-resort fallback'],
], col_widths=[Inches(0.8), Inches(0.9), Inches(1.5), Inches(1.1), Inches(2.2)], font_size=10)
add_caption('Table 3.4. SmartSpend 5-Provider LLM Failover Architecture (v2.9.10)')

add_body('Why context injection (not RAG):', before=8)
add_body(
    'A typical SmartSpend user has 20–50 expenses, 5–10 budgets, and 3–5 goals '
    '(~1,000–5,000 tokens). This fits entirely within any evaluated model\'s context '
    'window. Retrieval-Augmented Generation (RAG) uses a vector database and embedding '
    'similarity search — designed for large knowledge bases (thousands of documents). '
    'RAG adds unnecessary infrastructure overhead for SmartSpend\'s small per-user '
    'dataset (Davenport & Mittal, 2022). Full-context injection is simpler, more '
    'accurate, and works offline.'
)

add_heading('3.6  General LLMs vs Finance-Specialized LLMs — Comparative Study', 2)
add_body(
    'A key design question in SmartSpend\'s development was: should a finance-specialized '
    'LLM (trained specifically on financial data) be used instead of a general-purpose '
    'multilingual LLM? This section documents the structured comparative analysis and '
    'the research basis for the final selection.'
)

add_heading('Background: Categories of Financial AI Models', 3)
add_table([
    ['Category', 'Models', 'Training Data', 'Strengths', 'Weaknesses'],
    ['Financial NLP Models\n(not conversational)',
     'FinBERT, FLANG',
     'Financial news, SEC filings, earnings call transcripts',
     'Classification, sentiment analysis, named entity recognition in financial text',
     'Not conversational; cannot parse natural language expense input; no function calling; no Filipino'],
    ['Finance Instruction-Tuned LLMs',
     'FinGPT (Liu et al., 2023)\nFinMA / PIXIU\nInvestLM\nFinTral',
     'General LLM fine-tuned on 34+ financial data sources using LoRA/QLoRA; Bloomberg, Reuters, Yahoo Finance, SEC EDGAR',
     'Strong on financial classification, sentiment, and NER tasks; outperforms general models on FinBench classification sub-tasks',
     'Significantly weaker than frontier general models on reasoning, QA, and summarization (arXiv:2507.08015). Weak on multilingual (Filipino-English). No hosted free API. Self-hosting requires GPU.'],
    ['Financial Reasoning LLMs',
     'Fin-R1 (Liu et al., 2025)\n7B parameters, RL-trained',
     'General LLM with reinforcement learning on financial reasoning datasets (FinQA, ConvFinQA)',
     'SOTA on FinQA (avg 75.2) and ConvFinQA benchmarks; strong quantitative financial reasoning',
     'Self-host only; no API service; poor Filipino-English; trained on formal financial Q&A (not conversational expense logging); no function calling'],
    ['Domain-Specific Closed LLMs',
     'BloombergGPT (Wu et al., 2023)\n50B parameters',
     'Bloomberg\'s proprietary financial data corpus + general text',
     'Established that domain-specific training improves financial NLP tasks vs same-size general LLMs',
     'Closed-source; no API access; trained on Bloomberg market data (institutional, not consumer finance); no Filipino-English'],
    ['General-Purpose Frontier LLMs',
     'Gemini 3.1 Flash-Lite, GPT-5.6, Claude Fable 5, Groq LLaMA 3.3 70B',
     'Broad multilingual training including Southeast Asian languages, consumer apps, government documents, conversational text',
     'Best multilingual accuracy (Filipino-English); best instruction following and JSON reliability; fastest inference on free tiers; strong conversational financial reasoning',
     'Not specifically trained on financial markets; may have lower accuracy on complex institutional financial analysis (equity valuation, options pricing)'],
], col_widths=[Inches(1.3), Inches(1.3), Inches(1.4), Inches(1.5), Inches(1.5)], font_size=8.5)
add_caption('Table 3.5. Categories of Financial AI Models — Architecture and Capability Overview')

add_heading('Head-to-Head: General LLMs vs Finance-Specialized LLMs on SmartSpend\'s 5 Task Types', 3)
add_body(
    'The following table directly compares how each model category performs on '
    'SmartSpend\'s actual use cases — not on institutional financial benchmarks.'
)

add_table([
    ['SmartSpend Task', 'Finance-Specialized LLMs\n(FinGPT, Fin-R1, BloombergGPT)', 'General Multilingual LLMs\n(Gemini 3.1 Flash-Lite, LLaMA 3.3 70B)', 'Winner\nfor SmartSpend'],
    ['Filipino-English expense parsing\ne.g. "Nagbayad ako ng 150 sa Jollibee kaninang lunch"',
     '❌ WEAK\nNot trained on Filipino/Tagalog text. No coverage of local merchant names (Jollibee, SM Hypermarket, Palengke). Self-hosted models have no multilingual fine-tuning.',
     '✅ STRONG\nGoogle\'s training data covers Southeast Asian languages. Gemini correctly identifies Filipino food terms, merchant types, and Taglish constructions.',
     'General LLM\n(Gemini 3.1 Flash-Lite)'],
    ['Agentic action execution\n(parse intent → valid JSON for 1 of 31 action types)',
     '❌ WEAK\nFinGPT and Fin-R1 are not instruction-tuned for function calling or structured JSON output. BloombergGPT has no function calling capability.',
     '✅ STRONG\nGemini and GPT-5.6 have native function calling. LLaMA 3.3 70B reliably follows JSON schema instructions. SmartSpend\'s fallback parser brings success rate to ~99%.',
     'General LLM\n(Gemini / Groq)'],
    ['Philippine-specific financial advisory\ne.g. "How much SSS should I pay at ₱30K salary?"',
     '❌ WEAK\nFinGPT and BloombergGPT are trained on US/international financial markets (SEC filings, Bloomberg data). No Philippine government financial data (SSS, PhilHealth, Pag-IBIG, BIR TRAIN Law).',
     '✅ STRONG\nGemini models include Philippine government sources in training. Correctly computes SSS contributions, PhilHealth rates, BIR TRAIN Law tax brackets.',
     'General LLM\n(Gemini 3.5 Flash)'],
    ['Consumer personal finance reasoning\ne.g. "I spent ₱7,000 on food but my budget is ₱3,000. What should I do?"',
     '⚠️ MODERATE\nFin-R1 achieves SOTA on formal FinQA datasets. But FinQA tests structured financial Q&A, not conversational personal budgeting advice.',
     '✅ STRONG\nGeneral frontier LLMs excel at conversational advisory. Claude Fable 5 scored highest on the Hebbia Finance Benchmark for reading and synthesizing financial context (90.34% accuracy).',
     'General LLM\n(paid tier);\nGemini Flash-Lite\nfor free tier'],
    ['Financial calculations\n(FHS score, savings rate, budget utilization %)',
     'N/A — Financial calculations\nshould NEVER be done by any LLM\n(see note below)',
     'N/A — Financial calculations\nshould NEVER be done by any LLM\n(see note below)',
     '🔴 Neither —\nAlways application code\n(score_service.dart)'],
], col_widths=[Inches(1.5), Inches(2.2), Inches(2.2), Inches(0.8)], font_size=8.5)
add_caption(
    'Table 3.6. General LLMs vs Finance-Specialized LLMs on SmartSpend\'s 5 Task Types\n'
    'Sources: FrontierFinance Benchmark (Arcila et al., 2026; arXiv:2608.11683); '
    'Li et al. (2024); Liu et al. (2023, 2025); Wu et al. (2023); '
    'FinRCA-Bench (Xiao et al., 2026; arXiv:2608.18534); '
    'arXiv:2507.08015 (FinGPT evaluation); Hebbia Finance Benchmark (Claude Fable 5 result)'
)

add_body(
    '🔴 Important note on financial calculations: The FinDeepIndicator study (2026) and '
    'InvestLogicBench (arXiv:2608.06108, 2026) both found that LLMs — including frontier '
    'models — exhibit significant numerical degradation on financial arithmetic. '
    'SmartSpend\'s hybrid architecture addresses this directly: the LLM never computes '
    'the FHS, savings rate, or budget utilization. These are always computed '
    'deterministically by score_service.dart, then injected into the AI context as '
    'pre-calculated values. The LLM\'s role is exclusively to explain and advise — '
    'not to calculate.'
)

add_heading('Why Finance-Specialized LLMs Are Not the Right Choice for SmartSpend', 3)
add_table([
    ['Criterion', 'Finance-Specialized LLMs\n(FinGPT, Fin-R1, BloombergGPT)', 'General Multilingual LLMs\n(Gemini 3.1 Flash-Lite + Groq)'],
    ['Filipino-English support',
     '❌ None — trained on English financial corpora only',
     '✅ Native — Gemini trained on SEA languages including Tagalog/Filipino'],
    ['Philippine financial knowledge\n(SSS, PhilHealth, BIR TRAIN Law)',
     '❌ None — trained on US/international institutional data',
     '✅ Good — Gemini includes PH government sources; adequate for capstone scope'],
    ['Function calling / JSON output',
     '❌ None (FinGPT, Fin-R1, BloombergGPT do not support function calling)',
     '✅ Native function calling in Gemini and GPT-5.6; JSON reliable in LLaMA 3.3 70B'],
    ['Free API access for academic use',
     '✅ Self-hosted (requires GPU server)\n❌ No hosted free API for any finance LLM',
     '✅ Gemini Flash-Lite: 1,000/day free\n✅ Groq LLaMA 3.3 70B: ~14,400/day free'],
    ['Conversational personal finance\n(budgeting, goals, Filipino lifestyle)',
     '⚠️ Moderate on structured FinQA\n❌ Poor on conversational consumer queries',
     '✅ Excellent on conversational advisory — Claude Fable 5, Gemini Pro lead'],
    ['Financial market analysis\n(stock valuation, institutional finance)',
     '✅ Strong — specifically trained for this',
     '⚠️ Adequate for general queries; weaker on specialized institutional tasks'],
    ['Suitable for SmartSpend?',
     '❌ Not suitable — wrong domain (market data vs consumer finance), no Filipino support, no hosted API',
     '✅ Best choice — multilingual, free tier, reliable JSON, conversational strength'],
], col_widths=[Inches(1.7), Inches(2.4), Inches(2.4)], font_size=9)
add_caption('Table 3.7. Finance-Specialized vs General LLMs — SmartSpend Suitability Comparison')

add_body(
    'Core research conclusion (Part 11 of RESEARCH_BASIS.md): A financial application '
    'does not require a finance-specialized LLM as its primary language model. '
    'Finance-specialized models are optimized for institutional financial NLP tasks '
    '(market analysis, sentiment classification, earnings QA) — not for consumer personal '
    'finance in Filipino-English. For SmartSpend\'s workload, modern general-purpose '
    'multilingual frontier LLMs are both more capable and more accessible. '
    'Source: FrontierFinance Benchmark (Arcila et al., 2026); Li et al. (2024); '
    'FinRCA-Bench (Xiao et al., 2026).'
)

add_defense_box(
    'Have you considered using Fin-R1 or FinGPT since they\'re specifically for finance?',
    'Yes — both were evaluated. Fin-R1 (Liu et al., 2025) achieves SOTA on FinQA and ConvFinQA benchmarks (avg 75.2), which test formal structured financial Q&A. However, it (1) has no hosted API — self-hosting requires a GPU server; (2) has no Filipino-English training; (3) has no function calling support for SmartSpend\'s 31 agentic actions; and (4) is optimized for institutional financial reasoning, not conversational consumer budgeting. FinGPT is strong on financial classification and sentiment, but the arXiv:2507.08015 evaluation found it significantly weaker than frontier general models on reasoning, QA, and summarization. For SmartSpend\'s Filipino-English expense parsing and agentic action execution, Gemini 3.1 Flash-Lite outperforms both at zero cost with a hosted API.'
)
add_defense_box(
    'Does using a general LLM mean the AI gives less accurate financial advice?',
    'No — and this is an important distinction. "Financial accuracy" means different things for different tasks. For institutional tasks (stock valuation, options pricing, SEC filing analysis), finance-specialized LLMs have advantages. For consumer personal finance advisory (budgeting, savings plans, SSS contributions, debt payoff strategy) in Filipino-English, general frontier LLMs are superior because they have broader multilingual training and stronger conversational reasoning. The FrontierFinance Benchmark (2026) confirmed that tool harness architecture — how the LLM receives financial context — affects performance more than whether the model is finance-specialized. SmartSpend\'s full-context injection architecture (injecting all user financial data before each query) is the key design choice, not the specific model used.'
)

add_defense_box(
    'Why not use a hybrid — finance-specialized for calculations, general LLM for conversation?',
    'SmartSpend already implements this hybrid — but the "specialized" layer is application code (score_service.dart), not a finance LLM. Financial calculations (FHS score, savings rate, budget utilization) are always computed deterministically by the Dart financial engine, then handed to the general LLM as pre-calculated context. This is more reliable than any LLM for numerical computation — FinDeepIndicator (2026) and InvestLogicBench (2026) both found that even frontier LLMs exhibit significant numerical degradation on financial arithmetic. Application code is always more accurate for math than any LLM.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — APP FEATURE COMPARISON
# ══════════════════════════════════════════════════════════════════════════════

page_break()
add_heading('SECTION 4 — APP FEATURE COMPARISON', 1)

add_body(
    'A comprehensive review of 22 existing personal finance applications (14 international, '
    '8 Philippine-context) was conducted in August 2026 to identify SmartSpend\'s '
    'competitive positioning, validate feature necessity, and document genuine capability '
    'gaps that inform the research recommendations.'
)

# 4.1 International comparison
add_heading('4.1  SmartSpend vs International Apps — Core Feature Matrix', 2)

add_table([
    ['Feature', 'SmartSpend', 'YNAB', 'Monarch', 'Copilot', 'Rocket\nMoney', 'Simplifi', 'PocketGuard', 'Goodbudget', 'Spendee'],
    # Platform
    ['Android',          '✅','✅','✅','❌ iOS','✅','✅','✅','✅','✅'],
    ['Free tier',        '✅ Full','❌','❌','❌','✅ Ltd','❌','✅ Ltd','✅ Ltd','✅ Ltd'],
    ['Paid price',       'Free','$14.99/mo','$9.99/mo','$10.99/mo','$6–12/mo','$3.99/mo','$12.99/mo','$8/mo','$2.99/mo'],
    # Input
    ['Voice input',      '✅','❌','❌','❌','❌','❌','❌','❌','❌'],
    ['OCR receipt scan', '✅','❌','❌','❌','❌','❌','❌','❌','❌'],
    ['Batch screenshots','✅ 40+','❌','❌','❌','❌','❌','❌','❌','❌'],
    ['Bank / e-wallet sync','❌ manual','✅','✅','✅','✅','✅','✅','❌','✅ Paid'],
    # AI
    ['AI chat assistant','✅ 31 actions','❌','❌','✅ Basic','❌','❌','❌','❌','❌'],
    ['Financial Health Score','✅ 0–100','❌','❌','❌','❌','❌','❌','❌','❌'],
    ['Behavioral analysis','✅','❌','❌','❌','❌','❌','❌','❌','❌'],
    # Features
    ['Multi-period limits','✅ D/W/M/Y','❌','❌','❌','❌','❌','✅ Daily','❌','❌'],
    ['50/30/20 tracker', '✅','❌','❌','❌','❌','❌','✅','❌','❌'],
    ['Debt tracker',     '✅','✅','✅','❌','✅','✅','❌','❌','❌'],
    # Connectivity
    ['Offline mode',     '✅ Full','❌','❌','❌','❌','❌','❌','✅','❌'],
    ['Cloud sync',       '✅ Firebase','✅','✅','✅','✅','✅','✅','✅','✅'],
    ['PIN + biometric',  '✅','❌','❌','❌','❌','❌','❌','❌','❌'],
    # Gamification
    ['Achievement badges','✅ 23','❌','❌','❌','❌','❌','❌','❌','❌'],
    ['Daily quests',     '✅ 10','❌','❌','❌','❌','❌','❌','❌','❌'],
],
col_widths=[Inches(1.3)]+[Inches(0.65)]*8, font_size=8.5)
add_caption('Table 4.1. SmartSpend vs International Personal Finance Apps — Core Feature Comparison\nSources: PCMag Best Personal Finance Apps 2026; NerdWallet Best Budget Apps 2026; CNBC Best Budgeting Apps 2026')

# 4.2 Philippine competitors
page_break()
add_heading('4.2  SmartSpend vs Philippine-Context Apps', 2)

add_table([
    ['Feature', 'SmartSpend', 'BudgetPH\n(KindlyF)', 'Tarsi', 'Alkansya AI', 'Sentimo', 'SweldoWise', 'Lista PH', 'GCash\nPera Coach'],
    ['Platform',        'Android','PWA + Android','iOS + Android','iOS only','Android','Android','Android','Android (GCash)'],
    ['Cost',            'Free','Free','Free','~$2–3/mo','Free','Free','Free','Free (GCash verified)'],
    ['AI chat / actions','✅ 31 actions','✅ Insights only','❌','✅ Chat','❌','❌','❌','✅ Q&A + literacy'],
    ['Natural language input','✅','❌','❌','✅','❌','❌','❌','✅ Filipino/Taglish'],
    ['Voice input',     '✅','❌','❌','❌','❌','❌','❌','❌'],
    ['OCR / Barcode',   '✅','❌','✅ OCR','❌','❌','❌','❌','❌'],
    ['Batch screenshots','✅ 40+','❌','❌','❌','❌','❌','❌','❌'],
    ['Offline mode',    '✅ Full','✅','✅','❌','❌','❌','✅','❌'],
    ['Cloud sync',      '✅ Firebase','✅','❌','✅','❌','❌','❌','✅ GCash'],
    ['Financial Health Score','✅ 0–100 dual-mode','✅ Budget score (simpler)','❌','✅ Web-only','❌','❌','❌','❌'],
    ['Gamification',    '✅ 23 badges\n10 quests, streaks','✅ XP/levels\nstreaks','❌','❌','✅ Streaks\nKBoy mascot','❌','❌','❌'],
    ['50/30/20 tracker','✅','✅','❌','❌','❌','❌','❌','❌'],
    ['Paluwagan tracker','❌','✅ Full','❌','❌','❌','❌','❌','❌'],
    ['15th/30th cycle', '❌','✅','❌','❌','✅ Partial','✅','❌','❌'],
    ['GCash/Maya track','✅ Wallet balance','✅ CSV import','❌','❌','❌','❌','❌','✅ GCash balance'],
    ['SSS/PhilHealth',  '✅ AI compute + tracker','✅ Records','❌','❌','❌','❌','❌','❌'],
    ['Insurance tracker','✅','❌','❌','❌','❌','❌','❌','❌'],
    ['BIR TRAIN Law',   '✅ Calculator','❌','❌','❌','❌','❌','❌','❌'],
    ['Expense tracking','✅','✅','✅','✅','✅','✅','✅','❌ Advisory only'],
],
col_widths=[Inches(1.2)]+[Inches(0.79)]*8, font_size=8.3)
add_caption('Table 4.2. SmartSpend vs Philippine-Context Personal Finance Apps (August 2026)\nSources: BudgetPH (budget.kindlyf.com); Alkansya (alkansya.online); Sentimo (sentimoapp.com); SweldoWise (sweldowise.ph); GCash/Mynt Press Release (March 20, 2026)')

# 4.3 Where SmartSpend leads / gaps
add_heading('4.3  Where SmartSpend Leads — Unique Advantages', 2)

add_table([
    ['Advantage', 'Scope', 'Detail'],
    ['31 agentic AI actions',
     'Global — no equivalent found',
     'Every competitor has no AI or passive AI insights only. SmartSpend is the only personal finance app where the AI takes real autonomous actions: logging expenses, setting budgets, updating goals, computing contributions, planning salary splits, and 26 more.'],
    ['Batch screenshot import (40+ platforms)',
     'Global — no equivalent found',
     'Unique feature globally. Users can import 10 screenshots at once from Shopee, Steam, GCash, GrabFood, Netflix, BPI, and 34+ other platforms. Each is OCR\'d and AI-parsed with a platform-specific extraction prompt.'],
    ['Dual-mode FHS (Full + Lightweight)',
     'Only app with dual-mode',
     'Wingman Money (Australia) has a single-mode FHS. BudgetPH has a simpler budget score. SmartSpend\'s dual mode (Full for income trackers, Lightweight for students/freelancers) is the only published dual-mode behavioral FHS.'],
    ['Free + offline + cloud sync on Android',
     'Only free app with all three',
     'Tarsi and Goodbudget are offline but have no sync. All cloud apps require internet. SmartSpend combines full offline functionality (SQLite) with optional Firebase sync — entirely free.'],
    ['Always-free Android app',
     'Among Filipino AI finance apps',
     'BudgetPH is free but web-first. Alkansya AI requires a subscription (iOS only). SmartSpend is always free, always Android-first.'],
    ['Filipino-English agentic AI',
     'Only free Filipino AI finance app\nwith autonomous actions',
     'GCash Pera Coach is Filipino-language but advisory/literacy only (no expense tracking, no FHS, no agentic actions). Alkansya AI has Filipino chat but is iOS-only and subscription-based.'],
], col_widths=[Inches(1.6), Inches(1.4), Inches(3.5)], font_size=9)
add_caption('Table 4.3. SmartSpend Unique Competitive Advantages (August 2026)')

# 4.4 Gaps
add_heading('4.4  Acknowledged Gaps and Research Recommendations', 2)

add_table([
    ['Gap', 'Who Has It', 'Priority', 'Recommendation in Chapter IV'],
    ['Paluwagan tracker',
     'BudgetPH (full feature)',
     '🔴 High',
     'Highest-priority post-capstone feature. The existing debt and recurring transaction infrastructure provides a suitable architectural base. Rotating savings group tracking is uniquely Filipino and absent from all international apps.'],
    ['15th & 30th payday cycle awareness',
     'BudgetPH, SweldoWise, SweldoTrack, Sentimo',
     '🔴 High',
     'Payday-cycle-aware budgeting resets aligned with the Philippine standard of semi-monthly salary payments. Would significantly strengthen Filipino-first positioning.'],
    ['OFW remittance tracking',
     'BudgetPH (partial)',
     '🟡 Medium',
     'Add inbound international remittance tracking as a distinct income category. OFW remittances totalled ~$38.6B in 2024 — a large PH demographic.'],
    ['Couple / family wallet sharing',
     'Monarch (Goals 3.0), Wally, Goodbudget, Honeydue',
     '🟡 Medium',
     'Allow two users to view a shared household wallet. Important for the primary target population (parents managing household finances with spouse).'],
    ['iOS / web / PWA version',
     'All international competitors',
     '🟢 Low (post-capstone)',
     'Flutter supports web and iOS builds. Recommended for post-capstone to expand reach beyond Android users.'],
    ['Backend API proxy',
     'Production apps (e.g., Cleo, Monarch)',
     '🔴 High (security)',
     'Move LLM API key management to a server-side proxy (Firebase Cloud Function or Vercel) to eliminate device-side key exposure. Currently fetched from Firebase Remote Config.'],
    ['Bank sync (automated)',
     'YNAB, Monarch, Copilot, Rocket Money',
     '🟢 Low (not feasible)',
     'Not feasible in the Philippine context — BSP Open Finance (OFxPERA) only launched pilot in July 2025. Architecture is ready for integration when the framework matures.'],
], col_widths=[Inches(1.4), Inches(1.2), Inches(0.7), Inches(3.2)], font_size=9)
add_caption('Table 4.4. SmartSpend Acknowledged Gaps and Research Recommendations')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

page_break()
add_heading('SECTION 5 — FULL APA REFERENCE LIST', 1)
add_body('All references cited in this Technical Reference Document, in APA 7th edition format:')
doc.add_paragraph()

refs = [
    'Arcila, A., et al. (2026). FrontierFinance: A challenging benchmark for measuring frontier intelligence of finance agents. arXiv:2608.11683. https://arxiv.org/abs/2608.11683',
    'Ariely, D. (2008). Predictably irrational: The hidden forces that shape our decisions. HarperCollins.',
    'Bangko Sentral ng Pilipinas. (2021). 2021 Financial Inclusion Survey. BSP. https://www.bsp.gov.ph/Inclusive-Finance/Financial-Inclusion-Surveys/2021-FIS-Report.pdf',
    'Bangko Sentral ng Pilipinas. (2025). Consumer Finance and Inclusion Survey (CFIS) 2025. BSP.',
    'Bangor, A., Kortum, P., & Miller, J. (2009). Determining what individual SUS scores mean: Adding an adjective rating scale. Journal of Usability Studies, 4(3), 114–123.',
    'Bitrián, P., Buil, I., & Catalán, S. (2021). Making finance fun: The gamification of personal financial management apps. International Journal of Bank Marketing, 39(7), 1310–1332. https://doi.org/10.1108/IJBM-09-2020-0491',
    'Brooke, J. (1996). SUS: A "quick and dirty" usability scale. In P. W. Jordan, B. Thomas, B. A. Weerdmeester, & I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189–194). Taylor & Francis.',
    'Cambridge Judge Business School. (2025). From automation to autonomy: The agentic AI era of financial services. https://www.jbs.cam.ac.uk/2025/from-automation-to-autonomy-the-agentic-ai-era-of-financial-services/',
    'Consumer Financial Protection Bureau. (2017). Financial well-being scale: Scale development technical report. CFPB. https://files.consumerfinance.gov/f/documents/201705_cfpb_financial-well-being-scale-technical-report.pdf',
    'Creswell, J. W., & Plano Clark, V. L. (2011). Designing and conducting mixed methods research. Sage Publications.',
    'Davenport, T. H., & Mittal, N. (2022). All-in on AI: How smart companies win big with artificial intelligence. Harvard Business Review Press.',
    'Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340.',
    'Deloitte. (2026). Agentic AI boosts wealth management: How AI agents enhance productivity. https://www.deloitte.com/us/en/insights/industry/financial-services/financial-services-industry-predictions/2026/agentic-ai-wealth-management-productivity.html',
    'Duhigg, C. (2012). The power of habit: Why we do what we do in life and business. Random House.',
    'Elenvo AI. (2026). How a financial health score is calculated. https://www.elenvo.ai/methodology',
    'Ernst & Young. (2026a). Nearly half of global consumers now use AI to guide savings and investment decisions. https://www.ey.com/en_gl/newsroom/2026/04/nearly-half-of-global-consumers-now-use-ai-to-guide-savings-and-investment-decisions',
    'Ernst & Young. (2026b). EY survey: Autonomous AI is no longer theoretical as adoption grows. https://www.ey.com/en_nl/newsroom/2026/03/ey-survey-autonomous-ai-is-no-longer-theoretical-as-adoption-grows-despite-ongoing-trust-concerns',
    'Financial Health Network. (2021). FinHealth Score® Toolkit. https://finhealthnetwork.org/tools/financial-health-score/',
    'Financial Health Network. (2026). From insight to impact: The next phase of financial health measurement. https://finhealthnetwork.org/research/from-insight-to-impact-the-next-phase-of-financial-health-measurement/',
    'Flores, C. A. R. (2025). Financial freedom of Filipinos in personal finance management. Pantao: The International Journal of the Humanities and Social Sciences, 4(1). https://pantaojournal.com/2025/01/27/v4-i1-7/',
    'GCash / Mynt. (2026). GCash launches country\'s first AI financial coach embedded in e-wallet [Press release]. PR Newswire. https://www.prnewswire.com/apac/news-releases/ph-fintech-gcash-launches-countrys-first-ai-financial-coach-embedded-in-e-wallet-to-strengthen-financial-literacy-302718569.html',
    'Hean, O., Saha, U., & Saha, B. (2025). Can AI help with your personal finances? Applied Economics. https://doi.org/10.1080/00036846.2025.2450384',
    'IBM. (2025). Agentic AI in financial services: Navigating innovation. https://www.ibm.com/think/insights/agentic-ai-financial-services-ethical-adoption',
    'Inquiro. (2024). Financial literacy in the Philippines: Key statistics. https://inquiro.ph/financial-literacy-in-the-philippines-2024-key-statistics/',
    'Juniper Research. (2026). Gamification in banking: How game mechanics drive financial behavior change [Research report].',
    'Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. Econometrica, 47(2), 263–292.',
    'Li, Y., et al. (2024). Large language models in finance (FinLLMs). Neural Computing and Applications. https://doi.org/10.1007/s00521-024-10495-6',
    'Li, Z., et al. (2024). A survey of large language models for financial applications. arXiv:2406.11903. https://arxiv.org/abs/2406.11903',
    'Liu, X., et al. (2023). FinGPT: Open-source financial large language models. arXiv:2306.06031. https://arxiv.org/abs/2306.06031',
    'Liu, Z., et al. (2025). Fin-R1: A large language model for financial reasoning through reinforcement learning. arXiv:2503.16252. https://arxiv.org/abs/2503.16252',
    'Meyll, T., et al. (2025). Spendception: The psychological impact of digital payments on consumer purchase behavior and impulse buying. Behavioral Sciences, 15(3), 387. https://doi.org/10.3390/bs15030387',
    'Mindfulsuite. (2026). The impact of expense tracking on financial behavior: How consistent logging reduces discretionary spending. https://mindfulsuite.com/blog/expense-tracking-financial-behavior',
    'NielsenIQ. (2026). The new financial reality: How Filipino consumers are spending, saving, and banking in 2026. https://nielseniq.com/global/en/insights/report/2026/the-new-financial-reality-how-filipino-consumers-are-spending-saving-and-banking-in-2026/',
    'Nielsen, J. (2006). Progressive disclosure. Nielsen Norman Group. https://www.nngroup.com/articles/progressive-disclosure/',
    'Philippine Statistics Authority. (2021). Family Income and Expenditure Survey (FIES) 2021. PSA. https://www.psa.gov.ph',
    'Philippine Statistics Authority. (2025). Philippine Digital Economy Satellite Account (PDESA) 2025. PSA. https://psa.gov.ph',
    'Plaid. (2026). State of intelligent finance report — Spring 2026. https://plaid.com/blog/state-of-intelligent-finance-report-spring-2026/',
    'Ramsey, D. (2003). Financial peace revisited. Viking.',
    'Rateweb. (2026). Financial health score — how it works. https://rateweb.co.za/financial-health',
    'Social Weather Stations. (2026, March). SWS financial inclusion survey: Philippines financial inclusion rises to 58%. Cited in CoinGeek (2026). https://coingeek.com/10-point-surge-pushes-philippines-financial-inclusion-to-58/',
    'Stefanov, T., Stefanova, M., & Varbanova, S. (2024). Personal finance management application. TEM Journal, 13(3), 2066–2075. https://doi.org/10.18421/TEM133-34',
    'Strivecloud. (2026). Fintech app gamification: Data shows 22% boost in saving habits. https://strivecloud.io/blog/mobile-app-gamification-fintech',
    'Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257–285.',
    'Thaler, R. H., & Sunstein, C. R. (2008). Nudge: Improving decisions about health, wealth, and happiness. Yale University Press.',
    'UNSGSA. (2021). Measuring financial health: A framework for practitioners. https://www.unsgsa.org',
    'Warren, E., & Tyagi, A. W. (2005). All your worth: The ultimate lifetime money plan. Free Press.',
    'World Economic Forum. (2024). How agentic AI will transform financial services. https://www.weforum.org/stories/2024/12/agentic-ai-financial-services-autonomy-efficiency-and-inclusion/',
    'Wu, S., et al. (2023). BloombergGPT: A large language model for finance. arXiv:2303.17564. https://arxiv.org/abs/2303.17564',
    'Xiao, Z., et al. (2026). FinRCA-Bench: Benchmarking evidence retrieval and reasoning for financial AI systems. arXiv:2608.18534. https://arxiv.org/abs/2608.18534',
    'Yang, H., et al. (2023). FinGPT: Democratizing internet-scale data for financial large language models. arXiv:2307.10485. https://arxiv.org/abs/2307.10485',
    'Yomio. (2026). YNAB alternatives: Which budget app actually works in 2026? https://yomio.app/en/blog/ynab-alternatives',
]

for ref in refs:
    add_apa(ref)

# ── SAVE ───────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
size_kb = OUT.stat().st_size / 1024
print(f'✓ Saved: {OUT.name}  ({size_kb:.0f} KB)')
print(f'  Full path: {OUT.resolve()}')
print(f'  Sections: Cover, FHS (full + lightweight + score labels + comparison + defense), '
      f'FMS (formula + classification + comparison + defense), '
      f'LLM Benchmarking (19 models, 4 tasks, multi-provider arch, defense), '
      f'App Comparison (22 apps, leads, gaps), References (49 APA)')
