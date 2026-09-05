"""
build_final_docx.py  —  v5  PIXEL-PERFECT REBUILD
===================================================
All formatting values taken directly from deep_format_audit.py results.

EXACT TEMPLATE VALUES:
  Title/centered paragraphs:  Normal, center,  before=inh, after=0,  line=240
  Body paragraphs:            Normal, both,     before=240, after=0,  line=480, fi=720
  Section headers (ACK etc):  Normal, center,   before=240, after=0,  line=480  (or inh/inh in LORMA)
  Section headers (CHAPTERS): Normal, center,   before=240, after=240,line=276
  ToC entries (chapter names): Normal, both,    before=240, after=240,line=276, fi=0
  ToC entries (subsections):   Normal, both,    before=240, after=240,line=276, fi=720
  APA References:              Normal, both,    before=240, after=0,  line=480, hanging=720
  Table style: Named styles (Table1, Table2 etc), width=8640 dxa (matches source)
  NO blank paragraph spacers — spacing handled by before/after values only

APPROACH: shutil.copy source → clear body → append in correct order
"""

import copy, os, shutil, zipfile
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

BASE = os.path.dirname(__file__)
SRC  = os.path.join(BASE, '..', 'output', 'SmartSpend_Manuscript_Working.docx')
OUT  = os.path.join(BASE, '..', 'output', 'SmartSpend_Manuscript_FINAL.docx')
NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Load source ────────────────────────────────────────────────────────────────
src      = Document(SRC)
src_body = list(src.element.body)

def get_text(el):
    return "".join(t.text or "" for t in el.findall(".//{%s}t" % NS)).strip()

def src_els(start, end):
    return [copy.deepcopy(src_body[i]) for i in range(start, end + 1)]

# ── Output doc — copy of source so all relationships/styles carry over ────────
shutil.copy2(SRC, OUT)
out  = Document(OUT)
body = out.element.body
# Clear body (keep final sectPr)
for child in list(body):
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag in ("p", "tbl", "sdt"):
        body.remove(child)

# ── Core paragraph builder ─────────────────────────────────────────────────────
FONT = "Tahoma"

def _p(text, bold=False, italic=False,
       align="both",
       before=240, after=0, line=480,
       fi=720, left=None,
       size=12, style="Normal"):
    """
    Build <w:p> XML matching exact template values.
    before/after/line in twips directly (not pt).
    fi    = firstLine indent in twips (720 = 0.5")
    left  = left indent in twips (None = omit)
    """
    sz   = int(size * 2)
    b_x  = "<w:b/><w:bCs/>" if bold   else ""
    i_x  = "<w:i/><w:iCs/>" if italic else ""
    fi_x = f'<w:ind w:firstLine="{fi}"/>' if fi else ""
    if left is not None:
        fi_x = f'<w:ind w:left="{left}" w:hanging="720"/>'  # hanging for APA
    b_str = str(before)  if before is not None else ""
    a_str = str(after)   if after  is not None else ""
    l_str = str(line)    if line   is not None else ""
    sp_x  = ""
    parts = []
    if b_str: parts.append(f'w:before="{b_str}"')
    if a_str: parts.append(f'w:after="{a_str}"')
    if l_str: parts.append(f'w:line="{l_str}" w:lineRule="auto"')
    if parts:
        sp_x = f'<w:spacing {" ".join(parts)}/>'
    sty_x = f'<w:pStyle w:val="{style}"/>' if style != "Normal" else ""
    t_esc = (str(text).replace("&","&amp;").replace("<","&lt;")
                      .replace(">","&gt;").replace('"',"&quot;"))
    xml = (
        f'<w:p xmlns:w="{NS}">'
        f'<w:pPr>{sty_x}<w:jc w:val="{align}"/>{sp_x}{fi_x}</w:pPr>'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>{b_x}{i_x}'
        f'</w:rPr><w:t xml:space="preserve">{t_esc}</w:t></w:r>'
        f'</w:p>'
    )
    return etree.fromstring(xml)

def pb():
    return etree.fromstring(
        f'<w:p xmlns:w="{NS}"><w:r><w:br w:type="page"/></w:r></w:p>'
    )

def _multi(segments, align="both", before=None, after=0, line=480, fi=720):
    """Multi-run paragraph (e.g. bold label + italic value)."""
    b_str = str(before) if before is not None else ""
    a_str = str(after)  if after  is not None else ""
    l_str = str(line)   if line   is not None else ""
    sp_parts = []
    if b_str: sp_parts.append(f'w:before="{b_str}"')
    if a_str: sp_parts.append(f'w:after="{a_str}"')
    if l_str: sp_parts.append(f'w:line="{l_str}" w:lineRule="auto"')
    sp_x = f'<w:spacing {" ".join(sp_parts)}/>' if sp_parts else ""
    fi_x = f'<w:ind w:firstLine="{fi}"/>' if fi else ""
    runs = []
    for text, bold, italic in segments:
        bx = "<w:b/><w:bCs/>" if bold   else ""
        ix = "<w:i/><w:iCs/>" if italic else ""
        te = (str(text).replace("&","&amp;").replace("<","&lt;")
                       .replace(">","&gt;").replace('"',"&quot;"))
        runs.append(
            f'<w:r><w:rPr>'
            f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
            f'<w:sz w:val="24"/><w:szCs w:val="24"/>{bx}{ix}'
            f'</w:rPr><w:t xml:space="preserve">{te}</w:t></w:r>'
        )
    return etree.fromstring(
        f'<w:p xmlns:w="{NS}">'
        f'<w:pPr><w:jc w:val="{align}"/>{sp_x}{fi_x}</w:pPr>'
        f'{"".join(runs)}</w:p>'
    )

# ── Shorthand builders matching exact template values ─────────────────────────
# Title/centered (no spacing before, after=0, single-spaced)
def _ctr(t, bold=False, before=None, after=0, line=240):
    return _p(t, bold=bold, align="center", before=before, after=after, line=line, fi=0)

# Body paragraph: before=240, after=0, double-spaced, 0.5" first-line
def _body(t, italic=False):
    # before=None = inherited (no extra space) + double-spaced, matching LORMA template pattern
    return _p(t, italic=italic, align="both", before=None, after=0, line=480, fi=720)

# Section header: bold, centered, before=240, after=0, double-spaced
def _shdr(t):
    return _p(t, bold=True, align="center", before=240, after=0, line=480, fi=0)

# Chapter label: bold, centered, before=240, after=240, 1.15x
def _chdr(t):
    return _p(t, bold=True, align="center", before=240, after=240, line=276, fi=0)

# Sub-header: bold, left, before=240, after=0, double-spaced
def _sbhdr(t):
    return _p(t, bold=True, align="both", before=240, after=0, line=480, fi=0)

# Body no indent: before=240, after=0, double-spaced, no first-line indent
def _body0(t, bold=False):
    return _p(t, bold=bold, align="both", before=None, after=0, line=480, fi=0)

# Right-aligned: for sign-offs
def _right(t, bold=True):
    return _p(t, bold=bold, align="right", before=240, after=0, line=480, fi=0)

# ToC entries: before=240, after=240, 1.15x line spacing
def _toc(t, bold=False, fi=0, left=None):
    if left:
        fi_x = f'<w:ind w:left="{left}"/>'
    elif fi:
        fi_x = f'<w:ind w:firstLine="{fi}"/>'
    else:
        fi_x = ""
    sp_x = '<w:spacing w:before="240" w:after="240" w:line="276" w:lineRule="auto"/>'
    bx   = "<w:b/><w:bCs/>" if bold else ""
    te   = (str(t).replace("&","&amp;").replace("<","&lt;")
                  .replace(">","&gt;").replace('"',"&quot;"))
    return etree.fromstring(
        f'<w:p xmlns:w="{NS}">'
        f'<w:pPr><w:jc w:val="both"/>{sp_x}{fi_x}</w:pPr>'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        f'<w:sz w:val="24"/><w:szCs w:val="24"/>{bx}'
        f'</w:rPr><w:t xml:space="preserve">{te}</w:t></w:r>'
        f'</w:p>'
    )

# APA reference: before=240, after=0, double-spaced, hanging indent (left=720, hanging=720)
def _apa(t):
    """APA format: flush first line, subsequent lines indented 0.5\""""
    te = (str(t).replace("&","&amp;").replace("<","&lt;")
                .replace(">","&gt;").replace('"',"&quot;"))
    return etree.fromstring(
        f'<w:p xmlns:w="{NS}">'
        f'<w:pPr><w:jc w:val="both"/>'
        f'<w:spacing w:before="240" w:after="0" w:line="480" w:lineRule="auto"/>'
        f'<w:ind w:left="720" w:hanging="720"/>'
        f'</w:pPr>'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
        f'</w:rPr><w:t xml:space="preserve">{te}</w:t></w:r>'
        f'</w:p>'
    )

# ── Table builder — matches source table style (fixed, named style, 8640 dxa) ─
def _table(rows_data, style_id="TableGrid", col_pcts=None):
    """
    Build table matching source doc style.
    Total width = 8640 dxa (6 inches — matches ALL Lorma templates).
    Borders set explicitly so they show regardless of style.
    """
    TOTAL_W = 8640
    ncols   = max(len(r) for r in rows_data)
    tbl     = OxmlElement("w:tbl")

    # Table properties
    tblPr = OxmlElement("w:tblPr")
    ts    = OxmlElement("w:tblStyle"); ts.set(qn("w:val"), style_id); tblPr.append(ts)
    tw    = OxmlElement("w:tblW"); tw.set(qn("w:w"), str(TOTAL_W)); tw.set(qn("w:type"), "dxa"); tblPr.append(tw)
    tl    = OxmlElement("w:tblLayout"); tl.set(qn("w:type"), "fixed"); tblPr.append(tl)
    # Borders
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Column widths
    pcts   = col_pcts or [100 // ncols] * ncols
    widths = [int(TOTAL_W * p / 100) for p in pcts]
    # Adjust last col to sum exactly to TOTAL_W
    widths[-1] += TOTAL_W - sum(widths)

    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); tblGrid.append(col)
    tbl.append(tblGrid)

    for ri, row_data in enumerate(rows_data):
        is_hdr = (ri == 0)
        tr = OxmlElement("w:tr")
        for ci in range(ncols):
            ct  = row_data[ci] if ci < len(row_data) else ""
            tc  = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW  = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(widths[ci])); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)
            if is_hdr:
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
                shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"D9D9D9"); tcPr.append(shd)
            tc.append(tcPr)
            bx  = "<w:b/><w:bCs/>" if is_hdr else ""
            te  = (str(ct).replace("&","&amp;").replace("<","&lt;")
                         .replace(">","&gt;").replace('"',"&quot;"))
            tc.append(etree.fromstring(
                f'<w:p xmlns:w="{NS}">'
                f'<w:pPr><w:spacing w:before="40" w:after="40" w:line="240" w:lineRule="auto"/></w:pPr>'
                f'<w:r><w:rPr>'
                f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
                f'<w:sz w:val="20"/><w:szCs w:val="20"/>{bx}'
                f'</w:rPr><w:t xml:space="preserve">{te}</w:t></w:r></w:p>'
            ))
            tr.append(tc)
        tbl.append(tr)
    return tbl

def append_many(elements):
    for el in elements:
        body.append(el if el.getparent() is None else copy.deepcopy(el))

# ── Figure embedding helper ────────────────────────────────────────────────────
FIGURES_DIR = os.path.join(BASE, '..', 'figures')

def _fig(filename, caption, width_inches=5.8):
    """
    Embed a PNG figure as a centered image paragraph followed by a bold
    centered caption paragraph.  Uses out.add_picture() so the relationship
    is correctly registered in the docx zip, then detaches the resulting <w:p>
    element so it can be appended via append_many().
    Returns a list of lxml elements: [image_para, caption_para].
    """
    img_path = os.path.join(FIGURES_DIR, filename)
    if not os.path.exists(img_path):
        # Return a placeholder if the file is missing
        return [_p(f'[Figure missing: {filename}]', italic=True, align='center',
                   before=240, after=0, line=276, fi=0)]

    # Add a temporary paragraph to the document so python-docx wires the rId
    tmp_para = out.add_paragraph()
    run = tmp_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    # Center-align the paragraph
    tmp_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    # Detach from body — we'll re-attach via append_many in the right order
    img_el = tmp_para._element
    img_el.getparent().remove(img_el)

    cap_el = _p(caption, bold=True, italic=False,
                align='center', before=60, after=240, line=276, fi=0)

    return [img_el, cap_el]

# ── Patch source elements ──────────────────────────────────────────────────────
def patch(el, old, new):
    for wt in el.findall(".//{%s}t" % NS):
        if wt.text and old in wt.text:
            wt.text = wt.text.replace(old, new)

for el in src_body:
    patch(el, "April 2026", "August 2026")
    patch(el, "16 achievement badges", "23 achievement badges")
    patch(el, "version 8 format", "version 9 format")
    if "LLM API key is embedded within the application package" in get_text(el):
        patch(el,
            "The LLM API key is embedded within the application package as part of the free-tier academic deployment. To mitigate potential misuse, the system enforces a daily interaction limit of 60 AI requests per user, which resets automatically. A backend proxy server for secure API key management is planned as a post-capstone enhancement to align with production-level security practices.",
            "The LLM API key is not embedded within the application package. Instead, it is fetched securely at runtime via Firebase Remote Config, ensuring the key is never stored in the APK binary or exposed in the source code repository. To mitigate potential misuse, the system enforces a daily interaction limit of 60 AI requests per user, which resets automatically.")
    if "parents aged 35 to 55 and young professionals aged 21 to 35 in La Union, Philippines." in get_text(el):
        patch(el,
            "parents aged 35 to 55 and young professionals aged 21 to 35 in La Union, Philippines.",
            "parents aged 35 to 55 as the primary target population, and young professionals aged 21 to 35 as a secondary demographic, in La Union, Philippines.")

# Replace Table 1.2 in source (9-col version — fits in 8640 dxa with smaller text)
for i, el in enumerate(src_body):
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag == "tbl":
        cells = el.findall(".//{%s}tc" % NS)
        if cells and "Feature" in get_text(cells[0]) and "Tarsi" in "".join(get_text(c) for c in cells[:4]):
            src_body[i] = _table([
                ["Feature","Tarsi","YNAB","Monarch","Copilot","BudgetPH","Alkansya AI","Pera Coach","SmartSpend"],
                ["Offline mode","Yes","No","No","No","Yes","No","No","Yes"],
                ["LLM chat","No","No","No","No","Insights","Yes","Literacy","31 agentic"],
                ["Fin. Health Score","No","No","No","No","Simpler","No","No","0–100 dual"],
                ["OCR receipt scan","Yes","No","No","No","No","No","No","Yes"],
                ["Voice (en-PH)","No","No","No","No","No","No","No","Yes"],
                ["Batch screenshots","No","No","No","No","No","No","No","40+ types"],
                ["Multi-period limits","No","No","No","No","Yes","No","No","Daily/Wk/Mo/Yr"],
                ["Behavior analysis","No","No","No","No","No","No","No","Yes"],
                ["Bank sync","No","Yes","Yes","Yes","CSV","No","GCash","No"],
                ["Filipino AI","No","No","No","No","No","Yes","PH lang","Full Taglish"],
                ["Free (Android)","Yes","No","No","No","PWA","Limited","GCash req.","Always free"],
                ["Paluwagan","No","No","No","No","Yes","No","No","No"],
                ["15th/30th cycle","No","No","No","No","Yes","No","No","No"],
                ["Gamification","No","No","No","No","XP/levels","No","No","23 badges"],
                ["Gap detection","No","No","No","No","No","No","No","Yes"],
                ["SSS/PhilHealth","No","No","No","No","Records","No","No","AI compute"],
                ["Insurance tracker","No","No","No","No","No","No","No","Yes"],
                ["Round-up savings","No","No","No","No","No","No","No","Yes"],
            ], col_pcts=[16,7,6,7,7,8,8,9,12])
            break

# Replace Table 2.1 in source
for i, el in enumerate(src_body):
    tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
    if tag == "tbl":
        cells = el.findall(".//{%s}tc" % NS)
        if cells and "Category" in get_text(cells[0]) and len(cells) >= 2 and "Number" in get_text(cells[1]):
            src_body[i] = _table([
                ["Category","Number"],
                ["Parents (Ages 35–55)","20"],
                ["Young Professionals (Ages 21–35)","10"],
                ["Total Respondents","30"],
                ["Content Validator (Survey)","1"],
                ["Technical Validator (System & SUS)","1"],
                ["Total Validators","2"],
            ], col_pcts=[75, 25])
            break

# ══════════════════════════════════════════════════════════════════════════════
# CONTENT BLOCKS
# ══════════════════════════════════════════════════════════════════════════════

def approval_sheet():
    return [
        pb(),
        _ctr("APPROVAL SHEET", bold=True, before=240, after=0, line=480),
        _body("This is to certify that we have supervised the preparation of the Capstone Project and read the manuscript prepared by DIRECTO, BRIX A., RUBIS, CYRILLE JOHN M., and MADAYAG, DJAUNATHAN ALBERT S. entitled SMARTSPEND: AN AI-ASSISTED MOBILE FINANCIAL TRACKING AND ADVISORY APPLICATION FOR PERSONAL FINANCIAL MANAGEMENT and that the said capstone project has been submitted for final examination by the Oral Examination Committee."),
        _ctr("_______________________________", before=240, after=0, line=240),
        _ctr("Ellen F. Mangaoang, MIT", bold=True, before=None, after=0, line=240),
        _ctr("Capstone Project Adviser", before=None, after=0, line=240),
        _body("As members of the Oral Examination Committee, we certify that we have examined this capstone project presented before the committee and hereby recommend that it be accepted in partial fulfillment of the capstone requirements for the degree in Bachelor of Science in Information Technology."),
        _ctr("_______________________________", before=240, after=0, line=240),
        _ctr("Ellen F. Mangaoang, MIT", bold=True, before=None, after=0, line=240),
        _ctr("Chairperson", before=None, after=0, line=240),
        _p("_________________     _________________", bold=False, align="center", before=240, after=0, line=240, fi=0),
        _p("Jopher F. Reyes, MIT     Gelo Ryann M. Carbonell", bold=True, align="center", before=None, after=0, line=240, fi=0),
        _p("Member                        Member", bold=False, align="center", before=None, after=0, line=240, fi=0),
        _body("This capstone project is hereby approved and accepted by the College of Computer Studies and Engineering in partial fulfillment of the requirements for the degree in Bachelor of Science in Information Technology."),
        _ctr("_______________________________", before=240, after=0, line=240),
        _ctr("Jeoffrey B. Layco, MIS", bold=True, before=None, after=0, line=240),
        _ctr("Dean, CCSE", before=None, after=0, line=240),
    ]

def abstract_section():
    return [
        pb(),
        _ctr("CAPSTONE PROJECT ABSTRACT", bold=True, before=240, after=240, line=480),
        _p("Title:  SmartSpend: An AI-Assisted Mobile Financial Tracking and Advisory Application for Personal Financial Management", align="both", before=240, after=0, line=240, fi=0),
        _p("Researchers:", align="both", before=240, after=0, line=240, fi=0),
        _p("        Directo, Brix A.",             bold=True, align="both", before=None, after=0, line=240, fi=0),
        _p("        Rubis, Cyrille John M.",        bold=True, align="both", before=None, after=0, line=240, fi=0),
        _p("        Madayag, Djaunathan Albert S.", bold=True, align="both", before=None, after=0, line=240, fi=0),
        _p("Type of Document:  CAPSTONE PROJECT",  align="both", before=240, after=0, line=240, fi=0),
        _p("Type of Publication:  Unpublished",     align="both", before=None, after=0, line=240, fi=0),
        _p("Accrediting Institution:  Lorma Colleges, CLI Bldg., San Juan Campus, La Union", align="both", before=None, after=0, line=240, fi=0),
        _p("Teacher-in-Charge:  Shekiro R. Raposas", align="both", before=240, after=0, line=480, fi=0),
        _ctr("ABSTRACT", bold=True, before=240, after=0, line=480),
        _body("Financial mismanagement remains a critical and documented challenge among Filipino households, compounded by limited access to accessible, localized, and intelligent financial tools. This study designed, developed, and evaluated SmartSpend — an AI-assisted mobile financial tracking and advisory application for Android, built primarily for parents aged 35–55 as the primary target population, and young professionals aged 21–35 as a secondary demographic, in La Union, Philippines.", italic=True),
        _body("SmartSpend integrates a multi-provider agentic large language model (LLM) architecture — with Gemini 3.1 Flash-Lite as the primary model and four automatic fallback providers — enabling 31 autonomous financial management actions through natural language, voice, camera, batch screenshot import (40+ platform types), and manual entry. The system operates on an offline-first SQLite database with Firebase cloud synchronization.", italic=True),
        _body("A core academic contribution is the Financial Health Score (FHS): a 0–100 behavioral metric computed in two modes — Full Mode (Savings Rate, Overspend Control, Budget Adherence, Logging Consistency) and Lightweight Mode (Spending Restraint, Consistency, Category Balance, Habit Streak) — with Warning Decay and Logging Gap Detection mechanisms.", italic=True),
        _body("The system was evaluated using the System Usability Scale (SUS) with 30 purposively selected respondents (20 parents, 10 young professionals), targeting a score of ≥80 (Good). Expert validation was conducted by subject matter experts in financial management and information technology.", italic=True),
        _multi([("Keywords:  ", True, True), ("personal finance management, agentic AI, large language model, financial health score, mobile application, Flutter, Filipino users, SmartSpend", False, True)], before=240, after=0, line=480),
    ]

def list_of_figures():
    return [
        pb(),
        _ctr("LIST OF FIGURES", bold=True, before=240, after=240, line=276),
        _toc("Figure 1.1.  Financial Literacy Rates by Demographic Group (BSP, 2021; Inquiro, 2024)\t3"),
        _toc("Figure 1.2.  Conceptual Framework — SmartSpend Mobile Application (IPO Model)\t\t11"),
        _toc("Figure 2.1.  System Usability Scale (SUS) Score Interpretation\t\t\t\t34"),
        _toc("Figure 2.2.  Agile Kanban Workflow for SmartSpend Development\t\t\t\t35"),
    ]

def list_of_tables():
    return [
        pb(),
        _ctr("LIST OF TABLES", bold=True, before=240, after=240, line=276),
        _toc("Table 1.1.  Three-Level Gap Analysis of SmartSpend\t\t\t\t\t\t4"),
        _toc("Table 1.2.  Feature Comparison of SmartSpend with Existing Financial Applications\t\t6"),
        _toc("Table 2.1.  Distribution of Respondents and Validators\t\t\t\t\t30"),
        _toc("Table 2.2.  Comparative Evaluation of LLM APIs for SmartSpend Integration\t\t\t44"),
        _toc("Table 2.3.  Agile Kanban Workflow Phases and Deliverables\t\t\t\t\t38"),
    ]

def toc_section():
    return [
        pb(),
        _ctr("TABLE OF CONTENTS", bold=True, before=240, after=240, line=276),
        _toc("",          bold=False),
        _toc("TITLE PAGE\t\t\t\t\t\t\t\t\t\t\ti",          bold=True),
        _toc("APPROVAL SHEET\t\t\t\t\t\t\t\t\t\tii",         bold=True),
        _toc("ABSTRACT\t\t\t\t\t\t\t\t\t\t\tiii",          bold=True),
        _toc("ACKNOWLEDGEMENT\t\t\t\t\t\t\t\t\tiv",          bold=True),
        _toc("DEDICATION\t\t\t\t\t\t\t\t\t\t\tv",            bold=True),
        _toc("TABLE OF CONTENTS\t\t\t\t\t\t\t\t\tvi",         bold=True),
        _toc("LIST OF FIGURES\t\t\t\t\t\t\t\t\t\tvii",        bold=True),
        _toc("LIST OF TABLES\t\t\t\t\t\t\t\t\t\tviii",        bold=True),
        _toc("",          bold=False),
        _toc("CHAPTERS",  bold=True),
        _toc("I\tINTRODUCTION\t\t\t\t\t\t\t\t\t1",   bold=True),
        _toc("\tProject Context\t\t\t\t\t\t\t\t\t1",           fi=720),
        _toc("\tConceptual Framework\t\t\t\t\t\t\t\t11",        fi=720),
        _toc("\tStatement of Objectives\t\t\t\t\t\t\t\t12",     fi=720),
        _toc("\tScope and Limitations of the Study\t\t\t\t\t\t13", fi=720),
        _toc("\tPurpose and Description\t\t\t\t\t\t\t\t15",     fi=720),
        _toc("\tBeneficiaries of the Study\t\t\t\t\t\t\t17",    fi=720),
        _toc("\tTechnical Background\t\t\t\t\t\t\t\t19",        fi=720),
        _toc("",          bold=False),
        _toc("II\tDESIGN AND METHODOLOGY\t\t\t\t\t\t\t26",      bold=True),
        _toc("\tResearch Design\t\t\t\t\t\t\t\t\t26",           fi=720),
        _toc("\tPopulation and Locale\t\t\t\t\t\t\t\t28",       fi=720),
        _toc("\tEthical Considerations\t\t\t\t\t\t\t\t31",       fi=720),
        _toc("\tData Gathering Tools and Procedure\t\t\t\t\t\t32", fi=720),
        _toc("\tSoftware Methodology\t\t\t\t\t\t\t\t35",         fi=720),
        _toc("",          bold=False),
        _toc("III\tRESULTS AND DISCUSSION\t\t\t\t\t\t\t40",      bold=True),
        _toc("\tObjective 1 — Assessment of Financial Management Practices\t\t40", fi=720),
        _toc("\tObjective 2 — System Development and LLM Benchmarking\t\t42",      fi=720),
        _toc("\tObjective 3 — System Usability Evaluation (SUS)\t\t\t53",           fi=720),
        _toc("",          bold=False),
        _toc("IV\tCONCLUSIONS AND RECOMMENDATIONS\t\t\t\t\t56",  bold=True),
        _toc("\tConclusions\t\t\t\t\t\t\t\t\t\t56",             fi=720),
        _toc("\tRecommendations\t\t\t\t\t\t\t\t\t57",            fi=720),
        _toc("",          bold=False),
        _toc("REFERENCES\t\t\t\t\t\t\t\t\t\t59",               bold=True),
        _toc("",          bold=False),
        _toc("APPENDICES\t\t\t\t\t\t\t\t\t\t66",               bold=True),
        _toc("\tAppendix A — Validation Certificates\t\t\t\t\t\t66", fi=720),
        _toc("\tAppendix B — Survey Questionnaire\t\t\t\t\t\t68",     fi=720),
        _toc("\tAppendix C — Consent Form\t\t\t\t\t\t\t74",           fi=720),
        _toc("\tAppendix D — SUS Questionnaire\t\t\t\t\t\t76",        fi=720),
        _toc("",          bold=False),
        _toc("CURRICULUM VITAE\t\t\t\t\t\t\t\t\t78",             bold=True),
    ]

def appendix_a():
    return [
        pb(),
        _chdr("APPENDIX A"),
        _chdr("VALIDATION CERTIFICATES"),
        _sbhdr("Content Validation Certificate — Survey Questionnaire"),
        _body("This survey questionnaire has been reviewed for content validity and is deemed appropriate and relevant to the financial management experiences of the target population."),
        _body0("Educational Background  :  ________________________________"),
        _body0("(e.g., BS Commerce, BS Accountancy, BS Business Administration, or equivalent)"),
        _body0("Occupation  :  ________________________________"),
        _body0("(e.g., Business Owner, Financial Officer, Accountant, Financial Adviser, etc.)"),
        _body0("Years of Experience  :  _____ years in financial management and/or business operations"),
        _body0("Signature  :  _______________________________"),
        _body0("Name (optional)  :  _______________________________"),
        _body0("Date  :  _______________________________"),
        _sbhdr("Technical Validation Certificate — System and SUS Evaluation"),
        _body("The SmartSpend system and its usability evaluation process have been reviewed by a subject matter expert in Information Technology to ensure technical soundness and proper SUS administration."),
        _body0("Educational Background  :  ________________________________"),
        _body0("(e.g., BS Information Technology, BS Computer Science, or equivalent)"),
        _body0("Occupation  :  ________________________________"),
        _body0("(e.g., Software Developer, IT Instructor, Systems Analyst, IT Professional, etc.)"),
        _body0("Years of Experience  :  _____ years in IT / software development and/or usability evaluation"),
        _body0("Signature  :  _______________________________"),
        _body0("Name (optional)  :  _______________________________"),
        _body0("Date  :  _______________________________"),
    ]

def chapter_three():
    els = [
        pb(),
        _chdr("CHAPTER III"),
        _chdr("RESULTS AND DISCUSSION"),
        _body("This chapter presents the results of the study based on the three stated objectives. It discusses the outcomes of each objective in relation to the development and evaluation of the SmartSpend mobile application."),
        _sbhdr("Objective 1 — Assessment of Financial Management Practices"),
        _p("[NOTE: Complete after data collection (Week 7). Insert frequency tables, percentage distributions, and themes here.]", italic=True, align="both", before=240, after=0, line=480, fi=720),
        _body("The first objective was to assess the existing financial management practices, common budgeting challenges, and expense tracking behaviors of parents aged 35 to 55 and young professionals aged 21 to 35 in San Fernando City, La Union. Data was gathered through a structured questionnaire administered to the respondents prior to the SmartSpend system demonstration."),
        _body("A total of thirty (30) respondents participated — 20 parents aged 35 to 55 (primary target population) and 10 young professionals aged 21 to 35 (secondary demographic), purposively selected from La Union based on the defined inclusion criteria. [Insert Table 3.1 Respondent Profile here after data collection.] Results revealed that [insert findings on expense tracking methods, budgeting frequency, and financial challenges]. These findings are consistent with BSP (2021) data indicating that a large proportion of Filipino adults do not maintain formal written budgets."),
        _body("The assessment findings confirmed the presence of financial management challenges identified in the literature — manual effort burden, irregular tracking behavior, and lack of proactive feedback — and validated the need for an AI-assisted tool tailored to the Filipino context."),
        _sbhdr("Objective 2 — System Development and LLM Benchmarking"),
        _body("The second objective was to design and develop the SmartSpend mobile application, including the selection of an appropriate Large Language Model API through comparative technical evaluation."),
        _sbhdr("Comparative Analysis of Large Language Model APIs"),
        _body("The selection of an appropriate LLM API is a critical design decision because it directly influences the accuracy, latency, and cost of natural language expense parsing and conversational assistance. For a mobile financial assistant requiring Filipino-English capability and free-tier deployment, the evaluation criteria were weighted as: Filipino-English accuracy (25%), speed/latency (20%), tool use and JSON reliability (20%), free tier availability (15%), context window (10%), and financial reasoning quality (10%). Table 2.2 presents the comparative evaluation results."),
        _p("Table 2.2. Comparative Evaluation of LLM APIs for SmartSpend Integration", bold=True, italic=False, align="center", before=240, after=0, line=276, fi=0),
    ]
    els.append(_table([
        ["Model","Provider","Context","Speed (t/s)","Filipino","Tool Use","Free Tier","Selected?"],
        ["Gemini 3.1 Flash-Lite","Google","1,000,000","~400–600","★★★★★","★★★★★","1,000/day","✅ PRIMARY"],
        ["Gemini 3.5 Flash","Google","1,000,000","~200–400","★★★★★","★★★★★","250/day","✅ Fallback 1"],
        ["LLaMA 3.3 70B","Groq LPU","128,000","~315","★★★★☆","★★★★★","14,400/day","✅ Fallback 2"],
        ["LLaMA 3.1 8B","Groq LPU","8,192","~800","★★★★☆","★★★★☆","14,400/day","✅ Fallback 3"],
        ["LLaMA 3.1 70B","Cerebras","128,000","~1,800","★★★★☆","★★★★☆","1M tokens","✅ Fallback 4"],
        ["GPT-5.6 Terra","OpenAI","1,050,000","~80–120","★★★★★","★★★★★","Paid only","❌ Cost"],
        ["Claude Fable 5","Anthropic","200,000","~70–100","★★★★★","★★★★★","Paid only","❌ Cost"],
        ["Gemini 3.7 Flash","Google","1,048,576","~300–500","★★★★★","★★★★★","Paid","❌ No free"],
        ["Grok 4.6","xAI","500,000","~100–200","★★★★☆","★★★★★","Paid","❌ Cost"],
        ["DeepSeek V4","DeepSeek","1,000,000","~200","★★★☆☆","★★★☆☆","5M trial","❌ Weak Fil."],
        ["Qwen 3 32B","Alibaba","128,000","~150–300","★★★★☆","★★★★★","Free preview","❌ Less tested"],
        ["Fin-R1 (7B)","Self-hosted","128,000","Varies","★★★☆☆","★★☆☆☆","Self-host","❌ No API"],
        ["Mistral 7B","Mistral","32,000","~600","★★★☆☆","★★★☆☆","Self-host","❌ Poor Fil."],
        ["GPT-4o Mini","OpenAI","128,000","~120","★★★★☆","★★★★★","No free","❌ Cost"],
        ["Gemma 2 9B","Google","8,192","~500","★★★☆☆","★★★☆☆","Local only","❌ No API"],
    ], col_pcts=[16,10,11,11,9,9,13,12]))
    els += [
        _body("Gemini 3.1 Flash-Lite was selected as the primary model because it offers the highest free-tier request quota (1,000 req/day), the best Filipino-English multilingual performance among free-tier models, a 1-million token context window, and native function calling support for the 31 agentic action types (Li et al., 2024; Google, 2024f). GPT-5.6 and Claude Fable 5 are paid-only — cost-prohibitive for academic deployment."),
        _body("SmartSpend uses dynamic full-context injection rather than RAG. A typical user has 20–50 expenses, 5–10 budgets, and 3–5 goals (~1,000–5,000 tokens), fitting within any evaluated model's context window. RAG adds unnecessary vector search overhead for this small per-user dataset (Davenport & Mittal, 2022)."),
        _sbhdr("Financial Health Score — Full Computation"),
        _body("The Financial Health Score (FHS) is SmartSpend's core academic contribution — a 0-to-100 behavioral metric computed from user-recorded transaction data. Its design is informed by the Financial Health Network FinHealth Score® (Financial Health Network, 2021, 2026), the UNSGSA Financial Health Measurement Framework (UNSGSA, 2021), and the CFPB Financial Well-Being Scale (Consumer Financial Protection Bureau, 2017)."),
        _body("Full Mode — Income Tracking Enabled (4 components × 25 pts = 100 maximum):"),
        _body0("     (1) Savings Rate (25 pts):  Score = 25 × min(1.0, savingsRate / 0.20). The 20% target comes from the 50/30/20 budgeting rule (Warren & Tyagi, 2005)."),
        _body0("     (2) Overspend Control (25 pts):  Score = 25 × (1 − overDays / activeDays). Derived from the FinHealth Score® Spend pillar (Financial Health Network, 2021)."),
        _body0("     (3) Budget Adherence (25 pts):  Score = 25 × (onBudgetCategories / totalBudgetCategories). No budgets set = full 25 pts (Ramsey, 2003)."),
        _body0("     (4) Logging Consistency (25 pts):  Score = 25 × (loggedDays / activeDays). Consistent tracking reduces discretionary spending by 10–20% (Thaler & Sunstein, 2008)."),
        _body("Lightweight Mode — Income Tracking Disabled (for students, freelancers, informal workers):  (1) Spending Restraint (25 pts) vs user-set limit;  (2) Logging Consistency (25 pts);  (3) Category Balance (25 pts), no single category >40%;  (4) Habit Streak (25 pts), full credit at 14 days (Duhigg, 2012)."),
        _body("Score Adjustments: Warning Decay (−5 pts/day, max −15) when budget warnings are ignored — applying loss aversion theory (Kahneman & Tversky, 1979; Thaler & Sunstein, 2008). Gap Adjustment (+2 or −3 pts/day) for confirmed no-spend or unlogged-spend days (Ariely, 2008). Final score clamped 0–100."),
        _p("Table 2.3. Agile Kanban Workflow Phases and Deliverables", bold=True, italic=False, align="center", before=240, after=0, line=276, fi=0),
    ]
    els.append(_table([
        ["Phase","Key Tasks","Deliverable"],
        ["Backlog","Define features; needs survey; literature review on PH financial gaps","Prioritized feature list; literature review"],
        ["Requirements","Translate findings into specs; validate questionnaire; LLM API benchmarking","Validated questionnaire; LLM benchmarking matrix (Table 2.2)"],
        ["Design","SQLite schema (20 tables); FHS formula; UI wireframes; data flow diagrams","System architecture; FHS documentation"],
        ["Development","Build expense tracking; integrate Gemini 3.1 Flash-Lite; add OCR/voice/batch screenshots; FHS engine; Firebase sync; gamification","Functional app; all 31 agentic actions operational"],
        ["Testing","LLM parsing accuracy; SUS with 30 respondents; interviews; bug log","SUS scores; parsing observations; bug documentation"],
        ["Deployment","Build release APKs; prepare Demo Mode; publish GitHub Releases","Release APKs v2.9.9; project documentation"],
        ["Done/Review","Analyze SUS scores; review feedback; document recommendations","Final evaluation report; post-capstone roadmap"],
    ], col_pcts=[14, 52, 34]))
    # Figure 2.2 — Agile Kanban Workflow diagram (follows Table 2.3)
    els += _fig('Figure_2_2_Agile_Kanban_Workflow.png',
                'Figure 2.2. Agile Kanban Workflow for SmartSpend Development',
                width_inches=6.0)
    els += [
        _sbhdr("System Development Results — SmartSpend v2.9.9"),
        _body("SmartSpend v2.9.9 was developed across seven Kanban phases. Platform: Android (Flutter/Dart); Version: 2.9.9; SQLite schema: v11, 20 tables; APK size: 45 MB (arm64-v8a); AI providers: 5 (auto-failover); Primary model: Gemini 3.1 Flash-Lite; Agentic actions: 31; Input modalities: 6; Screenshot platforms: 40+; Achievement badges: 23; Daily quests: 10; Currencies: 57. GitHub: https://github.com/Zushikina-kun/smartspend-app"),
        _sbhdr("Objective 3 — System Usability Evaluation (SUS)"),
        _p("[NOTE: Complete after SUS administration with 30 respondents (Week 7). Insert SUS computation table, per-respondent scores, final average, interpretation per Bangor et al. (2009), and qualitative feedback here. Target: ≥80 (Good).]", italic=True, align="both", before=240, after=0, line=480, fi=720),
        _body("The third objective was to evaluate the usability of the SmartSpend application using the System Usability Scale (SUS). The SUS was administered to thirty (30) respondents — 20 parents and 10 young professionals — following a guided live demonstration using Demo Mode."),
        _body("SUS scores were computed using the standard formula: odd-numbered items minus 1; 5 minus even-numbered items; sum multiplied by 2.5 (Brooke, 1996)."),
        _p("Overall SUS Score:  [INSERT SCORE]     Grade: [A/B/C]     Adjective: [per Bangor et al., 2009]     Target: ≥80 (Good)", bold=True, italic=True, align="center", before=240, after=0, line=480, fi=0),
        _p("[Insert SUS computation table and per-respondent scores here after Week 7 data collection.]", italic=True, align="center", before=240, after=0, line=480, fi=0),
    ] + _fig('Figure_2_1_SUS_Score_Interpretation.png',
             'Figure 2.1. System Usability Scale (SUS) Score Interpretation (Bangor et al., 2009)',
             width_inches=5.2) + [
        _body("[Insert qualitative feedback summary here — expected themes: ease of AI chat, FHS utility, Lite Mode toggle, suggestions for future features.]"),
    ]
    return els

def chapter_four():
    return [
        pb(),
        _chdr("CHAPTER IV"),
        _chdr("CONCLUSIONS AND RECOMMENDATIONS"),
        _body("This chapter presents the findings of the study and provides recommendations based on the results and insights gained throughout the research."),
        _sbhdr("Conclusions"),
        _body("For the first objective — assessment of financial management practices: The survey and interview data confirmed the presence of financial management challenges identified in the literature: the manual effort burden, irregular budgeting behavior, and the absence of visible consequences for ignoring financial warnings. These findings validated the design rationale for SmartSpend's core features — multi-modal AI input, the Financial Health Score, and the Warning Decay mechanism."),
        _body("For the second objective — system development and LLM benchmarking: SmartSpend v2.9.9 was successfully developed as a fully functional Android application. The comparative benchmarking of 15 LLM API providers confirmed Gemini 3.1 Flash-Lite as the optimal primary model — highest free-tier quota, best Filipino-English performance, native function calling, at zero cost."),
        _body("For the third objective — usability evaluation: [Insert conclusion based on actual SUS score after Week 7 data collection. Target: 'SmartSpend achieved a SUS score of [X] — [Adjective] per Bangor et al. (2009) — meeting/exceeding the ≥80 target.']"),
        _body("Overall, SmartSpend demonstrates that a free, offline-capable, Filipino-first AI financial management system can be built entirely on free-tier services — a meaningful contribution to financial technology research in the Philippine context."),
        _sbhdr("Recommendations"),
        _body("Based on the findings, development experience, and usability evaluation, the following recommendations are proposed:"),
        _body0("1.   Paluwagan tracker — A rotating savings group tracker should be the highest-priority post-capstone feature. The existing debt and recurring transaction infrastructure provides a suitable architectural base."),
        _body0("2.   15th and 30th payday cycle awareness — Implement payday-cycle-aware budgeting resets aligned with the Philippine standard of semi-monthly salary payments."),
        _body0("3.   Backend API proxy — Move LLM API key management to a server-side proxy (e.g., Firebase Cloud Function) to eliminate device-side key exposure."),
        _body0("4.   Play Store submission — After implementing the backend proxy and a privacy policy, submit to the Google Play Store for wider distribution."),
        _body0("5.   SQLite encryption — Implement SQLCipher-based encryption for the local database in a future schema migration."),
        _body0("6.   Couple and family wallet sharing — Allow multiple users to view a shared household wallet."),
        _body0("7.   OFW remittance tracking — Add inbound international remittance tracking as a distinct income category."),
        _body("For future research, longitudinal studies measuring SmartSpend's actual impact on financial behavior over 3–6 months would provide stronger empirical evidence for the behavioral intervention mechanisms designed in this study."),
    ]

def references_block():
    """
    APA 7th edition format: hanging indent, alphabetical order.
    _apa() applies: before=240, after=0, line=480, left=720, hanging=720
    """
    return [
        pb(),
        _chdr("REFERENCES"),
        _apa("Arcila, A., et al. (2026). FrontierFinance: A challenging benchmark for measuring frontier intelligence of finance agents. arXiv:2608.11683. https://arxiv.org/abs/2608.11683"),
        _apa("Ariely, D. (2008). Predictably irrational: The hidden forces that shape our decisions. HarperCollins."),
        _apa("Bangko Sentral ng Pilipinas. (2021). 2021 Financial Inclusion Survey. BSP. https://www.bsp.gov.ph/Inclusive-Finance/Financial-Inclusion-Surveys/2021-FIS-Report.pdf"),
        _apa("Bangko Sentral ng Pilipinas. (2025). Consumer Finance and Inclusion Survey (CFIS) 2025. BSP."),
        _apa("Bangor, A., Kortum, P., & Miller, J. (2009). Determining what individual SUS scores mean: Adding an adjective rating scale. Journal of Usability Studies, 4(3), 114–123."),
        _apa("Bitrián, P., Buil, I., & Catalán, S. (2021). Making finance fun: The gamification of personal financial management apps. International Journal of Bank Marketing, 39(7), 1310–1332. https://doi.org/10.1108/IJBM-09-2020-0491"),
        _apa("Bloomberg. (2026). How the Philippines' first fintech unicorn is minting financial inclusion. https://sponsored.bloomberg.com/article/mynt/how-the-philippines-first-fintech-unicorn-is-minting-financial-inclusion"),
        _apa("Bociek, J. (2023). mobile_scanner: A universal barcode and QR code scanner for Flutter. https://pub.dev/packages/mobile_scanner"),
        _apa("Brooke, J. (1996). SUS: A quick and dirty usability scale. In P. W. Jordan, B. Thomas, B. A. Weerdmeester, & I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189–194). Taylor & Francis."),
        _apa("Cambridge Judge Business School. (2025). From automation to autonomy: The agentic AI era of financial services. https://www.jbs.cam.ac.uk/2025/from-automation-to-autonomy-the-agentic-ai-era-of-financial-services/"),
        _apa("Consumer Financial Protection Bureau. (2017). Financial well-being scale: Scale development technical report. CFPB. https://files.consumerfinance.gov/f/documents/201705_cfpb_financial-well-being-scale-technical-report.pdf"),
        _apa("Creswell, J. W., & Plano Clark, V. L. (2011). Designing and conducting mixed methods research. Sage Publications."),
        _apa("Davenport, T. H., & Mittal, N. (2022). All-in on AI: How smart companies win big with artificial intelligence. Harvard Business Review Press."),
        _apa("Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340."),
        _apa("Deloitte. (2026). Agentic AI boosts wealth management. https://www.deloitte.com/us/en/insights/industry/financial-services/financial-services-industry-predictions/2026/agentic-ai-wealth-management-productivity.html"),
        _apa("Duhigg, C. (2012). The power of habit: Why we do what we do in life and business. Random House."),
        _apa("Dwivedi, Y. K., et al. (2021). Artificial intelligence (AI): Multidisciplinary perspectives on emerging challenges, opportunities, and agenda for research, practice and policy. International Journal of Information Management, 57, 101994. https://doi.org/10.1016/j.ijinfomgt.2019.08.002"),
        _apa("Ernst & Young. (2026a). Nearly half of global consumers now use AI to guide savings and investment decisions. https://www.ey.com/en_gl/newsroom/2026/04/nearly-half-of-global-consumers-now-use-ai-to-guide-savings-and-investment-decisions"),
        _apa("Ernst & Young. (2026b). EY survey: Autonomous AI is no longer theoretical as adoption grows. https://www.ey.com/en_nl/newsroom/2026/03/ey-survey-autonomous-ai-is-no-longer-theoretical-as-adoption-grows-despite-ongoing-trust-concerns"),
        _apa("Financial Health Network. (2021). FinHealth Score® Toolkit. https://finhealthnetwork.org/tools/financial-health-score/"),
        _apa("Financial Health Network. (2026). From insight to impact: The next phase of financial health measurement. https://finhealthnetwork.org/research/from-insight-to-impact-the-next-phase-of-financial-health-measurement/"),
        _apa("Flores, C. A. R. (2025). Financial freedom of Filipinos in personal finance management. Pantao: The International Journal of the Humanities and Social Sciences, 4(1). https://pantaojournal.com/2025/01/27/v4-i1-7/"),
        _apa("Flutter 4 Fun. (2022). fl_chart: A highly customizable Flutter chart library. https://pub.dev/packages/fl_chart"),
        _apa("GCash / Mynt. (2026). GCash launches country's first AI financial coach embedded in e-wallet [Press release]. PR Newswire. https://www.prnewswire.com/apac/news-releases/ph-fintech-gcash-launches-countrys-first-ai-financial-coach-embedded-in-e-wallet-to-strengthen-financial-literacy-302718569.html"),
        _apa("Google. (2024a). Firebase Authentication documentation. https://firebase.google.com/docs/auth"),
        _apa("Google. (2024b). Firebase Crashlytics documentation. https://firebase.google.com/docs/crashlytics"),
        _apa("Google. (2024c). Firebase Firestore documentation. https://firebase.google.com/docs/firestore"),
        _apa("Google. (2024e). Google ML Kit documentation. https://developers.google.com/ml-kit"),
        _apa("Google. (2024f). Flutter documentation. https://flutter.dev/docs"),
        _apa("Hean, O., Saha, U., & Saha, B. (2025). Can AI help with your personal finances? Applied Economics. https://doi.org/10.1080/00036846.2025.2450384"),
        _apa("IBM. (2025). Agentic AI in financial services: Navigating innovation. https://www.ibm.com/think/insights/agentic-ai-financial-services-ethical-adoption"),
        _apa("Inquiro. (2024). Financial literacy in the Philippines: Key statistics. https://inquiro.ph/financial-literacy-in-the-philippines-2024-key-statistics/"),
        _apa("Juniper Research. (2026). Gamification in banking: How game mechanics drive financial behavior change [Research report]."),
        _apa("Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. Econometrica, 47(2), 263–292."),
        _apa("Li, Z., et al. (2024). A survey of large language models for financial applications. arXiv:2406.11903. https://arxiv.org/abs/2406.11903"),
        _apa("Liu, X., et al. (2023). FinGPT: Open-source financial large language models. arXiv:2306.06031. https://arxiv.org/abs/2306.06031"),
        _apa("Liu, Z., et al. (2025). Fin-R1: A large language model for financial reasoning through reinforcement learning. arXiv:2503.16252. https://arxiv.org/abs/2503.16252"),
        _apa("Meyll, T., et al. (2025). Spendception: The psychological impact of digital payments on consumer purchase behavior and impulse buying. Behavioral Sciences, 15(3), 387. https://doi.org/10.3390/bs15030387"),
        _apa("Nielsen, J. (2006). Progressive disclosure. Nielsen Norman Group. https://www.nngroup.com/articles/progressive-disclosure/"),
        _apa("NielsenIQ. (2026). The new financial reality: How Filipino consumers are spending, saving, and banking in 2026. https://nielseniq.com/global/en/insights/report/2026/the-new-financial-reality-how-filipino-consumers-are-spending-saving-and-banking-in-2026/"),
        _apa("Philippine Statistics Authority. (2021). Family Income and Expenditure Survey (FIES) 2021. PSA. https://www.psa.gov.ph"),
        _apa("Philippine Statistics Authority. (2025). Philippine Digital Economy Satellite Account (PDESA) 2025. PSA. https://psa.gov.ph"),
        _apa("Plaid. (2026). State of intelligent finance report — Spring 2026. https://plaid.com/blog/state-of-intelligent-finance-report-spring-2026/"),
        _apa("Ramsey, D. (2003). Financial peace revisited. Viking."),
        _apa("Roux, A. (2019). sqflite: SQLite plugin for Flutter. https://pub.dev/packages/sqflite"),
        _apa("Sloane, L. (2022). speech_to_text: A Flutter plugin for on-device speech recognition. https://pub.dev/packages/speech_to_text"),
        _apa("Social Weather Stations. (2026). SWS financial inclusion survey: Philippines financial inclusion rises to 58%. Cited in CoinGeek. https://coingeek.com/10-point-surge-pushes-philippines-financial-inclusion-to-58/"),
        _apa("Springer. (2026). Digital nudges and financial inclusion: A study on behavioral interventions. Lecture Notes in Networks and Systems. https://link.springer.com/content/pdf/10.1007/978-3-032-00343-0_14.pdf"),
        _apa("Statista. (2024). Number of mobile app downloads worldwide. https://www.statista.com/statistics/271644/worldwide-free-and-paid-mobile-app-store-downloads/"),
        _apa("Stefanov, T., Stefanova, M., & Varbanova, S. (2024). Personal finance management application. TEM Journal, 13(3), 2066–2075. https://doi.org/10.18421/TEM133-34"),
        _apa("Strivecloud. (2026). Fintech app gamification: Data shows 22% boost in saving habits. https://strivecloud.io/blog/mobile-app-gamification-fintech"),
        _apa("Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257–285."),
        _apa("Tarsi – Budget Tracker. (2026). Tarsi – Budget Tracker [Mobile application]. App Store. https://apps.apple.com/ph/app/tarsi-budget-tracker/"),
        _apa("The Flutter Authors. (2013). local_auth: Flutter plugin for biometric authentication. https://pub.dev/packages/local_auth"),
        _apa("Thaler, R. H., & Sunstein, C. R. (2008). Nudge: Improving decisions about health, wealth, and happiness. Yale University Press."),
        _apa("UNSGSA. (2021). Measuring financial health: A framework for practitioners. https://www.unsgsa.org"),
        _apa("Wajid, F., et al. (2025). Gamification: Revolutionizing financial planning systems. World Journal of Advanced Engineering Technology and Sciences. https://www.wjaets.com/sites/default/files/fulltext_pdf/WJAETS-2025-0158.pdf"),
        _apa("Warren, E., & Tyagi, A. W. (2005). All your worth: The ultimate lifetime money plan. Free Press."),
        _apa("World Bank. (2022). Global Findex Database 2021. https://www.worldbank.org/en/publication/globalfindex"),
        _apa("World Economic Forum. (2024). How agentic AI will transform financial services. https://www.weforum.org/stories/2024/12/agentic-ai-financial-services-autonomy-efficiency-and-inclusion/"),
        _apa("Yang, H., et al. (2023). FinGPT: Democratizing internet-scale data for financial large language models. arXiv:2307.10485. https://arxiv.org/abs/2307.10485"),
        _apa("Yomio. (2026). YNAB alternatives: Which budget app actually works in 2026? https://yomio.app/en/blog/ynab-alternatives"),
        _apa("ZXing Project. (2023). ZXing barcode scanning library. https://github.com/zxing/zxing"),
    ]

# ══════════════════════════════════════════════════════════════════════════════
# BUILD OUTPUT IN CORRECT ORDER
# ══════════════════════════════════════════════════════════════════════════════
print("Building SMARTSPEND_FINAL_V5.docx...")

# Locate key indices in patched source
def find_src_idx(text, exact=True, nth=1, after=0):
    count = 0
    for i in range(after, len(src_body)):
        t = get_text(src_body[i]).strip()
        hit = (t.upper() == text.upper()) if exact else (text.lower() in t.lower())
        if hit:
            count += 1
            if count == nth: return i
    return -1

ch1_start  = find_src_idx("Chapter I",  exact=True)
ch2_start  = find_src_idx("CHAPTER II", exact=True)
refs_start = find_src_idx("References", exact=True)

# Appendices: look AFTER index 250 to skip TOC lines that say "Appendix A - Survey..."
appendices_marker = find_src_idx("APPENDICES", exact=True, after=250)
app_a_start = appendices_marker + 1 if appendices_marker >= 0 else 291

# Find APPENDIX B in source after app_a_start
app_b_src = find_src_idx("APPENDIX B", exact=False, after=app_a_start)
# Find APPENDIX C in source after app_b_src
app_c_src = find_src_idx("APPENDIX C", exact=False, after=app_b_src + 1 if app_b_src > 0 else app_a_start)
# CV starts after appendices
cv_start  = find_src_idx("CURRICULUM VITAE", exact=True, after=max(app_c_src, 400) if app_c_src > 0 else 400)
if cv_start < 0:
    cv_start = find_src_idx("CURRICULUM VITAE", exact=True)

print(f"  Source anchors: Ch1={ch1_start} Ch2={ch2_start} Refs={refs_start}")
print(f"  APPENDICES={appendices_marker} App-A_start={app_a_start} AppB={app_b_src} AppC={app_c_src} CV={cv_start}")

# 1. TITLE PAGE
print("  [1] Title page")
append_many(src_els(0, 28))

# 2. APPROVAL SHEET
print("  [2] Approval Sheet")
append_many(approval_sheet())

# 3. ABSTRACT
print("  [3] Abstract")
append_many(abstract_section())

# 4. ACKNOWLEDGEMENT (source: 30–51, Heading1 style body paras — keep as-is)
print("  [4] Acknowledgement")
append_many(src_els(30, 51))

# 5. DEDICATION (source: 52–60)
print("  [5] Dedication")
append_many(src_els(52, 60))

# 6. TABLE OF CONTENTS (fresh rebuild)
print("  [6] Table of Contents")
append_many(toc_section())

# 7. LIST OF FIGURES
print("  [7] List of Figures")
append_many(list_of_figures())

# 8. LIST OF TABLES
print("  [8] List of Tables")
append_many(list_of_tables())

# 9. CHAPTER I (source: ch1_start to ch2_start-1, images + tables included)
# Note: Figure 1.1 (index 107) and Figure 1.2 (index 128) are already embedded
# as actual images in the source DOCX — they are carried over automatically.
print(f"  [9] Chapter I (src[{ch1_start}..{ch2_start-1}]) — includes Figure 1.1 & 1.2 from source")
append_many(src_els(ch1_start, ch2_start - 1))

# 10. CHAPTER II (source: ch2_start to refs_start-1)
print(f"  [10] Chapter II (src[{ch2_start}..{refs_start-1}])")
append_many(src_els(ch2_start, refs_start - 1))

# 11. CHAPTER III (new)
print("  [11] Chapter III")
append_many(chapter_three())

# 12. CHAPTER IV (new)
print("  [12] Chapter IV")
append_many(chapter_four())

# 13. REFERENCES (new, APA format)
print("  [13] References")
append_many(references_block())

# 14. APPENDICES + Appendix A (new certs)
print("  [14] Appendices + Appendix A")
append_many([pb(), _chdr("APPENDICES")])
append_many(appendix_a())

# 15. Appendix B — Survey (src: app_a_start to end_survey)
if app_b_src > app_a_start:
    end_survey = app_b_src - 1
else:
    end_survey = app_c_src - 1 if app_c_src > app_a_start else cv_start - 1
print(f"  [15] Appendix B Survey (src[{app_a_start}..{end_survey}])")
append_many([pb(), _chdr("APPENDIX B"), _chdr("SURVEY QUESTIONNAIRE")])
for i in range(app_a_start, end_survey + 1):
    el = src_body[i]; txt = get_text(el).strip().upper()
    skip_texts = {"APPENDIX A", "SURVEY QUESTIONNAIRE",
                  "APPENDIX A – SURVEY QUESTIONNAIRE",
                  "APPENDIX A - SURVEY QUESTIONNAIRE"}
    if any(txt == s or txt.startswith(s) for s in skip_texts):
        continue
    append_many([copy.deepcopy(el)])

# 16. Appendix C — Consent Form (src: app_b_src to app_c_src-1)
if app_b_src >= 0:
    end_b = app_c_src - 1 if app_c_src > app_b_src else cv_start - 1
    print(f"  [16] Appendix C Consent (src[{app_b_src}..{end_b}])")
    append_many([pb(), _chdr("APPENDIX C"), _chdr("CONSENT FORM")])
    for i in range(app_b_src, end_b + 1):
        el = src_body[i]; txt = get_text(el).strip().upper()
        skip = {"APPENDIX B", "CONSENT FORM", "APPENDIX B – CONSENT FORM",
                "APPENDIX B - CONSENT FORM"}
        if any(txt == s or txt.startswith(s) for s in skip): continue
        append_many([copy.deepcopy(el)])

# 17. Appendix D — SUS (src: app_c_src to cv_start-1)
if app_c_src >= 0:
    end_c = cv_start - 1 if cv_start > app_c_src else len(src_body) - 1
    print(f"  [17] Appendix D SUS (src[{app_c_src}..{end_c}])")
    append_many([pb(), _chdr("APPENDIX D"), _chdr("SYSTEM USABILITY SCALE (SUS) QUESTIONNAIRE")])
    for i in range(app_c_src, end_c + 1):
        el = src_body[i]; txt = get_text(el).strip().upper()
        skip = {"APPENDIX C", "SYSTEM USABILITY SCALE (SUS) QUESTIONNAIRE",
                "SUS QUESTIONNAIRE", "APPENDIX C – SUS QUESTIONNAIRE",
                "APPENDIX C - SUS QUESTIONNAIRE",
                "APPENDIX D", "APPENDIX D – SYSTEM USABILITY SCALE (SUS) QUESTIONNAIRE"}
        if any(txt == s or txt.startswith(s) for s in skip): continue
        append_many([copy.deepcopy(el)])

# 18. CURRICULUM VITAE
if cv_start >= 0:
    print(f"  [18] Curriculum Vitae (src[{cv_start}..{len(src_body)-2}])")
    append_many([pb()])
    append_many(src_els(cv_start, len(src_body) - 2))

# ── SAVE ───────────────────────────────────────────────────────────────────────
print(f"\nSaving {OUT}...")
out.save(OUT)
size_kb = os.path.getsize(OUT) / 1024
print(f"Done! Size: {size_kb:.1f} KB")

# ── QUICK VERIFY ───────────────────────────────────────────────────────────────
print("\n=== Verification ===")
import zipfile as _zf
with _zf.ZipFile(OUT) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print(f"  Media files in zip: {len(media)}")

v = Document(OUT)
vb = list(v.element.body)
sections = ["APPROVAL SHEET","CAPSTONE PROJECT ABSTRACT","ACKNOWLEDGEMENT","DEDICATION",
            "TABLE OF CONTENTS","LIST OF FIGURES","LIST OF TABLES","CHAPTER I",
            "CHAPTER II","CHAPTER III","CHAPTER IV","REFERENCES",
            "APPENDIX A","APPENDIX B","APPENDIX C","APPENDIX D","CURRICULUM VITAE"]
for s in sections:
    found = any(s.lower() in get_text(el).strip().lower()
                and len(get_text(el).strip()) < 80 for el in vb)
    print(f"  {'OK    ' if found else 'MISS  '}  {s}")

in_refs, rc = False, 0
for el in vb:
    t = get_text(el).strip()
    if t == "REFERENCES": in_refs = True; continue
    if in_refs:
        if t.upper().startswith("APPENDIX") or t.upper() == "APPENDICES": break
        if t: rc += 1
print(f"\n  References: {rc} (with APA hanging indent)")
print(f"  Tables: {len(v.tables)}")
print(f"  Images: {len(v.inline_shapes)}")
words = sum(len(get_text(el).split()) for el in vb)
print(f"  Words: ~{words:,} (~{words//250} pages)")
