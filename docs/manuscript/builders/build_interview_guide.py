"""
build_interview_guide.py
Generates SmartSpend_Interview_Guide.docx — a standalone printable
interview guide for Djaunathan to use during respondent sessions.
Run:  python build_interview_guide.py
Out:  docs/manuscript/output/SmartSpend_Interview_Guide.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT  = Path(__file__).parent / '..' / 'output' / 'SmartSpend_Interview_Guide.docx'
FONT   = 'Tahoma'
MAROON = RGBColor(0x5C, 0x0E, 0x24)
BLUE   = RGBColor(0x2A, 0x4A, 0x7F)
DARK   = RGBColor(0x22, 0x22, 0x22)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x1A, 0x6B, 0x3A)

doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(1.25)
sec.right_margin  = Inches(1.25)
doc.styles['Normal'].font.name = FONT
doc.styles['Normal'].font.size = Pt(11)


def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)


def h(text, size=13, color=None, center=False, before=10, after=4, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = color or MAROON


def body(text, size=11, italic=False, before=2, after=3, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent: p.paragraph_format.first_line_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.italic = italic


def question(num, text, notes_lines=3):
    """Render a numbered question with answer lines."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'Q{num}.  ')
    r1.font.name = FONT; r1.font.size = Pt(11); r1.font.bold = True
    r1.font.color.rgb = MAROON
    r2 = p.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(11); r2.font.bold = True

    for _ in range(notes_lines):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(1)
        r3 = p2.add_run('_' * 90)
        r3.font.name = FONT; r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)


def section_box(title):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, '5C0E24')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(title)
    r.font.name = FONT; r.font.size = Pt(11.5)
    r.font.bold = True; r.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def tip(text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, 'FFF8DC')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f'💡  {text}')
    r.font.name = FONT; r.font.size = Pt(10); r.font.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ── HEADER ────────────────────────────────────────────────────────────────────
h('SmartSpend — Semi-Structured Interview Guide', size=15, center=True, before=0)
h('Assessment of Financial Management Practices — Objective 1',
  size=11, color=BLUE, center=True, before=0, after=2, bold=False)
body('Lorma Colleges CCSE · BSIT 4th Year · AY 2026–2027, 1st Semester · Lucid Frame',
     size=10, italic=True, before=0, after=2)
body('QA Lead / Interviewer: Djaunathan Albert S. Madayag', size=10, before=0, after=8)

# Respondent info block
tbl = doc.add_table(rows=2, cols=4)
tbl.style = 'Table Grid'
labels = [('Respondent Code', 'R___'), ('Date', '_____________'),
          ('Group', '☐ Parent  ☐ Young Prof.'), ('Time Start', '_____________')]
for ci, (label, blank) in enumerate(labels):
    shade_cell(tbl.rows[0].cells[ci], 'F0E8EC')
    p = tbl.rows[0].cells[ci].paragraphs[0]
    r = p.add_run(label); r.font.name = FONT; r.font.size = Pt(9); r.font.bold = True
    p2 = tbl.rows[1].cells[ci].paragraphs[0]
    r2 = p2.add_run(blank); r2.font.name = FONT; r2.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── INTRODUCTION SCRIPT ───────────────────────────────────────────────────────
h('Interviewer Introduction (read this to the respondent):', size=11, color=BLUE, before=6, after=3)
body(
    '"Good day po! We are BSIT students from Lorma Colleges conducting a capstone research study. '
    'We are developing SmartSpend, an AI-powered financial tracking app for Filipino users. '
    'We would like to ask you a few questions about how you manage your finances. '
    'There are no right or wrong answers — we just want to understand your experience. '
    'This will take about 5–7 minutes. Your answers are completely confidential and will '
    'only be used for our research. Is it okay to proceed?"',
    italic=True, size=10.5
)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── PART I ────────────────────────────────────────────────────────────────────
section_box('PART I — Expense Tracking Practices  (2 min)')
tip('Listen for: notebook, mental tracking, spreadsheet, no tracking at all. Note exact words used.')

question(1, 'How do you currently keep track of your daily expenses?', notes_lines=2)
question(2, 'How often do you monitor or review your budget?\n     (Daily / Weekly / Monthly / Rarely / Never)', notes_lines=1)
question(3, 'Do you find it difficult or inconvenient to track expenses manually? Why?', notes_lines=2)

# ── PART II ───────────────────────────────────────────────────────────────────
section_box('PART II — Budgeting Challenges  (2 min)')
tip('Listen for: overspending, forgetting, no system, irregular income. These map directly to SmartSpend\'s value proposition.')

question(4, 'What are your most common difficulties when managing your monthly expenses?', notes_lines=2)
question(5, 'Have you ever exceeded your budget without realizing it until later?', notes_lines=1)
question(6, 'Do you currently have debts or installment payments? Do these affect your monthly budget?', notes_lines=2)
question(7, 'Do you sometimes run short of money before your next salary or allowance?', notes_lines=1)

# ── PART III ──────────────────────────────────────────────────────────────────
section_box('PART III — Savings Behavior  (1 min)')
tip('Listen for: no savings, irregular, percentage-based. This validates the Savings Rate component of FHS.')

question(8, 'Do you regularly set aside savings from your income? Approximately how much (% or amount)?', notes_lines=2)
question(9, 'Do you have a specific savings goal right now? (e.g., emergency fund, appliance, education)', notes_lines=2)

# ── PART IV ───────────────────────────────────────────────────────────────────
section_box('PART IV — Technology and AI Openness  (1–2 min)')
tip('Listen for: GCash, YNAB, Excel, nothing. This validates the gap and SmartSpend\'s relevance.')

question(10, 'Do you currently use any mobile app or tool to manage finances? If yes, which one?', notes_lines=2)
question(11, 'Would you use a free AI-powered financial tracking app on Android? What features matter most to you?', notes_lines=2)

# ── CLOSING ───────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(6)
h('Closing Script:', size=11, color=BLUE, before=4, after=2)
body(
    '"Thank you so much for your time! Your responses are very helpful for our research. '
    'We will now show you the SmartSpend app, then after the demo, we\'ll ask you to fill '
    'out a short usability questionnaire. It will only take a few more minutes."',
    italic=True, size=10.5
)

# ── NOTES ─────────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(4)
section_box('INTERVIEWER NOTES (fill after the session)')

for label in ['Key themes / patterns observed:', 'Notable quotes (write exact words):', 'Other observations:']:
    body(label, size=10)
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run('_' * 90)
        r.font.name = FONT; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Time End: _____________    Total Duration: _____________    '
              'Respondent Consent Obtained: ☐ Yes')
r.font.name = FONT; r.font.size = Pt(10)

doc.save(str(OUT))
print(f'✓ Saved: {OUT.name}  ({OUT.stat().st_size//1024} KB)')
