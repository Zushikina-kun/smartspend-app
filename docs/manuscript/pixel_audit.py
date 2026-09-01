"""
pixel_audit.py
Extracts EXACT formatting values from every non-empty paragraph in all 3 docs.
Shows font name, size, bold, italic, alignment, space-before, space-after,
line spacing, first-line indent, left indent — at the RUN level (actual applied),
not just the style level.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
BASE = os.path.dirname(__file__)

DOCS = {
    "TEMPLATE_LORMA": os.path.join(BASE, "templates", "TEMPLATE_LORMA_ACCESS_PLUS.docx"),
    "TEMPLATE_HABSS":  os.path.join(BASE, "templates", "TEMPLATE_HABSS.docx"),
    "WORKING_ORIG":    os.path.join(BASE, "SMARTSPEND_CAPSTONE_WORKING.docx"),
    "UPDATED":         os.path.join(BASE, "SMARTSPEND_UPDATED_MANUSCRIPT.docx"),
}

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT:    "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER:  "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT:   "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    None:                        "default(JUSTIFY)",
}

def twips_to_pt(tw):
    """Convert twips to points (20 twips = 1pt)."""
    try: return round(int(tw) / 20, 1)
    except: return None

def twips_to_in(tw):
    """Convert twips to inches (1440 twips = 1 inch)."""
    try: return round(int(tw) / 1440, 3)
    except: return None

def halfpt_to_pt(hp):
    """Convert half-points to points."""
    try: return round(int(hp) / 2, 1)
    except: return None

def get_pPr(para):
    """Extract paragraph-level properties from XML."""
    pPr  = para._p.find(qn("w:pPr"))
    if pPr is None:
        return {}
    result = {}

    # Alignment
    jc = pPr.find(qn("w:jc"))
    result["align"] = jc.get(qn("w:val")) if jc is not None else "default"

    # Spacing
    sp = pPr.find(qn("w:spacing"))
    if sp is not None:
        result["sb"]   = twips_to_pt(sp.get(qn("w:before")))   # space before (pt)
        result["sa"]   = twips_to_pt(sp.get(qn("w:after")))    # space after (pt)
        lsr = sp.get(qn("w:line"))
        lst = sp.get(qn("w:lineRule"))
        if lsr:
            if lst == "exact":
                result["ls"] = f"exact:{twips_to_pt(lsr)}pt"
            elif lst == "atLeast":
                result["ls"] = f"atLeast:{twips_to_pt(lsr)}pt"
            else:
                # auto: value is 240 = single, 360 = 1.5x, 480 = double
                result["ls"] = f"auto:{int(lsr)}"   # 240=single 360=1.5x 480=double
    else:
        result["sb"] = result["sa"] = result["ls"] = None

    # Indent
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        result["fi"]   = twips_to_in(ind.get(qn("w:firstLine")))   # first-line indent (in)
        result["hang"] = twips_to_in(ind.get(qn("w:hanging")))     # hanging indent (in)
        result["left"] = twips_to_in(ind.get(qn("w:left")))        # left indent (in)
    else:
        result["fi"] = result["hang"] = result["left"] = None

    return result

def get_rPr(run):
    """Extract run-level properties from XML."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return {}
    result = {}

    # Font
    rf = rPr.find(qn("w:rFonts"))
    if rf is not None:
        result["font"] = rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi")) or "?"
    else:
        result["font"] = None

    # Size (half-points)
    sz = rPr.find(qn("w:sz"))
    result["size"] = halfpt_to_pt(sz.get(qn("w:val"))) if sz is not None else None

    # Bold / Italic
    result["bold"]   = rPr.find(qn("w:b"))   is not None
    result["italic"] = rPr.find(qn("w:i"))   is not None

    # Color
    cl = rPr.find(qn("w:color"))
    result["color"] = cl.get(qn("w:val")) if cl is not None else None

    return result

def analyse_doc(label, path):
    if not os.path.exists(path):
        print(f"\n{'='*60}\n  {label}: FILE NOT FOUND\n{'='*60}")
        return

    doc = Document(path)
    sec = doc.sections[0]

    print(f"\n{'='*70}")
    print(f"  {label}")
    print(f"{'='*70}")

    # Page setup
    print(f"\n[PAGE]  {sec.page_width.inches:.3f}\"W x {sec.page_height.inches:.3f}\"H"
          f"  margins: T={sec.top_margin.inches:.3f}\" B={sec.bottom_margin.inches:.3f}\""
          f" L={sec.left_margin.inches:.3f}\" R={sec.right_margin.inches:.3f}\"")

    # Normal style inherited values
    try:
        ns  = doc.styles["Normal"]
        nf  = ns.font
        npf = ns.paragraph_format
        print(f"[Normal style] font={nf.name or 'none'} size={nf.size.pt if nf.size else 'none'}pt"
              f"  align={npf.alignment}  sb={npf.space_before}  sa={npf.space_after}"
              f"  ls={npf.line_spacing}")
    except:
        print("[Normal style] not accessible")

    # Detailed paragraph-by-paragraph audit (first 80 non-empty)
    print(f"\n[PARAGRAPHS]")
    print(f"  {'#':>4}  {'STYLE':20}  {'FONT':12}  {'SZ':5}  {'B':1}  {'I':1}"
          f"  {'ALIGN':7}  {'SB':5}  {'SA':5}  {'LS':12}  {'FI':6}  {'LEFT':6}  TEXT")
    print(f"  {'-'*4}  {'-'*20}  {'-'*12}  {'-'*5}  {'-'}  {'-'}"
          f"  {'-'*7}  {'-'*5}  {'-'*5}  {'-'*12}  {'-'*6}  {'-'*6}  {'-'*40}")

    count = 0
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue

        pp   = get_pPr(p)
        # Get run props from first run
        rp   = {}
        for run in p.runs:
            rp = get_rPr(run)
            if rp:
                break

        sty  = p.style.name[:20]
        font = (rp.get("font") or "inh")[:12]
        sz   = str(rp.get("size") or "inh")[:5]
        bold = "B" if rp.get("bold") else " "
        ital = "I" if rp.get("italic") else " "
        aln  = (pp.get("align") or "default")[:7]
        sb   = str(pp.get("sb") or "inh")[:5]
        sa   = str(pp.get("sa") or "inh")[:5]
        ls   = str(pp.get("ls") or "inh")[:12]
        fi   = str(pp.get("fi") or "")[:6]
        lft  = str(pp.get("left") or "")[:6]
        txt  = p.text.strip()[:45]

        print(f"  {i:>4}  {sty:20}  {font:12}  {sz:5}  {bold}  {ital}"
              f"  {aln:7}  {sb:5}  {sa:5}  {ls:12}  {fi:6}  {lft:6}  {txt}")
        count += 1
        if count >= 80:
            break

# Run on all docs
for label, path in DOCS.items():
    analyse_doc(label, path)

# ── FOCUSED DIFF: compare spacing between WORKING_ORIG and UPDATED ──────────
print(f"\n{'='*70}")
print("  FOCUSED DIFF — WORKING_ORIG vs UPDATED (same-index paragraphs)")
print(f"{'='*70}")

orig = Document(DOCS["WORKING_ORIG"])
upd  = Document(DOCS["UPDATED"])

diffs = 0
orig_paras = [p for p in orig.paragraphs if p.text.strip()]
upd_paras  = [p for p in upd.paragraphs  if p.text.strip()]

# Compare first 50 common paragraphs by text similarity
for i in range(min(50, len(orig_paras), len(upd_paras))):
    op = orig_paras[i]; up = upd_paras[i]
    if op.text.strip()[:30] != up.text.strip()[:30]:
        continue  # skip misaligned

    opp = get_pPr(op); upp = get_pPr(up)
    orp = {}; urp = {}
    for r in op.runs:
        orp = get_rPr(r)
        if orp: break
    for r in up.runs:
        urp = get_rPr(r)
        if urp: break

    issues = []
    if opp.get("align") != upp.get("align"):
        issues.append(f"align: {opp.get('align')} -> {upp.get('align')}")
    if opp.get("sb") != upp.get("sb"):
        issues.append(f"sb: {opp.get('sb')} -> {upp.get('sb')}")
    if opp.get("sa") != upp.get("sa"):
        issues.append(f"sa: {opp.get('sa')} -> {upp.get('sa')}")
    if opp.get("ls") != upp.get("ls"):
        issues.append(f"ls: {opp.get('ls')} -> {upp.get('ls')}")
    if opp.get("fi") != upp.get("fi"):
        issues.append(f"fi: {opp.get('fi')} -> {upp.get('fi')}")
    if orp.get("font") != urp.get("font") and urp.get("font"):
        issues.append(f"font: {orp.get('font')} -> {urp.get('font')}")
    if orp.get("size") != urp.get("size") and urp.get("size"):
        issues.append(f"size: {orp.get('size')} -> {urp.get('size')}")

    if issues:
        diffs += 1
        print(f"  Para {i}: \"{op.text.strip()[:40]}\"")
        for iss in issues:
            print(f"    DIFF: {iss}")

if diffs == 0:
    print("  No formatting differences found in first 50 common paragraphs.")
print(f"\n  Total diffs in first 50: {diffs}")
