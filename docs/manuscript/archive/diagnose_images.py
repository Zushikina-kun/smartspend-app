"""
Diagnose exactly why images appear missing in SMARTSPEND_FINAL_V4.docx
even though inline_shapes count = 7.
Also inspect the source working docx to understand how images are stored.
"""
import zipfile, os, shutil
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

BASE = os.path.dirname(__file__)

def inspect_doc(fname):
    path = os.path.join(BASE, fname)
    print(f"\n{'='*60}")
    print(f"  {fname}")
    print(f"{'='*60}")

    # Check zip contents (docx is a zip)
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        rels   = [n for n in names if "rels" in n]
        print(f"\n  ZIP contents ({len(names)} files):")
        print(f"    Media files: {len(media)}")
        for m in media[:10]:
            info = z.getinfo(m)
            print(f"      {m}  ({info.file_size:,} bytes)")
        print(f"    Relationship files: {len(rels)}")
        for r in rels[:5]:
            print(f"      {r}")

    # Check relationships
    doc = Document(path)
    part = doc.part
    print(f"\n  Part relationships:")
    for rel in part.rels.values():
        if "image" in rel.reltype.lower():
            try:
                blob_size = len(rel.target_part.blob)
                print(f"    {rel.rId}: {rel.reltype.split('/')[-1]} — {rel.target_part.partname} ({blob_size:,} bytes)")
            except:
                print(f"    {rel.rId}: {rel.reltype.split('/')[-1]} — {rel.target_partname} (BROKEN)")

    # Check inline shapes
    print(f"\n  inline_shapes: {len(doc.inline_shapes)}")
    for i, sh in enumerate(doc.inline_shapes):
        try:
            w = sh.width.inches
            h = sh.height.inches
            # Try to get the rId
            blip = sh._inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            rid  = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed') if blip is not None else "?"
            print(f"    IMG{i+1}: {w:.2f}\" x {h:.2f}\"  rId={rid}")
        except Exception as e:
            print(f"    IMG{i+1}: ERROR — {e}")

    # Check the actual XML for drawing elements
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = list(doc.element.body)
    drawing_count = sum(1 for el in body
                        for _ in el.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"))
    blip_count    = sum(1 for el in body
                        for _ in el.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))
    print(f"\n  XML drawing (inline) elements: {drawing_count}")
    print(f"  XML blip (image ref) elements: {blip_count}")

    # Show where drawings are
    print(f"\n  Drawing locations in body:")
    for i, el in enumerate(body):
        drawings = el.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
        if drawings:
            surrounding = [el.find(".//{%s}t" % NS)]
            # get text of neighboring elements
            txt = "".join(t.text or "" for t in el.findall(".//{%s}t" % NS))[:40]
            for d in drawings:
                extent = d.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent")
                blip   = d.find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                cx = extent.get("cx") if extent is not None else "?"
                cy = extent.get("cy") if extent is not None else "?"
                rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed") if blip is not None else "?"
                print(f"    body[{i:3d}]: {int(cx)/914400:.2f}\" x {int(cy)/914400:.2f}\"  rId={rid}  text='{txt}'")

# Inspect both files
inspect_doc("SMARTSPEND_CAPSTONE_WORKING.docx")
inspect_doc("SMARTSPEND_FINAL_V4.docx")

# Check if source element deep copy preserves image relationships
print("\n" + "="*60)
print("  RELATIONSHIP COPY TEST")
print("="*60)
src_doc = Document(os.path.join(BASE, "SMARTSPEND_CAPSTONE_WORKING.docx"))
v4_doc  = Document(os.path.join(BASE, "SMARTSPEND_FINAL_V4.docx"))

src_image_rels = {rel.rId: rel for rel in src_doc.part.rels.values()
                  if "image" in rel.reltype.lower()}
v4_image_rels  = {rel.rId: rel for rel in v4_doc.part.rels.values()
                  if "image" in rel.reltype.lower()}

print(f"\n  Source image relationships: {len(src_image_rels)}")
for rid, rel in src_image_rels.items():
    print(f"    {rid}: {rel.target_part.partname}")

print(f"\n  V4 image relationships: {len(v4_image_rels)}")
for rid, rel in v4_image_rels.items():
    try:
        print(f"    {rid}: {rel.target_part.partname}")
    except:
        print(f"    {rid}: BROKEN REFERENCE")
