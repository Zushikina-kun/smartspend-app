"""
Deep cross-examination of:
1. SMARTSPEND_CAPSTONE_WORKING.docx  (our main doc)
2. TEMPLATE_LORMA_ACCESS_PLUS.docx  (Lorma Access+ reference)
3. TEMPLATE_HABSS.docx              (HABSS reference)

Outputs a full formatting audit: fonts, sizes, margins, styles,
paragraph spacing, heading levels, table styles, etc.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, sys

DOCS = {
    "OUR_DOC":   "SMARTSPEND_CAPSTONE_WORKING.docx",
    "TEMPLATE1": "templates/TEMPLATE_LORMA_ACCESS_PLUS.docx",
    "TEMPLATE2": "templates/TEMPLATE_HABSS.docx",
}

ALIGN_MAP = {
    None: "default/inherited",
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
}

def pt_or(val, default="inherited"):
    try:
        return f"{val.pt:.1f}pt"
    except:
        return default

def inch_or(val, default="inherited"):
    try:
        return f"{val.inches:.3f}\""
    except:
        return default

def analyze(label, path):
    if not os.path.exists(path):
        print(f"\n{'='*60}")
        print(f"  {label}: FILE NOT FOUND — {path}")
        return

    doc = Document(path)
    sec = doc.sections[0]

    print(f"\n{'='*60}")
    print(f"  {label}: {path}")
    print(f"{'='*60}")

    # ── Page ──────────────────────────────────────────────────
    print("\n[PAGE SETUP]")
    print(f"  Size:    {sec.page_width.inches:.2f}\" x {sec.page_height.inches:.2f}\"")
    print(f"  Top:     {inch_or(sec.top_margin)}")
    print(f"  Bottom:  {inch_or(sec.bottom_margin)}")
    print(f"  Left:    {inch_or(sec.left_margin)}")
    print(f"  Right:   {inch_or(sec.right_margin)}")
    print(f"  Header:  {inch_or(sec.header_distance)}")
    print(f"  Footer:  {inch_or(sec.footer_distance)}")

    # ── Default / Normal style ─────────────────────────────────
    print("\n[DEFAULT BODY STYLE]")
    try:
        ns = doc.styles["Normal"]
        f  = ns.font
        pf = ns.paragraph_format
        print(f"  Font:       {f.name or 'not set'}")
        print(f"  Size:       {pt_or(f.size)}")
        print(f"  Bold:       {f.bold}")
        print(f"  Italic:     {f.italic}")
        print(f"  Align:      {ALIGN_MAP.get(pf.alignment,'?')}")
        print(f"  SpaceBefore:{pt_or(pf.space_before)}")
        print(f"  SpaceAfter: {pt_or(pf.space_after)}")
        print(f"  LineSpacing:{pt_or(pf.line_spacing)}")
        print(f"  FirstIndent:{inch_or(pf.first_line_indent)}")
        print(f"  LeftIndent: {inch_or(pf.left_indent)}")
    except Exception as e:
        print(f"  Error: {e}")

    # ── All heading styles ─────────────────────────────────────
    print("\n[HEADING STYLES]")
    for hn in ["Heading 1","Heading 2","Heading 3","Heading 4"]:
        try:
            hs = doc.styles[hn]
            f  = hs.font
            pf = hs.paragraph_format
            print(f"  {hn}: font={f.name or 'inh'} {pt_or(f.size)} bold={f.bold} "
                  f"align={ALIGN_MAP.get(pf.alignment,'inh')} "
                  f"sb={pt_or(pf.space_before)} sa={pt_or(pf.space_after)}")
        except:
            print(f"  {hn}: not defined")

    # ── Styles used with counts ────────────────────────────────
    styles_used = {}
    for p in doc.paragraphs:
        sn = p.style.name
        styles_used[sn] = styles_used.get(sn, 0) + 1
    print("\n[STYLES USED]")
    for s, c in sorted(styles_used.items(), key=lambda x: -x[1]):
        print(f"  {c:5d}x  {s}")

    # ── Actual paragraph formatting sample ────────────────────
    print("\n[PARAGRAPH SAMPLE — first 60 non-empty, with actual run formatting]")
    count = 0
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        # Get actual run-level font (overrides style)
        fn, fsz, bold, ital = "inh", "inh", "inh", "inh"
        if p.runs:
            r    = p.runs[0]
            fn   = r.font.name or "inh"
            fsz  = pt_or(r.font.size)
            bold = str(r.font.bold) if r.font.bold is not None else "inh"
            ital = str(r.font.italic) if r.font.italic is not None else "inh"
        align = ALIGN_MAP.get(p.paragraph_format.alignment, "inh")
        sb    = pt_or(p.paragraph_format.space_before)
        sa    = pt_or(p.paragraph_format.space_after)
        ls    = pt_or(p.paragraph_format.line_spacing)
        fi    = inch_or(p.paragraph_format.first_line_indent)
        print(f"  {i:4d} [{p.style.name}] {fn}/{fsz}/B={bold}/I={ital} "
              f"align={align} sb={sb} sa={sa} ls={ls} fi={fi}")
        print(f"       \"{p.text[:70]}\"")
        count += 1
        if count >= 60:
            break

    # ── Tables ─────────────────────────────────────────────────
    print(f"\n[TABLES — {len(doc.tables)}]")
    for i, tbl in enumerate(doc.tables):
        style = tbl.style.name if tbl.style else "none"
        rows  = len(tbl.rows)
        cols  = len(tbl.columns) if tbl.columns else "?"
        hdr   = [c.text.strip()[:20] for c in tbl.rows[0].cells] if tbl.rows else []
        # Table cell font
        cfont, csize = "inh", "inh"
        if tbl.rows and tbl.rows[0].cells:
            cp = tbl.rows[0].cells[0].paragraphs
            if cp and cp[0].runs:
                cfont = cp[0].runs[0].font.name or "inh"
                csize = pt_or(cp[0].runs[0].font.size)
        print(f"  T{i+1}: style={style} {rows}r x {cols}c  "
              f"cellfont={cfont}/{csize}  hdr={hdr}")

    # ── Images ─────────────────────────────────────────────────
    print(f"\n[INLINE IMAGES — {len(doc.inline_shapes)}]")
    for i, sh in enumerate(doc.inline_shapes):
        try:
            w = f"{sh.width.inches:.2f}\""
            h = f"{sh.height.inches:.2f}\""
        except:
            w = h = "?"
        print(f"  IMG{i+1}: {w} x {h}  type={sh.type}")

    # ── References section detection ──────────────────────────
    print("\n[SECTION DETECTION]")
    markers = {
        "Chapter I": 0, "Chapter II": 0, "Chapter III": 0, "Chapter IV": 0,
        "References": 0, "Appendix": 0, "Curriculum Vitae": 0,
        "Abstract": 0, "Acknowledgement": 0,
    }
    for p in doc.paragraphs:
        for m in markers:
            if m.lower() in p.text.lower() and len(p.text.strip()) < 60:
                markers[m] += 1
    for m, c in markers.items():
        status = "✅ found" if c > 0 else "❌ MISSING"
        print(f"  {status}  {m} ({c} occurrences)")

# Run on all three
for label, path in DOCS.items():
    analyze(label, path)

print("\n" + "="*60)
print("CROSS-EXAMINATION COMPLETE")
print("="*60)
