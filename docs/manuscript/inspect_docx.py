from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import os, json

doc = Document('SMARTSPEND_CAPSTONE_WORKING.docx')

# Page setup
sec = doc.sections[0]
print("=== PAGE SETUP ===")
print(f"  Page:    {sec.page_width.inches:.2f}\" x {sec.page_height.inches:.2f}\"")
print(f"  Top:     {sec.top_margin.inches:.3f}\"")
print(f"  Bottom:  {sec.bottom_margin.inches:.3f}\"")
print(f"  Left:    {sec.left_margin.inches:.3f}\"")
print(f"  Right:   {sec.right_margin.inches:.3f}\"")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables:     {len(doc.tables)}")
print(f"  InlineShapes: {len(doc.inline_shapes)}")

# Styles used
styles_used = {}
for p in doc.paragraphs:
    sn = p.style.name
    styles_used[sn] = styles_used.get(sn, 0) + 1
print("\n=== STYLES USED ===")
for s, c in sorted(styles_used.items(), key=lambda x: -x[1]):
    print(f"  {c:4d}x  {s}")

# Detailed style definitions
print("\n=== KEY STYLE DEFINITIONS ===")
for sname in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Body Text", "List Paragraph"]:
    try:
        s = doc.styles[sname]
        f = s.font
        pf = s.paragraph_format
        sz = f"{f.size.pt:.0f}pt" if f.size else "inherited"
        fname = f.name or "inherited"
        align = str(pf.alignment) if pf.alignment else "inherited"
        sba = f"{pf.space_before.pt:.0f}" if pf.space_before else "0"
        saf = f"{pf.space_after.pt:.0f}" if pf.space_after else "0"
        print(f"  {sname}: font={fname} {sz} bold={f.bold} align={align} sb={sba} sa={saf}")
    except:
        print(f"  {sname}: not found")

# First 40 paragraphs with full detail
print("\n=== PARAGRAPH DETAILS (first 40 non-empty) ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip():
        continue
    run_info = ""
    if p.runs:
        r = p.runs[0]
        fn   = r.font.name or "inherited"
        fsz  = f"{r.font.size.pt:.0f}pt" if r.font.size else "inh"
        bold = "B" if r.font.bold else ""
        ital = "I" if r.font.italic else ""
        run_info = f" font={fn}/{fsz}{bold}{ital}"
    align = str(p.paragraph_format.alignment) if p.paragraph_format.alignment else ""
    print(f"  [{p.style.name}{run_info} {align}] {p.text[:80]}")
    count += 1
    if count >= 40:
        break

# Tables preview
print(f"\n=== TABLES ({len(doc.tables)}) ===")
for i, tbl in enumerate(doc.tables[:8]):
    if tbl.rows:
        row0 = [c.text.strip()[:25] for c in tbl.rows[0].cells]
        print(f"  T{i+1} ({len(tbl.rows)} rows x {len(tbl.columns)} cols): {row0}")

# Inline shapes (images)
print(f"\n=== IMAGES ({len(doc.inline_shapes)}) ===")
for i, sh in enumerate(doc.inline_shapes[:10]):
    w = sh.width.inches if sh.width else 0
    h = sh.height.inches if sh.height else 0
    print(f"  IMG{i+1}: {w:.2f}\" x {h:.2f}\" type={sh.type}")
